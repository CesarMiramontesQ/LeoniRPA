from fastapi import FastAPI, Request, Depends, Form, UploadFile, File, Query
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.gzip import GZipMiddleware
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import RedirectResponse, JSONResponse, FileResponse, StreamingResponse, Response
from datetime import datetime, timedelta
from contextlib import asynccontextmanager
from typing import Optional, List
from decimal import Decimal, InvalidOperation
from sqlalchemy.ext.asyncio import AsyncSession
from app.auth.router import router as auth_router, get_current_user, AuthenticationError, require_roles
from app.db.init_db import init_db
from app.db.base import get_db
from app.db.models import User, ExecutionStatus, MasterUnificadoVirtualOperacion, ClienteGrupo, Cliente, Venta
from app.db import crud
from app.bom.service import load_bom
from app.bom.schemas import LoadBomInput, LoadBomResponse
from app.bom.parse_sap_export import parse_sap_bom_txt
from app.core.config import settings
import threading
import asyncio
import subprocess
import json
from pathlib import Path
import sys
import socket
import platform
import os
import logging
import zipfile

try:
    from io import BytesIO
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
    from reportlab.lib import colors
    _REPORTLAB_AVAILABLE = True
except ImportError:
    _REPORTLAB_AVAILABLE = False

# Certificado C.O.: plantilla Excel + Office/LibreOffice para PDF
_CERT_PLANTILLA_XLSX = Path(__file__).resolve().parent / "Plantilla calificados.xlsx"
_CERT_FIRMA_IMG = Path(__file__).resolve().parent / "firma.png"
# Plantilla Word para componentes con ICR < 60% (no calificados USMCA)
_CERT_PLANTILLA_NO_CALIFICADOS_DOCX = Path(__file__).resolve().parent / "Plantilla no calificados.docx"


def _find_libreoffice() -> Optional[str]:
    """Devuelve la ruta del ejecutable de LibreOffice (soffice) o None si no se encuentra.
    En Mac busca en /Applications y ~/Applications (incl. LibreOffice X.X.app).
    """
    exe = os.environ.get("LIBREOFFICE_PATH") or os.environ.get("SOFFICE_PATH")
    if exe and Path(exe).exists():
        return exe
    if platform.system() == "Darwin":
        # Rutas fijas
        for p in (
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            "/Applications/OpenOffice.app/Contents/MacOS/soffice",
        ):
            if Path(p).exists():
                return p
        # Cualquier LibreOffice*.app en Applications (p. ej. "LibreOffice 24.2.0.app")
        for base in ("/Applications", Path.home() / "Applications"):
            base_path = Path(base)
            if not base_path.is_dir():
                continue
            for app in base_path.glob("LibreOffice*.app"):
                soffice = app / "Contents" / "MacOS" / "soffice"
                if soffice.exists():
                    return str(soffice)
            # OpenOffice
            soffice_oo = base_path / "OpenOffice.app" / "Contents" / "MacOS" / "soffice"
            if soffice_oo.exists():
                return str(soffice_oo)
    if platform.system() == "Windows":
        for p in (
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ):
            if Path(p).exists():
                return p
    import shutil
    return shutil.which("soffice") or shutil.which("libreoffice")


def _convert_xlsx_to_pdf_with_excel(xlsx_path: Path, pdf_path: Path, sheet_name: str = "C.O.") -> bool:
    """Convierte la hoja indicada del xlsx a PDF usando Microsoft Excel (solo Windows). Devuelve True si tuvo éxito."""
    if platform.system() != "Windows":
        return False
    try:
        import win32com.client
        xl = win32com.client.DispatchEx("Excel.Application")
        xl.Visible = False
        xl.DisplayAlerts = False
        wb = xl.Workbooks.Open(str(xlsx_path.resolve()))
        try:
            sheet = wb.Worksheets(sheet_name)
            sheet.ExportAsFixedFormat(0, str(pdf_path.resolve()))
        finally:
            wb.Close(SaveChanges=False)
        xl.Quit()
        return pdf_path.exists()
    except Exception:
        return False


def _convert_xlsx_to_pdf_with_excel_mac(xlsx_path: Path, pdf_path: Path) -> tuple[bool, str]:
    """Convierte el workbook xlsx a PDF usando Microsoft Excel en macOS (AppleScript).
    Devuelve (True, '') si tuvo éxito, (False, mensaje_error) si falló."""
    if platform.system() != "Darwin":
        return False, ""
    xlsx_posix = str(xlsx_path.resolve())
    pdf_posix = str(pdf_path.resolve())
    # Escribir script en archivo temporal para evitar problemas de escape con rutas
    import tempfile
    with tempfile.NamedTemporaryFile(mode="w", suffix=".scpt", delete=False, encoding="utf-8") as f:
        script_path = Path(f.name)
    try:
        # Script: abrir xlsx, guardar active workbook como PDF (ruta HFS para Excel)
        xlsx_esc = xlsx_posix.replace("\\", "\\\\").replace('"', '\\"')
        pdf_esc = pdf_posix.replace("\\", "\\\\").replace('"', '\\"')
        script_content = f'''set xlsxPath to "{xlsx_esc}"
set pdfPath to "{pdf_esc}"
set xlsxFile to POSIX file xlsxPath
set pdfFile to POSIX file pdfPath
tell application "Microsoft Excel"
    open xlsxFile
    delay 1
    tell active workbook
        save workbook as filename (pdfFile as text) file format PDF file format with overwrite
        close saving no
    end tell
end tell
'''
        scripts_to_try = [
            script_content,
            script_content.replace(
                "save workbook as filename (pdfFile as text) file format PDF file format with overwrite",
                "save workbook as filename pdfPath file format PDF file format with overwrite",
            ),
            script_content.replace(
                "save workbook as filename (pdfFile as text) file format PDF file format with overwrite",
                "save workbook as filename (pdfFile as text) file format PDF file format",
            ),
        ]
        last_err = "No se pudo generar el PDF con Excel."
        for app_name in ("Microsoft Excel", "Excel"):
            for script_variant in scripts_to_try:
                try:
                    script_text = script_variant.replace(
                        'tell application "Microsoft Excel"', f'tell application "{app_name}"'
                    )
                    script_path.write_text(script_text, encoding="utf-8")
                    result = subprocess.run(
                        ["osascript", str(script_path)],
                        capture_output=True,
                        timeout=90,
                        text=True,
                    )
                    if result.returncode == 0 and pdf_path.exists():
                        return True, ""
                    if result.returncode != 0:
                        last_err = (result.stderr or result.stdout or "").strip() or f"Exit code {result.returncode}"
                        logger.warning("Excel Mac (%s): %s", app_name, last_err[:300])
                except subprocess.TimeoutExpired:
                    last_err = "Timeout al ejecutar Excel (90 s). Cierre otros workbooks en Excel."
                    logger.warning("Excel Mac (%s): timeout", app_name)
                except Exception as e:
                    last_err = str(e)
                    logger.warning("Excel Mac (%s): %s", app_name, e)
        return False, last_err
    finally:
        try:
            script_path.unlink(missing_ok=True)
        except Exception:
            pass

# Para manejo de zonas horarias
try:
    from zoneinfo import ZoneInfo
except ImportError:
    # Fallback para Python < 3.9
    from backports.zoneinfo import ZoneInfo

# Meses en español para filtrar virtuales
MESES_VIRTUALES_ES = (
    "Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio",
    "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"
)

# Inicializar base de datos al iniciar
@asynccontextmanager
async def lifespan(app: FastAPI):
    """Inicializa la base de datos al iniciar la aplicación."""
    await init_db()
    yield

app = FastAPI(lifespan=lifespan)
templates = Jinja2Templates(directory="templates")
logger = logging.getLogger("bom_update")
logger.setLevel(logging.INFO)
if not logger.handlers:
    _handler = logging.StreamHandler()
    _handler.setLevel(logging.INFO)
    _handler.setFormatter(logging.Formatter("%(asctime)s | %(levelname)s | %(message)s"))
    logger.addHandler(_handler)
logger.propagate = False

# CORS: permite acceso desde otras máquinas en la red.
# allow_origins=["*"] admite cualquier origen; en producción se puede restringir.
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/static", StaticFiles(directory="static"), name="static")
app.add_middleware(GZipMiddleware)

# Exception handler para redirecciones de autenticación
@app.exception_handler(AuthenticationError)
async def authentication_exception_handler(request: Request, exc: AuthenticationError):
    """Redirige a login cuando hay error de autenticación."""
    return RedirectResponse(url="/auth/login", status_code=302)

# Incluir router de autenticación
app.include_router(auth_router)

# Redirigir raíz a dashboard
@app.get("/")
async def root():
    return RedirectResponse(url="/dashboard", status_code=302)


@app.get("/dashboard")
async def dashboard(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Dashboard principal - requiere autenticación. Incluye estadísticas, actividad reciente y pendientes."""
    from sqlalchemy import select, func, distinct
    from app.db.models import PaisOrigenMaterial

    mes_actual = MESES_VIRTUALES_ES[datetime.now().month - 1]
    stats = {
        "total_ventas": await crud.count_ventas(db),
        "total_compras": await crud.count_compras(db),
        "total_virtuales": await crud.count_master_unificado_virtuales(db, mes=mes_actual),
        "total_clientes": await crud.count_clientes(db),
        "total_proveedores": await crud.count_proveedores(db),
        "total_materiales": await crud.count_materiales(db),
        "total_precios_compra": await crud.count_precios_materiales(db),
    }
    # Materiales con país de origen pendiente
    r_pend = await db.execute(
        select(func.count(distinct(PaisOrigenMaterial.numero_material))).where(
            PaisOrigenMaterial.pais_origen == "Pendiente"
        )
    )
    materiales_pendientes = r_pend.scalar() or 0
    # Últimas ejecuciones de compras y ventas
    ejecuciones_compras = await crud.list_executions(db, limit=5)
    ejecuciones_ventas = await crud.list_sales_executions(db, limit=5)

    año_actual = datetime.now().year
    años_disponibles = list(range(año_actual, año_actual - 6, -1))
    return templates.TemplateResponse(
        "dashboard.html",
        {
            "request": request,
            "current_user": current_user,
            "active_page": "dashboard",
            "stats": stats,
            "materiales_pendientes": materiales_pendientes,
            "ejecuciones_compras": ejecuciones_compras,
            "ejecuciones_ventas": ejecuciones_ventas,
            "años_disponibles": años_disponibles,
            "mes_actual_nombre": mes_actual,
        },
    )


@app.get("/ventas")
async def ventas(request: Request, current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    """Página de ventas - requiere autenticación."""
    # Obtener el historial de ejecuciones para mostrar en la tabla
    executions = await crud.list_sales_executions(db, user_id=current_user.id, limit=50)
    
    return templates.TemplateResponse(
        "ventas.html",
        {
            "request": request,
            "active_page": "ventas",
            "current_user": current_user,
            "executions": executions
        }
    )


@app.get("/ventas-registros")
async def ventas_registros(request: Request, current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    """Página de registros de ventas - requiere autenticación."""
    total_ventas = await crud.count_ventas(db, only_with_sales_km=True)
    
    return templates.TemplateResponse(
        "ventas_registros.html",
        {
            "request": request,
            "active_page": "ventas_registros",
            "current_user": current_user,
            "total_ventas": total_ventas
        }
    )


@app.get("/analisis-icr")
async def analisis_icr(request: Request, current_user: User = Depends(get_current_user)):
    """Vista Análisis ICR: búsqueda por número de cliente y números de parte adquiridos (ventas con precios ICR)."""
    return templates.TemplateResponse(
        "analisis_icr.html",
        {
            "request": request,
            "active_page": "analisis_icr",
            "current_user": current_user,
        }
    )


@app.get("/api/analisis-icr/partes")
async def api_analisis_icr_partes(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    codigo_cliente: Optional[int] = None,
):
    """
    Devuelve el detalle Análisis ICR para un cliente: nombre, lista de partes con
    part_number, description, tariff_schedule, origin. Solo ventas con precios ICR válidos.
    """
    if codigo_cliente is None:
        return JSONResponse(
            {"error": "Falta el parámetro codigo_cliente"},
            status_code=400,
        )
    detalle = await crud.get_analisis_icr_detalle(db, codigo_cliente=codigo_cliente)
    # Original Components = números de parte con ICR > 60%; Non-Original = el resto
    partes_over_60 = detalle.get("partes_icr_over_60", 0)
    total_partes = detalle.get("total", 0)
    detalle["original_components"] = partes_over_60
    detalle["non_original_components"] = total_partes - partes_over_60
    detalle["conforming_part_numbers"] = detalle["total"]
    return JSONResponse(detalle)


# Datos fijos para el certificado C.O. (LEONI)
_CERT_CO_LEONI = {
    "certifier_company": "LEONI Cable S.A. de C.V.",
    "certifier_address_1": "Av. Río Conchos 9700",
    "certifier_address_2": "Parque Industrial Cuauhtemoc",
    "certifier_address_3": "Cd. Cuauhtemoc, Chih. Mexico CP 31543",
    "certifier_tax_id": "LCA981211QN0",
    "certifier_phone": "625 59 0 20 00",
    "certifier_email": "adalberto.blasco1@leonicables.com",
    "exporter_company": "LEONI Cable S.A. de C.V.",
    "exporter_address_1": "Av. Río Conchos 9700",
    "exporter_address_2": "Parque Industrial Cuauhtemoc",
    "exporter_address_3": "Cd. Cuauhtemoc, Chih. Mexico CP 31543",
    "exporter_tax_id": "LCA981211QN0",
    "exporter_phone": "625 59 0 20 00",
    "exporter_email": "obed.erives@leonicables.com",
    "producer_company": "LEONI Cable S.A. de C.V.",
    "producer_address_1": "Av. Río Conchos 9700",
    "producer_address_2": "Parque Industrial Cuauhtemoc",
    "producer_address_3": "Cd. Cuauhtemoc, Chih. Mexico CP 31543",
    "producer_tax_id": "LCA981211QN0",
    "producer_phone": "625 59 0 20 00",
    "producer_email": "obed.erives@leonicables.com",
    "certifier_name": "Adalberto Blasco",
    "certifier_title": "CUSTOMS MANAGER",
    "certifier_type": "EXPORTER",
}


def _render_certificado_co_pdf_reportlab(context: dict) -> bytes:
    """Genera el certificado C.O. en PDF con ReportLab cuando LibreOffice/Excel no están disponibles."""
    if not _REPORTLAB_AVAILABLE:
        raise RuntimeError("reportlab no está instalado; ejecute: pip install reportlab")
    out = BytesIO()
    doc = SimpleDocTemplate(out, pagesize=letter, leftMargin=0.7 * inch, rightMargin=0.7 * inch, topMargin=0.6 * inch, bottomMargin=0.6 * inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(name="CertTitle", parent=styles["Heading1"], fontSize=14, alignment=1, spaceAfter=2)
    sub_style = ParagraphStyle(name="CertSub", parent=styles["Normal"], fontSize=12, alignment=1, spaceAfter=1)
    body_style = ParagraphStyle(name="CertBody", parent=styles["Normal"], fontSize=9, spaceAfter=1)
    bold_style = ParagraphStyle(name="CertBold", parent=styles["Normal"], fontSize=10, fontName="Helvetica-Bold", alignment=1, spaceAfter=8)
    story = []
    story.append(Paragraph("United States Mexico Canada Agreement – USMCA", title_style))
    story.append(Paragraph("CERTIFICATION OF ORIGEN", sub_style))
    story.append(Paragraph("T-MEC", bold_style))
    story.append(Spacer(1, 6))
    certifier_block = "<b>1.- CERTIFIER, NAME, AND ADDRESS</b><br/>" + f"{context.get('certifier_company', '')}<br/>{context.get('certifier_address_1', '')}<br/>{context.get('certifier_address_2', '')}<br/>{context.get('certifier_address_3', '')}<br/>TAX IDENTIFICATION NUMBER: {context.get('certifier_tax_id', '')}<br/>TELEPHONE EMAIL<br/>" + f"{context.get('certifier_phone', '')} {context.get('certifier_email', '')}"
    exporter_block = "<b>2.- EXPORTER NAME, AND ADDRESS</b><br/>" + f"{context.get('exporter_company', '')}<br/>{context.get('exporter_address_1', '')}<br/>{context.get('exporter_address_2', '')}<br/>{context.get('exporter_address_3', '')}<br/>TAX IDENTIFICATION NUMBER: {context.get('exporter_tax_id', '')}<br/>TELEPHONE EMAIL<br/>" + f"{context.get('exporter_phone', '')} {context.get('exporter_email', '')}"
    t1 = Table([[Paragraph(certifier_block, body_style), Paragraph(exporter_block, body_style)]], colWidths=[3.25 * inch, 3.25 * inch])
    t1.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (1, 0), (1, 0), 20)]))
    story.append(t1)
    story.append(Spacer(1, 6))
    producer_block = "<b>3.- PRODUCER NAME, ADDRESS AND EMAIL</b><br/>" + f"{context.get('producer_company', '')}<br/>{context.get('producer_address_1', '')}<br/>{context.get('producer_address_2', '')}<br/>{context.get('producer_address_3', '')}<br/>TAX IDENTIFICATION NUMBER: {context.get('producer_tax_id', '')}<br/>TELEPHONE EMAIL<br/>" + f"{context.get('producer_phone', '')} {context.get('producer_email', '')}"
    importer_block = "<b>4.- IMPORTER NAME, ADDRESS AND EMAIL.</b><br/>" + f"<b>{context.get('cliente_nombre') or '—'}</b><br/>Customer # {context.get('codigo_cliente') or '—'}<br/>" + f"{context.get('importer_address') or 'VARIOUS (optional customer address)'}<br/>TAX IDENTIFICATION NUMBER:<br/>EMAIL<br/>"
    t2 = Table([[Paragraph(producer_block, body_style), Paragraph(importer_block, body_style)]], colWidths=[3.25 * inch, 3.25 * inch])
    t2.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (1, 0), (1, 0), 20)]))
    story.append(t2)
    story.append(Spacer(1, 6))
    partes = context.get("partes") or []
    data = [["5.", "DESCRIPTION OF GOOD (S)", "6. HS TARIFF CLASSIFICATION NUMBER", "7. CRITERION OF ORIGIN", "8. ORIGIN COUNTRY"]]
    for p in partes:
        desc = (p.get("part_number") or "—")
        if p.get("description"):
            d = (p.get("description") or "")[:80]
            desc += " — " + (d + "..." if len((p.get("description") or "")) > 80 else d)
        data.append(["", desc, p.get("tariff_schedule") or "SEE ATTACHED", "B", p.get("origin") or "MX"])
    if not partes:
        data.append(["", "—", "SEE ATTACHED", "—", "—"])
    t3 = Table(data, colWidths=[0.3 * inch, 2.3 * inch, 1.3 * inch, 1.0 * inch, 1.0 * inch])
    t3.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")), ("GRID", (0, 0), (-1, -1), 0.5, colors.black), ("FONTSIZE", (0, 0), (-1, -1), 9), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("LEFTPADDING", (0, 0), (-1, -1), 4), ("RIGHTPADDING", (0, 0), (-1, -1), 4)]))
    story.append(t3)
    story.append(Spacer(1, 6))
    story.append(Paragraph(f"<b>9. BLANKET PERIOD (MM/DD/YY)</b><br/>FROM: {context.get('blanket_period_from', '')}  TO: {context.get('blanket_period_to', '')}", body_style))
    story.append(Spacer(1, 6))
    story.append(Paragraph("I certify that the goods described in this document qualify as originating and the information contained in this document is true and accurate.", body_style))
    story.append(Paragraph("I assume responsibility for proving such representations and agree to maintain and present upon request or to make available during a verification visit, documentation necessary to support this certification.", body_style))
    story.append(Paragraph(f"This certification consists of {context.get('num_pages', 1)} page(s), including all attachments.", body_style))
    story.append(Spacer(1, 12))
    left_sig = "<b>CERTIFIER'S SIGNATURE</b><br/><br/><b>CERTIFIER'S NAME (PRINT OR TYPE)</b><br/>" + (context.get("certifier_name") or "") + "<br/><br/><b>DATE (MM/DD/YY)</b><br/>" + (context.get("certification_date") or "")
    right_sig = "<b>COMPANY NAME</b><br/>" + (context.get("certifier_company") or "") + "<br/><br/><b>CERTIFIER'S TITLE</b><br/>" + (context.get("certifier_title") or "") + "<br/><br/><b>CERTIFIER TYPE (IMPORTER, EXPORTER, PRODUCER)</b><br/>" + (context.get("certifier_type") or "")
    t4 = Table([[Paragraph(left_sig, body_style), Paragraph(right_sig, body_style)]], colWidths=[3.25 * inch, 3.25 * inch])
    t4.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (1, 0), (1, 0), 20)]))
    story.append(t4)
    doc.build(story)
    return out.getvalue()


def _render_certificado_co_xlsx(context: dict) -> bytes:
    """
    Genera el certificado C.O. en Excel: rellena la plantilla (hoja C.O.) y devuelve el archivo .xlsx en bytes.
    Solo crea el Excel; no convierte a PDF.
    """
    import shutil
    from io import BytesIO
    from openpyxl import load_workbook

    if not _CERT_PLANTILLA_XLSX.exists():
        raise FileNotFoundError(
            f"No se encontró la plantilla: {_CERT_PLANTILLA_XLSX}. Coloque 'Plantilla calificados.xlsx' en la raíz del proyecto."
        )

    with BytesIO() as buf:
        with open(_CERT_PLANTILLA_XLSX, "rb") as f:
            buf.write(f.read())
        buf.seek(0)
        wb = load_workbook(buf, read_only=False, data_only=False)

    if "C.O." not in wb.sheetnames:
        wb.close()
        raise ValueError("La plantilla no contiene la hoja 'C.O.'")
    ws = wb["C.O."]

    try:
        from openpyxl.cell.cell import MergedCell
    except ImportError:
        MergedCell = type(None)

    def _cell_to_write(ws, row: int, col: int):
        cell = ws.cell(row=row, column=col)
        if not isinstance(cell, MergedCell):
            return (row, col)
        for rng in ws.merged_cells.ranges:
            if rng.min_row <= row <= rng.max_row and rng.min_col <= col <= rng.max_col:
                return (rng.min_row, rng.min_col)
        return (row, col)

    def _set_cell(ws, row: int, col: int, value):
        r, c = _cell_to_write(ws, row, col)
        ws.cell(row=r, column=c).value = value

    def _ajustar_ancho_columnas(ws, columnas: tuple, fila_ini: int, fila_fin: int, ancho_min=8, ancho_max=55):
        """Ajusta el ancho de las columnas para que se muestre todo el texto (según contenido en fila_ini..fila_fin)."""
        try:
            from openpyxl.utils import get_column_letter
            for col_idx in columnas:
                max_len = ancho_min
                for row in range(fila_ini, fila_fin + 1):
                    val = ws.cell(row=row, column=col_idx).value
                    if val is not None:
                        max_len = max(max_len, len(str(val)))
                w = min(ancho_max, max(max_len * 1.23, ancho_min))
                ws.column_dimensions[get_column_letter(col_idx)].width = w
        except Exception as e:
            logger.warning("No se pudo ajustar ancho de columnas: %s", e)

    cliente_nombre = (context.get("cliente_nombre") or "").strip() or "VARIOUS (optional customer address)"
    _set_cell(ws, 16, 5, cliente_nombre)

    # En C.O. no se rellenan componentes (se deja como en plantilla, p. ej. SEE ATTACHED). Solo cliente y número de páginas.
    # La firma va ya en la plantilla; no se inserta aquí.

    # Número de páginas del certificado: C.O. (1) + C.O. 3 (2) = 2 cuando existe hoja C.O. 3.
    num_pages = int(context.get("num_pages") or (2 if "C.O. 3" in wb.sheetnames else 1))
    text_pages = "This certification consists of"
    text_suffix = "page(s), including all attachments."
    for row in range(1, min(ws.max_row + 1, 120)):
        for col in range(1, min(ws.max_column + 1, 30)):
            val = ws.cell(row=row, column=col).value
            if isinstance(val, str) and text_pages in val and "page(s)" in val:
                r, c = _cell_to_write(ws, row, col)
                ws.cell(row=r, column=c).value = f"{text_pages} {num_pages} {text_suffix}"
                break
        else:
            continue
        break

    # Hoja C.O. 3: tabla con Description, Customer Part Number, Leoni Part Number, Leoni Part Name, NUMBER (tariff), B, MX.
    # Filas 9-17 en plantilla (9 filas de formato fijo). Si hay más de 9 materiales, insertar filas a partir de la 18.
    if "C.O. 3" in wb.sheetnames:
        ws_co3 = wb["C.O. 3"]
        fila_ini_co3 = 9
        fila_fin_tabla_co3 = 17
        num_filas_co3 = 9
        columnas_co3 = (2, 3, 4, 5, 6, 8, 9)  # B=Description, C=Customer Part#, D=Leoni Part#, E=Leoni Part Name, F=NUMBER, H=B, I=origin
        partes_co3 = context.get("partes") or []
        num_partes_co3 = len(partes_co3)
        filas_extra_co3 = max(0, num_partes_co3 - num_filas_co3)
        if filas_extra_co3 > 0:
            fila_ins_co3 = fila_fin_tabla_co3 + 1
            ws_co3.insert_rows(fila_ins_co3, filas_extra_co3)
            for rng in list(ws_co3.merged_cells.ranges):
                if rng.min_row >= fila_ins_co3:
                    rng.shift(row_shift=filas_extra_co3, col_shift=0)
            # Merge F:G en cada fila nueva (igual que en la plantilla).
            for r in range(fila_ins_co3, fila_ins_co3 + filas_extra_co3):
                ws_co3.merge_cells(start_row=r, start_column=6, end_row=r, end_column=7)
            rd3 = ws_co3.row_dimensions
            altura_co3 = 15.0
            if fila_fin_tabla_co3 in rd3 and getattr(rd3[fila_fin_tabla_co3], "height", None) is not None:
                altura_co3 = rd3[fila_fin_tabla_co3].height
            for row_idx in sorted(rd3.keys(), reverse=True):
                if row_idx >= fila_ins_co3:
                    ref = rd3[row_idx]
                    new_row = row_idx + filas_extra_co3
                    if getattr(ref, "height", None) is not None:
                        rd3[new_row].height = ref.height
                    if getattr(ref, "hidden", None) is not None:
                        rd3[new_row].hidden = ref.hidden
                    del rd3[row_idx]
            for r in range(fila_ins_co3, fila_ins_co3 + filas_extra_co3):
                rd3[r].height = altura_co3
            from copy import copy
            # Copiar formato (bordes incluidos) de la fila de referencia a cada fila nueva; incluir F y G (6 y 7).
            columnas_con_formato_co3 = (2, 3, 4, 5, 6, 7, 8, 9)
            for r in range(fila_ins_co3, fila_ins_co3 + filas_extra_co3):
                for col in columnas_con_formato_co3:
                    src = ws_co3.cell(row=fila_fin_tabla_co3, column=col)
                    tgt = ws_co3.cell(row=r, column=col)
                    if hasattr(src, "_style") and getattr(src, "_style", None) is not None:
                        tgt._style = copy(src._style)
        # Limpiar filas de datos: 9-17 y las dos filas 18-19 de la plantilla. Si se insertaron filas, 18 y 19
        # quedaron desplazadas abajo (18+filas_extra_co3, 19+filas_extra_co3), hay que vaciarlas también.
        ultima_fila_a_limpiar = 19 + filas_extra_co3
        for row in range(fila_ini_co3, ultima_fila_a_limpiar + 1):
            for col in columnas_co3:
                _set_cell(ws_co3, row, col, "")
        for i, p in enumerate(partes_co3):
            row = fila_ini_co3 + i
            numero = (p.get("part_number") or "").strip()
            customer_part = (p.get("customer_part_number") or numero or "").strip()
            _set_cell(ws_co3, row, 2, "Electrical Cable")
            _set_cell(ws_co3, row, 3, customer_part)
            _set_cell(ws_co3, row, 4, numero)
            descripcion = (p.get("description") or "").strip()
            if descripcion:
                descripcion = descripcion[:120].strip()
            _set_cell(ws_co3, row, 5, descripcion or "")
            _set_cell(ws_co3, row, 6, (p.get("tariff_schedule") or "").strip() or "")
            _set_cell(ws_co3, row, 8, "B")
            _set_cell(ws_co3, row, 9, "MX")
        # Ajustar ancho de columnas en C.O. 3 para que se vea todo el texto.
        fila_fin_datos_co3 = fila_ini_co3 + max(num_partes_co3, 1) - 1
        _ajustar_ancho_columnas(ws_co3, (2, 3, 4, 5, 6, 7, 8, 9), 8, fila_fin_datos_co3)

    out = BytesIO()
    wb.save(out)
    wb.close()
    return out.getvalue()


def _render_certificado_co_pdf(context: dict) -> bytes:
    """
    Genera el PDF del certificado C.O.: primero genera el Excel con _render_certificado_co_xlsx
    y luego lo convierte a PDF con LibreOffice o Microsoft Excel.
    """
    import tempfile
    xlsx_bytes = _render_certificado_co_xlsx(context)
    libreoffice = _find_libreoffice()
    use_excel_win = platform.system() == "Windows"
    use_excel_mac = platform.system() == "Darwin"
    if not libreoffice and not use_excel_win and not use_excel_mac:
        raise RuntimeError(
            "Para generar el PDF instale LibreOffice o use Windows/Mac con Microsoft Excel."
        )

    with tempfile.TemporaryDirectory(prefix="cert_co_") as tmpdir:
        tmpdir_path = Path(tmpdir)
        tmp_xlsx = tmpdir_path / "certificado_co.xlsx"
        tmp_xlsx.write_bytes(xlsx_bytes)
        pdf_path = tmpdir_path / "certificado_co.pdf"

        if libreoffice:
            subprocess.run(
                [libreoffice, "--headless", "--convert-to", "pdf", "--outdir", str(tmpdir_path), str(tmp_xlsx)],
                capture_output=True,
                timeout=60,
                cwd=str(tmpdir_path),
            )
            if not pdf_path.exists():
                raise RuntimeError(
                    "LibreOffice no generó el PDF. Compruebe que LibreOffice esté instalado correctamente."
                )
            from pypdf import PdfReader, PdfWriter
            reader = PdfReader(pdf_path)
            writer = PdfWriter()
            # Incluir C.O. (pág 1) y C.O. 3 (pág 2) en el mismo PDF.
            for i in range(min(2, len(reader.pages))):
                writer.add_page(reader.pages[i])
            out = BytesIO()
            writer.write(out)
            return out.getvalue()
        if use_excel_win:
            # Exportar hoja C.O. y hoja C.O. 3 a PDFs temporales y unirlos en uno solo.
            from pypdf import PdfReader, PdfWriter
            pdf_co = tmpdir_path / "co.pdf"
            pdf_co3 = tmpdir_path / "co3.pdf"
            ok1 = _convert_xlsx_to_pdf_with_excel(tmp_xlsx, pdf_co, "C.O.")
            ok2 = _convert_xlsx_to_pdf_with_excel(tmp_xlsx, pdf_co3, "C.O. 3")
            if ok1:
                writer = PdfWriter()
                reader1 = PdfReader(pdf_co)
                writer.add_page(reader1.pages[0])
                if ok2 and pdf_co3.exists():
                    reader2 = PdfReader(pdf_co3)
                    if len(reader2.pages) > 0:
                        writer.add_page(reader2.pages[0])
                out = BytesIO()
                writer.write(out)
                return out.getvalue()
        if use_excel_mac:
            ok, err_msg = _convert_xlsx_to_pdf_with_excel_mac(tmp_xlsx, pdf_path)
            if ok:
                from pypdf import PdfReader, PdfWriter
                reader = PdfReader(pdf_path)
                writer = PdfWriter()
                # Excel en Mac guarda todo el libro: incluir C.O. (pág 1) y C.O. 3 (pág 2).
                for i in range(min(2, len(reader.pages))):
                    writer.add_page(reader.pages[i])
                out = BytesIO()
                writer.write(out)
                return out.getvalue()
            raise RuntimeError(
                "No se pudo generar el PDF con Microsoft Excel. "
                + (err_msg if err_msg else "")
                + " Compruebe que Excel esté instalado. Alternativa: instale LibreOffice."
            )
        if use_excel_win:
            raise RuntimeError(
                "No se pudo generar el PDF con Microsoft Excel. Alternativa: instale LibreOffice."
            )
        raise RuntimeError("No se pudo convertir la plantilla Excel a PDF.")


def _render_no_calificados_docx(context: dict) -> bytes:
    """
    Rellena la plantilla Word 'Plantilla no calificados.docx' con los componentes con ICR < 60%.
    Context debe tener: codigo_cliente, cliente_nombre, partes (lista con part_number, description,
    tariff_schedule, customer_part_number), blanket_year (opcional).
    Devuelve el .docx en bytes.
    """
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT

    if not _CERT_PLANTILLA_NO_CALIFICADOS_DOCX.exists():
        raise FileNotFoundError(
            f"No se encontró la plantilla: {_CERT_PLANTILLA_NO_CALIFICADOS_DOCX}. "
            "Coloque 'Plantilla no calificados.docx' en la raíz del proyecto."
        )
    doc = Document(str(_CERT_PLANTILLA_NO_CALIFICADOS_DOCX))
    codigo_cliente = context.get("codigo_cliente") or ""
    cliente_nombre = (context.get("cliente_nombre") or "").strip() or "—"
    partes = context.get("partes") or []
    ahora = context.get("_ahora") or datetime.now()
    blanket_year = context.get("blanket_year") or ahora.year

    # Reemplazar XXXXX y año en la primera tabla (filas 0 y 1: Customer Code, Customer Name, Blanket Year)
    tbl0 = doc.tables[0]
    xxxx_count = 0
    for row in tbl0.rows[:2]:
        for cell in row.cells:
            full = cell.text
            if "XXXXX" in full:
                xxxx_count += 1
                if xxxx_count == 1:
                    cell.text = full.replace("XXXXX", str(codigo_cliente))
                else:
                    cell.text = full.replace("XXXXX", cliente_nombre)
            if "2026" in full:
                cell.text = full.replace("2026", str(blanket_year))

    # Tabla: fila 2 = encabezado, filas 3+ = datos. Columnas: Leoni Part Number, Leoni Part Name, Customer Part Number (cross reference), Description, HTS, Comments
    num_cols = 6
    data_start_row = 3
    header_row = tbl0.rows[2]
    num_data_rows_template = max(0, len(tbl0.rows) - data_start_row)
    for i, p in enumerate(partes):
        row_idx = data_start_row + i
        part_number = (p.get("part_number") or "").strip()
        descripcion = (p.get("description") or "").strip() or "Electrical Cable"
        if len(descripcion) > 120:
            descripcion = descripcion[:120].strip()
        # Customer Part Number: viene del cross reference (customer_part_number); si no hay, se usa part_number
        customer_part = (p.get("customer_part_number") or part_number or "").strip()
        tariff = (p.get("tariff_schedule") or "").strip() or "8544.49"
        if row_idx < len(tbl0.rows):
            row = tbl0.rows[row_idx]
            cells = row.cells
            if len(cells) >= num_cols:
                cells[0].text = part_number
                cells[1].text = descripcion
                cells[2].text = customer_part
                cells[3].text = "Electrical Cable"
                cells[4].text = tariff
                cells[5].text = "Not USMCA Complaint"
        else:
            new_row = tbl0.add_row()
            vals = [part_number, descripcion, customer_part, "Electrical Cable", tariff, "Not USMCA Complaint"]
            for c in range(min(num_cols, len(new_row.cells))):
                new_row.cells[c].text = vals[c]

    # Vaciar filas de datos sobrantes si hay menos partes que filas plantilla
    for row_idx in range(data_start_row + len(partes), len(tbl0.rows)):
        row = tbl0.rows[row_idx]
        for cell in row.cells[:num_cols]:
            cell.text = ""

    # Tabla con letra más pequeña para que quepa en el documento y centrada
    tabla_font_pt = 8
    tbl0.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in tbl0.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(tabla_font_pt)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


def _render_no_calificados_pdf(context: dict) -> bytes:
    """Genera el PDF de la declaración no calificados: docx rellenado y conversión con LibreOffice."""
    import tempfile
    docx_bytes = _render_no_calificados_docx(context)
    libreoffice = _find_libreoffice()
    if not libreoffice:
        _mac = " (en Mac: brew install --cask libreoffice o descarga desde libreoffice.org). " if platform.system() == "Darwin" else " "
        raise RuntimeError(
            "Para generar el PDF hace falta LibreOffice (soffice)."
            + _mac
            + "Si está en otra ruta, defina la variable de entorno LIBREOFFICE_PATH. "
            "La descarga del .docx está disponible sin PDF."
        )
    with tempfile.TemporaryDirectory(prefix="no_calificados_") as tmpdir:
        tmpdir_path = Path(tmpdir)
        tmp_docx = tmpdir_path / "no_calificados.docx"
        tmp_docx.write_bytes(docx_bytes)
        pdf_path = tmpdir_path / "no_calificados.pdf"
        subprocess.run(
            [libreoffice, "--headless", "--convert-to", "pdf", "--outdir", str(tmpdir_path), str(tmp_docx)],
            capture_output=True,
            timeout=60,
            cwd=str(tmpdir_path),
        )
        if not pdf_path.exists():
            raise RuntimeError("LibreOffice no generó el PDF. Compruebe la plantilla e instalación.")
        return pdf_path.read_bytes()


@app.get("/api/analisis-icr/certificado-pdf")
async def api_analisis_icr_certificado_pdf(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    codigo_cliente: Optional[int] = None,
    numero_parte: List[str] = Query(default=[], description="Números de parte seleccionados (columna Select)"),
):
    """
    Genera el certificado de origen (C.O.) en PDF para los números de parte seleccionados.
    Requiere al menos un número de parte seleccionado en la columna Select.
    """
    if codigo_cliente is None:
        return JSONResponse({"error": "Falta el parámetro codigo_cliente"}, status_code=400)
    selected = [str(n).strip() for n in numero_parte if n is not None and str(n).strip()]
    if not selected:
        return JSONResponse(
            {"error": "No hay números de parte seleccionados. Seleccione al menos uno en la columna Select."},
            status_code=400,
        )
    detalle = await crud.get_analisis_icr_detalle(db, codigo_cliente=codigo_cliente)
    todas_partes = detalle.get("partes") or []
    # Filtrar solo las partes seleccionadas, manteniendo el orden de selección
    partes_para_certificado = []
    for np in selected:
        for p in todas_partes:
            if (p.get("part_number") or "").strip() == np:
                partes_para_certificado.append(p)
                break
    if not partes_para_certificado:
        return JSONResponse(
            {"error": "Ninguno de los números de parte seleccionados pertenece a este cliente."},
            status_code=400,
        )
    # Customer Part Number desde cross_reference para la hoja C.O. 3
    part_numbers = [(p.get("part_number") or "").strip() for p in partes_para_certificado]
    cross_ref = await crud.get_cross_reference_por_cliente_materiales(db, str(codigo_cliente), part_numbers)
    for p in partes_para_certificado:
        np = (p.get("part_number") or "").strip()
        p["customer_part_number"] = (cross_ref.get(np) or np or "").strip()

    ahora = datetime.now()
    def _fmt_date(d):
        return f"{d.month}/{d.day}/{d.strftime('%y')}"
    certification_date = _fmt_date(ahora)
    blanket_from = _fmt_date(ahora.replace(month=1, day=1))
    blanket_to = _fmt_date(ahora.replace(month=12, day=31))
    context = {
        **_CERT_CO_LEONI,
        "codigo_cliente": codigo_cliente,
        "cliente_nombre": detalle.get("cliente_nombre") or str(codigo_cliente),
        "importer_address": None,
        "partes": partes_para_certificado,
        "blanket_period_from": blanket_from,
        "blanket_period_to": blanket_to,
        "certification_date": certification_date,
        "num_pages": 2,  # C.O. (pág 1) + C.O. 3 (pág 2)
        "_ahora": ahora,
    }
    try:
        pdf_bytes = await asyncio.to_thread(_render_certificado_co_pdf, context)
    except Exception as e:
        logger.exception("Error generando PDF certificado C.O.")
        return JSONResponse(
            {"error": f"Error al generar el PDF: {e!s}"},
            status_code=500,
        )
    filename = f"certificado_co_{codigo_cliente}.pdf"
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/api/analisis-icr/certificado-excel")
async def api_analisis_icr_certificado_excel(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    codigo_cliente: Optional[int] = None,
    numero_parte: List[str] = Query(default=[], description="Números de parte seleccionados (columna Select)"),
):
    """
    Genera el certificado de origen (C.O.) en Excel (.xlsx) para revisión.
    Mismos parámetros que el certificado PDF; solo devuelve el archivo Excel rellenado.
    """
    if codigo_cliente is None:
        return JSONResponse({"error": "Falta el parámetro codigo_cliente"}, status_code=400)
    selected = [str(n).strip() for n in numero_parte if n is not None and str(n).strip()]
    if not selected:
        return JSONResponse(
            {"error": "No hay números de parte seleccionados. Seleccione al menos uno en la columna Select."},
            status_code=400,
        )
    detalle = await crud.get_analisis_icr_detalle(db, codigo_cliente=codigo_cliente)
    todas_partes = detalle.get("partes") or []
    partes_para_certificado = []
    for np in selected:
        for p in todas_partes:
            if (p.get("part_number") or "").strip() == np:
                partes_para_certificado.append(p)
                break
    if not partes_para_certificado:
        return JSONResponse(
            {"error": "Ninguno de los números de parte seleccionados pertenece a este cliente."},
            status_code=400,
        )
    # Customer Part Number desde cross_reference para la hoja C.O. 3
    part_numbers = [(p.get("part_number") or "").strip() for p in partes_para_certificado]
    cross_ref = await crud.get_cross_reference_por_cliente_materiales(db, str(codigo_cliente), part_numbers)
    for p in partes_para_certificado:
        np = (p.get("part_number") or "").strip()
        p["customer_part_number"] = (cross_ref.get(np) or np or "").strip()

    ahora = datetime.now()

    def _fmt_date(d):
        return f"{d.month}/{d.day}/{d.strftime('%y')}"
    blanket_from = _fmt_date(ahora.replace(month=1, day=1))
    blanket_to = _fmt_date(ahora.replace(month=12, day=31))
    context = {
        **_CERT_CO_LEONI,
        "codigo_cliente": codigo_cliente,
        "cliente_nombre": detalle.get("cliente_nombre") or str(codigo_cliente),
        "importer_address": None,
        "partes": partes_para_certificado,
        "blanket_period_from": blanket_from,
        "blanket_period_to": blanket_to,
        "certification_date": _fmt_date(ahora),
        "num_pages": 2,  # C.O. (pág 1) + C.O. 3 (pág 2)
        "_ahora": ahora,
    }
    try:
        xlsx_bytes = await asyncio.to_thread(_render_certificado_co_xlsx, context)
    except Exception as e:
        logger.exception("Error generando Excel certificado C.O.")
        return JSONResponse(
            {"error": f"Error al generar el Excel: {e!s}"},
            status_code=500,
        )
    filename = f"certificado_co_{codigo_cliente}.xlsx"
    return Response(
        content=xlsx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/api/analisis-icr/certificado-pack")
async def api_analisis_icr_certificado_pack(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    codigo_cliente: Optional[int] = None,
    numero_parte: List[str] = Query(default=[], description="Números de parte seleccionados (columna Select)"),
):
    """
    Genera un ZIP con los documentos según la selección:
    - Si hay partes que califican (ICR >= 60%): incluye certificado Excel + PDF (solo esas partes).
    - Si hay partes que no califican (ICR < 60% o sin ICR): incluye además Word + PDF de no calificados (cliente completo).
    Si solo hay partes que califican: ZIP con certificado_co.xlsx y certificado_co.pdf.
    """
    if codigo_cliente is None:
        return JSONResponse({"error": "Falta el parámetro codigo_cliente"}, status_code=400)
    selected = [str(n).strip() for n in numero_parte if n is not None and str(n).strip()]
    if not selected:
        return JSONResponse(
            {"error": "Seleccione al menos un número de parte en la columna Select para generar el certificado."},
            status_code=400,
        )
    detalle = await crud.get_analisis_icr_detalle(db, codigo_cliente=codigo_cliente)
    todas_partes = detalle.get("partes") or []
    partes_seleccionadas = []
    for np in selected:
        for p in todas_partes:
            if (p.get("part_number") or "").strip() == np:
                partes_seleccionadas.append(p)
                break
    if not partes_seleccionadas:
        return JSONResponse(
            {"error": "Ninguno de los números de parte seleccionados pertenece a este cliente."},
            status_code=400,
        )

    def _califica(p):
        icr = p.get("icr")
        if icr is None:
            return False
        try:
            return float(icr) >= 60
        except (TypeError, ValueError):
            return False

    partes_que_califican = [p for p in partes_seleccionadas if _califica(p)]
    hay_no_calificados = any(not _califica(p) for p in partes_seleccionadas)

    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        ahora = datetime.now()

        if partes_que_califican:
            part_numbers = [(p.get("part_number") or "").strip() for p in partes_que_califican]
            cross_ref = await crud.get_cross_reference_por_cliente_materiales(db, str(codigo_cliente), part_numbers)
            for p in partes_que_califican:
                np = (p.get("part_number") or "").strip()
                p["customer_part_number"] = (cross_ref.get(np) or np or "").strip()
            def _fmt_date(d):
                return f"{d.month}/{d.day}/{d.strftime('%y')}"
            cert_context = {
                **_CERT_CO_LEONI,
                "codigo_cliente": codigo_cliente,
                "cliente_nombre": detalle.get("cliente_nombre") or str(codigo_cliente),
                "importer_address": None,
                "partes": partes_que_califican,
                "blanket_period_from": _fmt_date(ahora.replace(month=1, day=1)),
                "blanket_period_to": _fmt_date(ahora.replace(month=12, day=31)),
                "certification_date": _fmt_date(ahora),
                "num_pages": 2,
                "_ahora": ahora,
            }
            xlsx_bytes = await asyncio.to_thread(_render_certificado_co_xlsx, cert_context)
            zf.writestr(f"certificado_co_{codigo_cliente}.xlsx", xlsx_bytes)
            pdf_bytes = await asyncio.to_thread(_render_certificado_co_pdf, cert_context)
            zf.writestr(f"certificado_co_{codigo_cliente}.pdf", pdf_bytes)

        if hay_no_calificados:
            todas_partes = detalle.get("partes") or []
            partes_no_calificados = [
                p for p in todas_partes
                if p.get("icr") is None or (isinstance(p.get("icr"), (int, float)) and float(p["icr"]) < 60)
            ]
            if partes_no_calificados:
                part_numbers = [(p.get("part_number") or "").strip() for p in partes_no_calificados]
                cross_ref = await crud.get_cross_reference_por_cliente_materiales(db, str(codigo_cliente), part_numbers)
                for p in partes_no_calificados:
                    np = (p.get("part_number") or "").strip()
                    p["customer_part_number"] = (cross_ref.get(np) or np or "").strip()
                no_cert_context = {
                    "codigo_cliente": codigo_cliente,
                    "cliente_nombre": detalle.get("cliente_nombre") or str(codigo_cliente),
                    "partes": partes_no_calificados,
                    "blanket_year": ahora.year,
                    "_ahora": ahora,
                }
                docx_bytes = await asyncio.to_thread(_render_no_calificados_docx, no_cert_context)
                zf.writestr(f"no_calificados_{codigo_cliente}.docx", docx_bytes)
                try:
                    pdf_nc_bytes = await asyncio.to_thread(_render_no_calificados_pdf, no_cert_context)
                    zf.writestr(f"no_calificados_{codigo_cliente}.pdf", pdf_nc_bytes)
                except Exception as e:
                    logger.warning("No se pudo incluir PDF no calificados en el pack: %s", e)

    buf.seek(0)
    filename_zip = f"certificado_pack_{codigo_cliente}.zip"
    return Response(
        content=buf.getvalue(),
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{filename_zip}"'},
    )


@app.get("/api/analisis-icr/no-calificados-docx")
async def api_analisis_icr_no_calificados_docx(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    codigo_cliente: Optional[int] = None,
):
    """
    Genera la declaración jurada (Word) para componentes con ICR < 60% (no calificados USMCA).
    Incluye todos los números de parte del cliente con ICR < 60% o sin ICR calculado.
    """
    if codigo_cliente is None:
        return JSONResponse({"error": "Falta el parámetro codigo_cliente"}, status_code=400)
    detalle = await crud.get_analisis_icr_detalle(db, codigo_cliente=codigo_cliente)
    todas_partes = detalle.get("partes") or []
    partes_no_calificados = [
        p for p in todas_partes
        if p.get("icr") is None or (isinstance(p.get("icr"), (int, float)) and float(p["icr"]) < 60)
    ]
    if not partes_no_calificados:
        return JSONResponse(
            {"error": "No hay componentes con ICR < 60% para este cliente."},
            status_code=400,
        )
    part_numbers = [(p.get("part_number") or "").strip() for p in partes_no_calificados]
    # Customer Part Number en la plantilla Word: se obtiene del cross reference (customer_material por material/cliente)
    cross_ref = await crud.get_cross_reference_por_cliente_materiales(db, str(codigo_cliente), part_numbers)
    for p in partes_no_calificados:
        np = (p.get("part_number") or "").strip()
        p["customer_part_number"] = (cross_ref.get(np) or np or "").strip()
    ahora = datetime.now()
    context = {
        "codigo_cliente": codigo_cliente,
        "cliente_nombre": detalle.get("cliente_nombre") or str(codigo_cliente),
        "partes": partes_no_calificados,
        "blanket_year": ahora.year,
        "_ahora": ahora,
    }
    try:
        docx_bytes = await asyncio.to_thread(_render_no_calificados_docx, context)
    except Exception as e:
        logger.exception("Error generando Word no calificados.")
        return JSONResponse(
            {"error": f"Error al generar el documento: {e!s}"},
            status_code=500,
        )
    filename = f"no_calificados_{codigo_cliente}.docx"
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/api/analisis-icr/no-calificados-pdf")
async def api_analisis_icr_no_calificados_pdf(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    codigo_cliente: Optional[int] = None,
):
    """
    Genera la declaración jurada en PDF para componentes con ICR < 60% (no calificados USMCA).
    Requiere LibreOffice instalado para la conversión docx → pdf.
    """
    if codigo_cliente is None:
        return JSONResponse({"error": "Falta el parámetro codigo_cliente"}, status_code=400)
    detalle = await crud.get_analisis_icr_detalle(db, codigo_cliente=codigo_cliente)
    todas_partes = detalle.get("partes") or []
    partes_no_calificados = [
        p for p in todas_partes
        if p.get("icr") is None or (isinstance(p.get("icr"), (int, float)) and float(p["icr"]) < 60)
    ]
    if not partes_no_calificados:
        return JSONResponse(
            {"error": "No hay componentes con ICR < 60% para este cliente."},
            status_code=400,
        )
    part_numbers = [(p.get("part_number") or "").strip() for p in partes_no_calificados]
    # Customer Part Number en la plantilla Word: se obtiene del cross reference (customer_material por material/cliente)
    cross_ref = await crud.get_cross_reference_por_cliente_materiales(db, str(codigo_cliente), part_numbers)
    for p in partes_no_calificados:
        np = (p.get("part_number") or "").strip()
        p["customer_part_number"] = (cross_ref.get(np) or np or "").strip()
    ahora = datetime.now()
    context = {
        "codigo_cliente": codigo_cliente,
        "cliente_nombre": detalle.get("cliente_nombre") or str(codigo_cliente),
        "partes": partes_no_calificados,
        "blanket_year": ahora.year,
        "_ahora": ahora,
    }
    try:
        pdf_bytes = await asyncio.to_thread(_render_no_calificados_pdf, context)
    except Exception as e:
        logger.exception("Error generando PDF no calificados.")
        return JSONResponse(
            {"error": f"Error al generar el PDF: {e!s}"},
            status_code=500,
        )
    filename = f"no_calificados_{codigo_cliente}.pdf"
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/api/reportes/partes-no-calificados-icr")
async def api_reportes_partes_no_calificados_icr(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Reporte Excel: todos los números de parte que no cumplen ICR > 60% por cliente.
    Origen: ventas con sales_km y precios (exmetal/full_metal) no vacíos ni 0;
    para cada (codigo_cliente, part_number) se calcula ICR y se listan los que tienen ICR < 60 o sin ICR.
    """
    from io import BytesIO
    from sqlalchemy import select
    from app.db.models import Parte
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    import logging
    logger = logging.getLogger(__name__)

    try:
        pares = await crud.list_todos_clientes_partes_con_ventas_icr(db)
    except Exception as e:
        logger.exception("Error en list_todos_clientes_partes_con_ventas_icr")
        return JSONResponse(
            {"error": f"Error al obtener datos para el reporte: {str(e)}"},
            status_code=500,
        )
    if not pares:
        return JSONResponse(
            {"error": "No hay ventas con sales_km y precios ICR válidos para generar el reporte."},
            status_code=400,
        )

    # Calcular ICR por cada (cliente, parte) y filtrar no conformes (ICR < 60 o None)
    filas = []
    try:
        for p in pares:
            icr = await crud.get_icr_para_parte(db, p["codigo_cliente"], p["part_number"])
            if icr is not None and float(icr) > 60:
                continue
            filas.append({
                "codigo_cliente": p["codigo_cliente"],
                "cliente_nombre": p["cliente_nombre"],
                "part_number": p["part_number"],
                "icr": icr,
            })
    except Exception as e:
        logger.exception("Error al calcular ICR en reporte partes no calificados")
        return JSONResponse(
            {"error": f"Error al calcular ICR para el reporte: {str(e)}"},
            status_code=500,
        )

    if not filas:
        return JSONResponse(
            {"error": "No hay componentes con ICR < 60% (todos cumplen o no se pudo calcular)."},
            status_code=400,
        )

    try:
        # Descripción y tarifa desde tabla Parte
        part_numbers = list({f["part_number"] for f in filas})
        partes_query = await db.execute(
            select(Parte.numero_parte, Parte.descripcion, Parte.fraccion).where(Parte.numero_parte.in_(part_numbers))
        )
        partes_map = {}
        for r in partes_query.all():
            pn = (r[0] and str(r[0])).strip()
            if pn:
                partes_map[pn] = {"description": (r[1] or "").strip() or "—", "tariff_schedule": (r[2] or "").strip() or "—"}

        for f in filas:
            info = partes_map.get(f["part_number"]) or {}
            f["description"] = info.get("description", "—")
            f["tariff_schedule"] = info.get("tariff_schedule", "—")

        # Excel
        wb = Workbook()
        ws = wb.active
        ws.title = "Partes ICR < 60%"
        headers = ["Codigo Cliente", "Cliente Nombre", "Part Number", "Description", "Tariff Schedule", "ICR"]
        header_fill = PatternFill(fill_type="solid", start_color="4472C4", end_color="4472C4")
        header_font = Font(bold=True, color="FFFFFF")
        for col_idx, h in enumerate(headers, start=1):
            c = ws.cell(row=1, column=col_idx, value=h)
            c.fill = header_fill
            c.font = header_font
        for row_idx, f in enumerate(filas, start=2):
            ws.cell(row=row_idx, column=1, value=f["codigo_cliente"])
            ws.cell(row=row_idx, column=2, value=f["cliente_nombre"])
            ws.cell(row=row_idx, column=3, value=f["part_number"])
            ws.cell(row=row_idx, column=4, value=f["description"])
            ws.cell(row=row_idx, column=5, value=f["tariff_schedule"])
            ws.cell(row=row_idx, column=6, value=f["icr"] if f["icr"] is not None else "—")
        for col in range(1, len(headers) + 1):
            ws.column_dimensions[ws.cell(row=1, column=col).column_letter].width = 18

        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)
        filename = f"reporte_partes_icr_menor_60_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
        return Response(
            content=buf.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        logger.exception("Error al generar Excel en reporte partes no calificados")
        return JSONResponse(
            {"error": f"Error al generar el Excel: {str(e)}"},
            status_code=500,
        )


@app.get("/analisis-icr/material")
async def analisis_icr_material(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    codigo_cliente: Optional[int] = None,
    numero_parte: Optional[str] = None,
):
    """Página de detalle de un material/número de parte dentro del contexto del cliente, con ítems del BOM."""
    if not codigo_cliente or not (numero_parte or "").strip():
        return RedirectResponse(url="/analisis-icr", status_code=302)

    numero_parte = (numero_parte or "").strip()
    # Nombre del cliente
    cliente_nombre = None
    from sqlalchemy import select
    cliente_row = await db.execute(
        select(Cliente.nombre).where(Cliente.codigo_cliente == codigo_cliente)
    )
    cliente_nombre = cliente_row.scalar_one_or_none()
    if cliente_nombre is None:
        v = await db.execute(
            select(Venta.cliente).where(
                Venta.codigo_cliente == codigo_cliente,
                Venta.cliente.isnot(None),
            ).limit(1)
        )
        cliente_nombre = (v.scalar_one_or_none() or "").strip() or str(codigo_cliente)

    bom_data = await crud.get_bom_items_for_parte(db, numero_parte)
    parte_info = bom_data.get("parte")
    if not parte_info:
        parte_info = {"numero_parte": numero_parte, "descripcion": None, "fraccion": None}

    items_orig = bom_data.get("items_originating") or []
    items_non_orig = bom_data.get("items_non_originating") or []
    # Lista unificada de ítems del BOM para la tabla resumen (cada uno con clave "tipo")
    items_bom = [dict(item, tipo="Originating") for item in items_orig] + [dict(item, tipo="Non-Originating") for item in items_non_orig]

    # Totales Value por breakdown (suma de columna Value de cada tabla)
    def _sum_value(items):
        total = 0.0
        for item in items:
            for p in (item.get("proveedores") or []):
                qty = item.get("qty")
                pct = p.get("porcentaje_compra")
                precio = p.get("precio_compra")
                if qty is not None and pct is not None and precio is not None:
                    total += float(qty) * float(pct) / 100 * float(precio)
        return round(total, 3)

    total_originating_value = _sum_value(items_orig)
    total_non_originating_value = _sum_value(items_non_orig)

    # F.O.B USD value = último precio de venta de ese número de parte a ese cliente
    fob_total_value = await crud.get_ultimo_precio_venta_cliente_parte(db, codigo_cliente, numero_parte)
    if fob_total_value is not None:
        fob_total_value = round(fob_total_value, 3)

    # Markup = F.O.B USD value - Total Originating Supplies - Total Non-Originating Supplies
    markup_value = None
    if fob_total_value is not None:
        markup_value = round(
            fob_total_value - total_originating_value - total_non_originating_value,
            3,
        )

    # Regional index = ((Total Originating Supplies + Markup) / F.O.B USD value) * 100
    regional_index = None
    if fob_total_value is not None and fob_total_value != 0 and markup_value is not None:
        regional_index = round(
            (total_originating_value + markup_value) / fob_total_value * 100,
            2,
        )

    response = templates.TemplateResponse(
        "analisis_icr_material.html",
        {
            "request": request,
            "active_page": "analisis_icr",
            "current_user": current_user,
            "codigo_cliente": codigo_cliente,
            "cliente_nombre": cliente_nombre,
            "numero_parte": numero_parte,
            "parte_info": parte_info,
            "items_bom": items_bom,
            "items_originating": items_orig,
            "items_non_originating": items_non_orig,
            "fob_total_value": fob_total_value,
            "total_originating_value": total_originating_value,
            "total_non_originating_value": total_non_originating_value,
            "markup_value": markup_value,
            "regional_index": regional_index,
        }
    )
    # Evitar caché del navegador para que cambios en porcentaje_compra (pais_origen_material) se vean al recargar
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    response.headers["Pragma"] = "no-cache"
    return response


@app.get("/precios-venta")
async def precios_venta(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Página de precios de venta - requiere autenticación."""
    total_precios_venta = await crud.count_precios_venta(db)
    return templates.TemplateResponse(
        "precios_venta.html",
        {
            "request": request,
            "active_page": "precios_venta",
            "current_user": current_user,
            "total_precios_venta": total_precios_venta,
        },
    )


@app.get("/api/precios-venta")
async def api_precios_venta(
    request: Request,
    q: Optional[str] = None,
    limit: int = 25,
    offset: int = 0,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """API para listar precios_venta con búsqueda y paginación."""
    from sqlalchemy import select

    limit = max(1, min(limit, 200))
    offset = max(0, offset)
    q = (q or "").strip()

    rows_db = await crud.list_precios_venta(db, search=q or None, limit=limit, offset=offset)
    total = await crud.count_precios_venta(db, search=q or None)

    # Resolver grupo por código de cliente (si existe en cliente_grupo).
    codigos_cliente = sorted({row.codigo_cliente for row in rows_db if row.codigo_cliente is not None})
    grupos_por_cliente = {}
    if codigos_cliente:
        grupos_result = await db.execute(
            select(
                ClienteGrupo.codigo_cliente,
                ClienteGrupo.grupo,
            )
            .where(ClienteGrupo.codigo_cliente.in_(codigos_cliente))
            .order_by(
                ClienteGrupo.codigo_cliente.asc(),
                ClienteGrupo.updated_at.desc(),
                ClienteGrupo.id.desc(),
            )
        )
        for codigo_cliente, grupo in grupos_result.all():
            if codigo_cliente not in grupos_por_cliente:
                grupos_por_cliente[codigo_cliente] = grupo

    rows = [
        {
            "id": row.id,
            "codigo_cliente": row.codigo_cliente,
            "nombre_cliente": row.cliente.nombre if row.cliente else None,
            "grupo_cliente": grupos_por_cliente.get(row.codigo_cliente),
            "numero_parte": row.numero_parte,
            "descripcion_parte": row.parte.descripcion if row.parte else None,
            "tipo_cable": row.tipo_cable,
            "precio_venta": float(row.precio_venta) if row.precio_venta is not None else None,
            "comentario": row.comentario,
            "comentario_2": row.comentario_2,
            "comentario_3": row.comentario_3,
            "updated_at": row.updated_at.isoformat() if row.updated_at else None,
        }
        for row in rows_db
    ]

    return {
        "ok": True,
        "q": q or None,
        "limit": limit,
        "offset": offset,
        "total": int(total),
        "rows": rows,
    }


@app.post("/api/precios-venta/{precio_venta_id}/actualizar")
async def api_actualizar_precio_venta(
    precio_venta_id: int,
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Actualiza solo precio_venta y comentarios de un registro de precios_venta.
    """
    payload = await request.json()

    precio_venta_raw = payload.get("precio_venta")
    comentario = payload.get("comentario")
    comentario_2 = payload.get("comentario_2")
    comentario_3 = payload.get("comentario_3")

    precio_venta_value = None
    if precio_venta_raw is not None and str(precio_venta_raw).strip() != "":
        try:
            precio_venta_value = Decimal(str(precio_venta_raw).strip())
        except (InvalidOperation, ValueError):
            return JSONResponse(
                status_code=400,
                content={"ok": False, "mensaje": "El campo precio_venta no es válido."},
            )

    item = await crud.update_precio_venta(
        db=db,
        precio_venta_id=precio_venta_id,
        precio_venta=precio_venta_value,
        comentario=(comentario or None),
        comentario_2=(comentario_2 or None),
        comentario_3=(comentario_3 or None),
        user_id=current_user.id,
    )

    if not item:
        return JSONResponse(
            status_code=404,
            content={"ok": False, "mensaje": "Registro de precio de venta no encontrado."},
        )

    return {
        "ok": True,
        "mensaje": "Precio de venta actualizado correctamente.",
        "row": {
            "id": item.id,
            "codigo_cliente": item.codigo_cliente,
            "numero_parte": item.numero_parte,
            "tipo_cable": item.tipo_cable,
            "precio_venta": float(item.precio_venta) if item.precio_venta is not None else None,
            "comentario": item.comentario,
            "comentario_2": item.comentario_2,
            "comentario_3": item.comentario_3,
            "updated_at": item.updated_at.isoformat() if item.updated_at else None,
        },
    }


@app.post("/api/precios-venta/{precio_venta_id}/actualizar-comentarios")
async def api_actualizar_comentarios_precio_venta(
    precio_venta_id: int,
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Actualiza únicamente comentarios de un registro de precios_venta.
    """
    payload = await request.json()
    comentario = payload.get("comentario")
    comentario_2 = payload.get("comentario_2")
    comentario_3 = payload.get("comentario_3")

    item = await crud.update_comentarios_precio_venta(
        db=db,
        precio_venta_id=precio_venta_id,
        comentario=(comentario or None),
        comentario_2=(comentario_2 or None),
        comentario_3=(comentario_3 or None),
        user_id=current_user.id,
    )

    if not item:
        return JSONResponse(
            status_code=404,
            content={"ok": False, "mensaje": "Registro de precio de venta no encontrado."},
        )

    return {
        "ok": True,
        "mensaje": "Comentarios actualizados correctamente.",
        "row": {
            "id": item.id,
            "comentario": item.comentario,
            "comentario_2": item.comentario_2,
            "comentario_3": item.comentario_3,
            "updated_at": item.updated_at.isoformat() if item.updated_at else None,
        },
    }


@app.post("/api/precios-venta/actualizar-desde-ventas")
async def api_actualizar_precios_venta_desde_ventas(
    request: Request,
    current_user: User = Depends(require_roles(["admin"])),
    db: AsyncSession = Depends(get_db),
):
    """
    Sincroniza precio_venta desde ventas usando:
    - match por (codigo_cliente, numero_parte == producto_condensado)
    - solo ventas con sales_km y precio_full_metal_km
    - toma la última venta por combinación
    """
    try:
        resultado = await crud.sincronizar_precios_venta_desde_ventas(
            db=db,
            user_id=current_user.id,
        )
        partes_msg = []
        if resultado.get("creados", 0):
            partes_msg.append(f"{resultado['creados']} creado(s)")
        if resultado.get("actualizados", 0):
            partes_msg.append(f"{resultado['actualizados']} actualizado(s)")
        partes_msg.append(f"{resultado.get('sin_cambios', 0)} sin cambios")
        partes_msg.append(f"{resultado.get('sin_venta', 0)} sin venta aplicable")
        if resultado.get("creados_omitidos", 0):
            partes_msg.append(f"{resultado['creados_omitidos']} omitido(s) (cliente o parte no existe)")
        mensaje = "✓ Actualización completada. " + ", ".join(partes_msg) + "."
        return {
            "ok": True,
            "mensaje": mensaje,
            **resultado,
        }
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={
                "ok": False,
                "mensaje": f"Error al actualizar precios desde ventas: {str(e)}",
            },
        )


@app.get("/api/precios-venta/historial")
async def api_precios_venta_historial(
    request: Request,
    current_user: User = Depends(require_roles(["admin"])),
    db: AsyncSession = Depends(get_db),
):
    """Devuelve los últimos movimientos del historial de precios de venta. Solo administradores."""
    historial = await crud.list_precio_venta_historial(db, limit=5, offset=0)
    items = []
    for h in historial:
        items.append(
            {
                "id": h.id,
                "precio_venta_id": h.precio_venta_id,
                "codigo_cliente": h.codigo_cliente,
                "numero_parte": h.numero_parte,
                "tipo_cable": h.tipo_cable,
                "operacion": h.operacion,
                "user_email": h.user.email if h.user else None,
                "user_nombre": h.user.nombre if h.user else None,
                "datos_antes": h.datos_antes,
                "datos_despues": h.datos_despues,
                "campos_modificados": h.campos_modificados,
                "detalle": h.detalle,
                "created_at": h.created_at.isoformat() if h.created_at else None,
            }
        )
    return {"ok": True, "movimientos": items}


@app.get("/api/ventas")
async def api_ventas(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    search: Optional[str] = None,
    cliente: Optional[str] = None,
    codigo_cliente: Optional[int] = None,
    periodo_inicio: Optional[str] = None,
    periodo_fin: Optional[str] = None,
    producto: Optional[str] = None,
    planta: Optional[str] = None,
    limit: int = 100,
    offset: int = 0
):
    """API para obtener ventas con filtros y paginación."""
    # Convertir fechas de string a datetime si están presentes
    periodo_inicio_dt = None
    periodo_fin_dt = None
    
    if periodo_inicio:
        try:
            periodo_inicio_dt = datetime.strptime(periodo_inicio, "%Y-%m-%d")
        except ValueError:
            pass
    
    if periodo_fin:
        try:
            periodo_fin_dt = datetime.strptime(periodo_fin, "%Y-%m-%d")
        except ValueError:
            pass
    
    ventas = await crud.list_ventas(
        db=db,
        limit=limit,
        offset=offset,
        search=search,
        cliente=cliente,
        codigo_cliente=codigo_cliente,
        periodo_inicio=periodo_inicio_dt,
        periodo_fin=periodo_fin_dt,
        producto=producto,
        planta=planta,
        only_with_sales_km=True
    )
    
    total = await crud.count_ventas(
        db=db,
        search=search,
        cliente=cliente,
        codigo_cliente=codigo_cliente,
        periodo_inicio=periodo_inicio_dt,
        periodo_fin=periodo_fin_dt,
        producto=producto,
        planta=planta,
        only_with_sales_km=True
    )
    
    # Convertir a diccionarios para JSON
    ventas_dict = []
    for venta in ventas:
        venta_data = {
            "id": venta.id,
            "cliente": venta.cliente,
            "codigo_cliente": venta.codigo_cliente,
            "grupo": venta.grupo.grupo if venta.grupo else None,
            "unidad_negocio": venta.unidad_negocio,
            "periodo": venta.periodo.strftime("%Y-%m-%d") if venta.periodo else None,
            "producto_condensado": venta.producto_condensado,
            "region_asc": venta.region_asc,
            "planta": venta.planta,
            "ship_to_party": venta.ship_to_party,
            "producto": venta.producto,
            "descripcion_producto": venta.descripcion_producto,
            "turnover_wo_metal": float(venta.turnover_wo_metal) if venta.turnover_wo_metal else None,
            "oe_turnover_like_fi": float(venta.oe_turnover_like_fi) if venta.oe_turnover_like_fi else None,
            "copper_sales_cuv": float(venta.copper_sales_cuv) if venta.copper_sales_cuv else None,
            "cu_sales_effect": float(venta.cu_sales_effect) if venta.cu_sales_effect else None,
            "cu_result": float(venta.cu_result) if venta.cu_result else None,
            "quantity_oe_to_m": float(venta.quantity_oe_to_m) if venta.quantity_oe_to_m else None,
            "quantity_oe_to_ft": float(venta.quantity_oe_to_ft) if venta.quantity_oe_to_ft else None,
            "cu_weight_techn_cut": float(venta.cu_weight_techn_cut) if venta.cu_weight_techn_cut else None,
            "cu_weight_sales_cuv": float(venta.cu_weight_sales_cuv) if venta.cu_weight_sales_cuv else None,
            "conversion_ft_a_m": float(venta.conversion_ft_a_m) if venta.conversion_ft_a_m else None,
            "sales_total_mts": float(venta.sales_total_mts) if venta.sales_total_mts else None,
            "sales_km": float(venta.sales_km) if venta.sales_km else None,
            "precio_exmetal_km": float(venta.precio_exmetal_km) if venta.precio_exmetal_km else None,
            "precio_full_metal_km": float(venta.precio_full_metal_km) if venta.precio_full_metal_km else None,
            "precio_exmetal_m": float(venta.precio_exmetal_m) if venta.precio_exmetal_m else None,
            "precio_full_metal_m": float(venta.precio_full_metal_m) if venta.precio_full_metal_m else None,
            "created_at": venta.created_at.isoformat() if venta.created_at else None
        }
        ventas_dict.append(venta_data)
    
    return JSONResponse({
        "ventas": ventas_dict,
        "total": total,
        "limit": limit,
        "offset": offset
    })


@app.get("/clientes")
async def clientes(request: Request, current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    """Página de clientes - requiere autenticación."""
    # Cargar los clientes desde la base de datos
    clientes_data = await crud.list_clientes(db, limit=5000)
    
    # Calcular estadísticas
    total_clientes = await crud.count_clientes(db)
    
    # Obtener países únicos para filtros
    paises = await crud.get_paises_clientes(db)
    total_paises = len(paises)
    
    # Obtener última actualización
    if clientes_data:
        ultima_actualizacion = max([c.updated_at for c in clientes_data if c.updated_at], default=None)
        if ultima_actualizacion:
            ultima_actualizacion = ultima_actualizacion.strftime('%d/%m/%Y %H:%M')
    else:
        ultima_actualizacion = None
    
    return templates.TemplateResponse(
        "clientes.html",
        {
            "request": request,
            "active_page": "clientes",
            "current_user": current_user,
            "clientes": clientes_data,
            "total_clientes": total_clientes,
            "paises": paises,
            "total_paises": total_paises,
            "ultima_actualizacion": ultima_actualizacion
        }
    )


@app.get("/grupos-clientes")
async def grupos_clientes(request: Request, current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    """Página de grupos de clientes - requiere autenticación."""
    # Cargar los grupos de clientes desde la base de datos
    grupos_data = await crud.list_cliente_grupos(db, limit=10000)
    
    # Calcular estadísticas
    total_grupos = await crud.count_cliente_grupos(db)
    
    # Obtener grupos únicos para filtros
    grupos_unicos = sorted(set([g.grupo for g in grupos_data if g.grupo]), key=lambda x: x or "")
    grupos_viejos_unicos = sorted(set([g.grupo_viejo for g in grupos_data if g.grupo_viejo]), key=lambda x: x or "")
    
    return templates.TemplateResponse(
        "grupos_clientes.html",
        {
            "request": request,
            "active_page": "grupos_clientes",
            "current_user": current_user,
            "grupos": grupos_data,
            "total_grupos": total_grupos,
            "grupos_unicos": grupos_unicos,
            "grupos_viejos_unicos": grupos_viejos_unicos
        }
    )


def convertir_a_cdmx(dt: Optional[datetime]) -> Optional[datetime]:
    """Convierte una fecha UTC a la zona horaria de Ciudad de México."""
    if dt is None:
        return None
    
    # Si la fecha ya tiene timezone info, convertirla
    if dt.tzinfo is not None:
        return dt.astimezone(ZoneInfo("America/Mexico_City"))
    else:
        # Asumir que está en UTC si no tiene timezone
        dt_utc = dt.replace(tzinfo=ZoneInfo("UTC"))
        return dt_utc.astimezone(ZoneInfo("America/Mexico_City"))


# Agregar filtro personalizado a Jinja2 para convertir fechas a CDMX
def datetime_cdmx_filter(dt: Optional[datetime], format_str: str = '%d/%m/%Y %H:%M') -> str:
    """Filtro de Jinja2 para convertir datetime a CDMX y formatear."""
    if dt is None:
        return 'N/A'
    dt_cdmx = convertir_a_cdmx(dt)
    return dt_cdmx.strftime(format_str)


def to_iso_filter(value):
    """Filtro Jinja2: convierte datetime a ISO string; si ya es str o otro tipo, lo devuelve como string."""
    if value is None:
        return ''
    if hasattr(value, 'isoformat'):
        return value.isoformat()
    return str(value)


# Registrar el filtro en el entorno de templates
templates.env.filters['datetime_cdmx'] = datetime_cdmx_filter
templates.env.filters['to_iso'] = to_iso_filter


@app.get("/compras")
async def compras(request: Request, current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    """Página de compras - requiere autenticación."""
    # Admin ve ejecuciones de todos; operador solo las propias
    if current_user.rol == "admin":
        executions = await crud.list_executions(db, limit=5)
    else:
        executions = await crud.list_executions(db, user_id=current_user.id, limit=5)
    
    return templates.TemplateResponse(
        "compras.html",
        {
            "request": request,
            "active_page": "compras",
            "current_user": current_user,
            "executions": executions
        }
    )


@app.get("/todas-compras")
async def todas_compras(request: Request, current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    """Página para ver todas las compras con filtros - requiere autenticación."""
    # Obtener el total de compras para mostrar en la estadística
    total_compras = await crud.count_compras(db)
    
    return templates.TemplateResponse(
        "todas_compras.html",
        {
            "request": request,
            "active_page": "todas_compras",
            "current_user": current_user,
            "total_compras": total_compras
        }
    )


@app.get("/boms")
async def boms(request: Request, current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    """Página de BOMs - Lista de materiales - requiere autenticación."""
    # Cargar los BOMs desde la base de datos
    boms_data = await crud.list_bom_flat(db, limit=5000)
    
    # Calcular estadísticas
    total_registros = len(boms_data)
    partes_unicas = len(set(bom.fg_part_no for bom in boms_data))
    materiales_unicos = len(set(bom.material for bom in boms_data))
    
    return templates.TemplateResponse(
        "boms.html",
        {
            "request": request,
            "active_page": "boms",
            "current_user": current_user,
            "boms": boms_data,
            "total_registros": total_registros,
            "partes_unicas": partes_unicas,
            "materiales_unicos": materiales_unicos
        }
    )


@app.get("/actualizar-boms")
async def actualizar_boms(request: Request, current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    """Página Actualizar BOMs - requiere autenticación."""
    from app.db.models import Parte, Bom, BomRevision
    from sqlalchemy import select, func, distinct
    # Contar partes y última actualización real desde tabla partes
    total_partes = await db.execute(select(func.count()).select_from(Parte))
    cantidad_numeros_parte = total_partes.scalar() or 0
    ultima = await db.execute(select(func.max(Parte.created_at)).select_from(Parte))
    ultima_ts = ultima.scalar()
    ultima_actualizacion = ultima_ts.strftime("%d/%m/%Y %H:%M") if ultima_ts else "—"

    # Estadística: partes que ya tienen al menos una revisión de BOM
    partes_con_bom_q = await db.execute(
        select(func.count(distinct(Parte.id)))
        .select_from(Parte)
        .join(Bom, Bom.parte_id == Parte.id)
        .join(BomRevision, BomRevision.bom_id == Bom.id)
    )
    numeros_parte_con_bom = partes_con_bom_q.scalar() or 0

    # Estadística: partes no válidas
    partes_no_validas_q = await db.execute(
        select(func.count()).select_from(Parte).where(Parte.valido.is_(False))
    )
    numeros_parte_no_validos = partes_no_validas_q.scalar() or 0

    # Estadísticas por diferencia (solo partes con BOM/revisión)
    bom_count_sq = (
        select(func.count(BomRevision.id))
        .select_from(BomRevision)
        .join(Bom, Bom.id == BomRevision.bom_id)
        .where(Bom.parte_id == Parte.id)
        .correlate(Parte)
        .scalar_subquery()
    )
    diff_gt0_q = await db.execute(
        select(func.count())
        .select_from(Parte)
        .where(
            bom_count_sq > 0,
            Parte.diferencia.is_not(None),
            Parte.diferencia > 0,
        )
    )
    numeros_parte_diferencia_gt0 = diff_gt0_q.scalar() or 0

    diff_le0_q = await db.execute(
        select(func.count())
        .select_from(Parte)
        .where(
            bom_count_sq > 0,
            Parte.diferencia.is_not(None),
            Parte.diferencia <= 0,
        )
    )
    numeros_parte_diferencia_le0 = diff_le0_q.scalar() or 0

    diff_na_q = await db.execute(
        select(func.count())
        .select_from(Parte)
        .where(
            bom_count_sq > 0,
            Parte.diferencia.is_(None),
        )
    )
    numeros_parte_diferencia_na = diff_na_q.scalar() or 0

    return templates.TemplateResponse(
        "actualizar_boms.html",
        {
            "request": request,
            "active_page": "actualizar_boms",
            "current_user": current_user,
            "cantidad_numeros_parte": cantidad_numeros_parte,
            "ultima_actualizacion": ultima_actualizacion,
            "numeros_parte_con_bom": numeros_parte_con_bom,
            "numeros_parte_no_validos": numeros_parte_no_validos,
            "numeros_parte_diferencia_gt0": numeros_parte_diferencia_gt0,
            "numeros_parte_diferencia_le0": numeros_parte_diferencia_le0,
            "numeros_parte_diferencia_na": numeros_parte_diferencia_na,
        }
    )


@app.get("/api/actualizar-boms/descargar-reporte")
async def api_descargar_reporte_actualizar_boms(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Descarga reporte BOM por componente para partes con diferencia <= 0."""
    import io
    from sqlalchemy import select
    from sqlalchemy.orm import aliased
    from app.db.models import Parte, Bom, BomRevision, BomItem, PesoNeto
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    headers = [
        "Producto",
        "Insumo",
        "Incorporado",
        "Desperdicio",
        "Merma",
        "Fecha Inicial",
        "Fecha Final",
        "Habilitado",
        "Sustituto",
    ]

    producto = aliased(Parte)
    insumo = aliased(Parte)
    fecha_descarga = datetime.now().date()

    query = (
        select(
            producto.numero_parte.label("producto"),
            insumo.numero_parte.label("insumo"),
            BomItem.qty.label("qty"),
            (producto.qty_total - PesoNeto.kgm).label("diferencia"),
        )
        .select_from(BomItem)
        .join(BomRevision, BomRevision.id == BomItem.bom_revision_id)
        .join(Bom, Bom.id == BomRevision.bom_id)
        .join(producto, producto.id == Bom.parte_id)
        .join(insumo, insumo.id == BomItem.componente_id)
        .join(PesoNeto, PesoNeto.numero_parte == producto.numero_parte)
        .where(
            BomRevision.effective_to.is_(None),
            (producto.qty_total - PesoNeto.kgm) <= 0,
        )
        .order_by(producto.numero_parte, insumo.numero_parte)
    )
    result = await db.execute(query)
    insumos_ajuste = ("3012", "3013", "3014", "3015", "3016", "3017")

    wb = Workbook()
    ws = wb.active
    ws.title = "Reporte"

    purple_fill = PatternFill(fill_type="solid", start_color="7030A0", end_color="7030A0")
    header_font = Font(bold=True, color="FFFFFF")

    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = purple_fill
        ws.column_dimensions[cell.column_letter].width = max(14, len(header) + 2)

    row_idx = 2
    producto_actual = None
    ajuste_actual = None

    def _append_insumos_ajuste(producto_no: str, incorporado_ajuste: float) -> int:
        nonlocal row_idx
        if not producto_no or incorporado_ajuste is None:
            return 0
        agregados = 0
        for insumo_ajuste in insumos_ajuste:
            ws.cell(row=row_idx, column=1, value=producto_no)
            ws.cell(row=row_idx, column=2, value=insumo_ajuste)
            ws.cell(row=row_idx, column=3, value=incorporado_ajuste)
            ws.cell(row=row_idx, column=4, value=0)
            ws.cell(row=row_idx, column=5, value=incorporado_ajuste * 0.03)
            ws.cell(row=row_idx, column=6, value=fecha_descarga)
            ws.cell(row=row_idx, column=7, value="")
            ws.cell(row=row_idx, column=8, value=1)
            ws.cell(row=row_idx, column=9, value="")
            row_idx += 1
            agregados += 1
        return agregados

    for producto_no, insumo_no, qty, diferencia in result.all():
        if producto_actual is not None and producto_no != producto_actual:
            _append_insumos_ajuste(producto_actual, ajuste_actual)
            ajuste_actual = None

        qty_value = float(qty or 0)
        incorporado = qty_value / 1000.0
        ws.cell(row=row_idx, column=1, value=producto_no or "")
        ws.cell(row=row_idx, column=2, value=insumo_no or "")
        ws.cell(row=row_idx, column=3, value=incorporado)
        ws.cell(row=row_idx, column=4, value=0)
        ws.cell(row=row_idx, column=5, value=incorporado * 0.03)
        ws.cell(row=row_idx, column=6, value=fecha_descarga)
        ws.cell(row=row_idx, column=7, value="")
        ws.cell(row=row_idx, column=8, value=1)
        ws.cell(row=row_idx, column=9, value="")
        row_idx += 1
        producto_actual = producto_no
        if producto_no and diferencia is not None and float(diferencia) < 0:
            ajuste_actual = abs(float(diferencia)) / 6.0

    if producto_actual is not None:
        _append_insumos_ajuste(producto_actual, ajuste_actual)

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": 'attachment; filename="reporte_actualizar_boms.xlsx"'},
    )


@app.get("/api/actualizar-boms/descargar-tablas")
async def api_descargar_tablas_actualizar_boms(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Descarga un Excel con las 4 tablas de BOM: partes, bom, bom_revision y bom_item."""
    import io
    import pandas as pd
    from sqlalchemy import select
    from app.db.models import Parte, Bom, BomRevision, BomItem

    def _to_rows(objs, model):
        cols = [c.name for c in model.__table__.columns]
        rows = []
        for o in objs:
            row = {}
            for c in cols:
                v = getattr(o, c, None)
                if hasattr(v, "isoformat"):
                    row[c] = v.isoformat()
                else:
                    row[c] = v
            rows.append(row)
        return rows, cols

    partes_result = await db.execute(select(Parte).order_by(Parte.id))
    bom_result = await db.execute(select(Bom).order_by(Bom.id))
    revision_result = await db.execute(select(BomRevision).order_by(BomRevision.id))
    item_result = await db.execute(select(BomItem).order_by(BomItem.id))

    partes_rows, partes_cols = _to_rows(partes_result.scalars().all(), Parte)
    bom_rows, bom_cols = _to_rows(bom_result.scalars().all(), Bom)
    revision_rows, revision_cols = _to_rows(revision_result.scalars().all(), BomRevision)
    item_rows, item_cols = _to_rows(item_result.scalars().all(), BomItem)

    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        pd.DataFrame(partes_rows, columns=partes_cols).to_excel(
            writer, index=False, sheet_name="partes"
        )
        pd.DataFrame(bom_rows, columns=bom_cols).to_excel(
            writer, index=False, sheet_name="bom"
        )
        pd.DataFrame(revision_rows, columns=revision_cols).to_excel(
            writer, index=False, sheet_name="bom_revision"
        )
        pd.DataFrame(item_rows, columns=item_cols).to_excel(
            writer, index=False, sheet_name="bom_item"
        )
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": 'attachment; filename="actualizar_boms_tablas.xlsx"'},
    )


@app.post("/api/actualizar-boms/recalcular-diferencia")
async def api_recalcular_diferencia_partes(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Recalcula y persiste diferencia (qty_total - kgm) en la tabla partes."""
    resumen_qty = await crud.recalcular_qty_total_partes(db)
    resumen_diff = await crud.recalcular_diferencia_partes(db)
    await db.commit()
    resumen = {}
    resumen.update(resumen_qty)
    resumen.update(resumen_diff)
    return {
        "ok": True,
        "mensaje": "Qty total recalculada y después diferencia registrada.",
        "resumen": resumen,
    }


@app.post("/api/actualizar-boms/load-bom", response_model=LoadBomResponse)
async def api_load_bom(
    payload: LoadBomInput,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Carga o actualiza un BOM: partes, BOM, revisiones e items. Transaccional."""
    return await load_bom(db, payload)


@app.post("/api/actualizar-boms/actualizar-parte")
async def api_actualizar_bom_parte(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Actualiza desde SAP un solo número de parte."""
    try:
        data = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"ok": False, "mensaje": "JSON inválido."})

    numero_parte = (data.get("parte_no") or "").strip()
    if not numero_parte:
        return JSONResponse(status_code=400, content={"ok": False, "mensaje": "parte_no es requerido."})

    vbs_path = Path(settings.BOM_VBS_PATH or str(Path(__file__).resolve().parent / "bom.vbs")).resolve()
    export_dir_raw = settings.BOM_EXPORT_DIR or ""
    export_dir = str(Path(export_dir_raw).resolve()) if export_dir_raw else ""
    timeout_sec = settings.BOM_VBS_TIMEOUT_SEC

    if not export_dir_raw or not Path(export_dir).is_dir():
        return JSONResponse(
            status_code=400,
            content={"ok": False, "mensaje": "Configura BOM_EXPORT_DIR y asegúrate de que la carpeta exista."},
        )
    if platform.system() != "Windows":
        return JSONResponse(
            status_code=501,
            content={"ok": False, "mensaje": "La ejecución del script SAP (bom.vbs) solo está soportada en Windows."},
        )

    try:
        # Una sola parte: cerrar SAP al terminar
        result = await asyncio.to_thread(
            subprocess.run,
            ["cscript", "//nologo", str(vbs_path), numero_parte, export_dir, "1"],
            capture_output=True,
            timeout=timeout_sec,
            cwd=str(vbs_path.parent),
            text=True,
        )
        if result.returncode != 0:
            err = (result.stderr or "").strip() or (result.stdout or "").strip()
            mensaje = f"Script VBS falló (código {result.returncode}): {err[:250]}"
            try:
                await crud.set_parte_valido(db, numero_parte, False)
                await db.commit()
                mensaje += " Se marcó la parte como no válida."
            except Exception:
                await db.rollback()
            return JSONResponse(status_code=400, content={"ok": False, "mensaje": mensaje})

        archivo = Path(export_dir) / f"{numero_parte}.txt"
        for _ in range(15):
            await asyncio.sleep(1)
            if archivo.is_file():
                break
        if not archivo.is_file():
            return JSONResponse(
                status_code=400,
                content={"ok": False, "mensaje": "No se generó el archivo .txt para esa parte."},
            )

        contenido = archivo.read_text(encoding="utf-8", errors="replace")
        payload = parse_sap_bom_txt(contenido)
        if not payload:
            return JSONResponse(status_code=400, content={"ok": False, "mensaje": "No se pudo parsear el archivo exportado."})

        resp = await load_bom(db, payload)
        if resp.ok:
            await crud.set_parte_valido(db, numero_parte, True)
        await db.commit()

        return {
            "ok": bool(resp.ok),
            "parte_no": numero_parte,
            "estado": "sin_cambios" if getattr(resp, "sin_cambios", False) else ("actualizado" if resp.ok else "error"),
            "mensaje": resp.mensaje,
            "items_insertados": getattr(resp, "items_insertados", None),
        }
    except subprocess.TimeoutExpired:
        await db.rollback()
        return JSONResponse(status_code=408, content={"ok": False, "mensaje": f"Timeout ({timeout_sec}s) ejecutando SAP."})
    except Exception as e:
        await db.rollback()
        return JSONResponse(status_code=500, content={"ok": False, "mensaje": f"Error actualizando parte: {e}"})


@app.post("/api/actualizar-boms/ejecutar-actualizacion")
async def api_ejecutar_actualizacion_boms(
    limit: Optional[int] = None,
    reset_all: bool = False,
    only_invalid: bool = False,
    only_diff_gt0: bool = False,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Para cada numero_parte activo que ya tiene BOM: ejecuta bom.vbs (SAP), lee el .txt,
    parsea y actualiza el BOM en la base de datos.
    Requiere Windows, SAP GUI abierto y BOM_EXPORT_DIR configurado.
    limit: opcional; si se pasa, solo procesa los primeros N partes (útil para pruebas).
    """
    vbs_path = Path(settings.BOM_VBS_PATH or str(Path(__file__).resolve().parent / "bom.vbs")).resolve()
    export_dir_raw = settings.BOM_EXPORT_DIR or ""
    export_dir = str(Path(export_dir_raw).resolve()) if export_dir_raw else ""
    timeout_sec = settings.BOM_VBS_TIMEOUT_SEC
    if not export_dir_raw or not Path(export_dir).is_dir():
        return JSONResponse(
            status_code=400,
            content={
                "ok": False,
                "mensaje": "Configura BOM_EXPORT_DIR (carpeta donde SAP guarda los .txt) y asegúrate de que exista.",
                "total": 0,
                "procesados": 0,
                "con_cambios": 0,
                "sin_cambios": 0,
                "errores": 0,
                "detalle": [],
            },
        )
    if platform.system() != "Windows":
        return JSONResponse(
            status_code=501,
            content={
                "ok": False,
                "mensaje": "La ejecución del script SAP (bom.vbs) solo está soportada en Windows.",
                "total": 0,
                "procesados": 0,
                "con_cambios": 0,
                "sin_cambios": 0,
                "errores": 0,
                "detalle": [],
            },
        )
    reset_info = None
    if reset_all:
        reset_info = await crud.reset_bom_para_reproceso(db)
        await db.commit()
        logger.info("BOM reproceso desde cero solicitado: %s", reset_info)

    if only_invalid and only_diff_gt0:
        return JSONResponse(
            status_code=400,
            content={
                "ok": False,
                "mensaje": "Solo puedes usar uno de estos filtros: only_invalid u only_diff_gt0.",
                "total": 0,
                "procesados": 0,
                "con_cambios": 0,
                "sin_cambios": 0,
                "errores": 0,
                "detalle": [],
            },
        )
    if only_invalid:
        part_numbers = await crud.list_partes_numeros_no_validos(db, limit=limit)
    elif only_diff_gt0:
        part_numbers = await crud.list_partes_numeros_diferencia_gt0(db, limit=limit)
    else:
        part_numbers = await crud.list_partes_numeros(db, limit=limit)
    total = len(part_numbers)
    procesados = 0
    con_cambios = 0
    sin_cambios = 0
    errores = 0
    detalle = []
    for idx_parte, numero_parte in enumerate(part_numbers):
        estado = "error"
        mensaje = ""
        es_ultima_parte = (idx_parte == len(part_numbers) - 1)
        args_bom = ["cscript", "//nologo", str(vbs_path), numero_parte, export_dir]
        if es_ultima_parte:
            args_bom.append("1")  # cerrar SAP al final del ultimo numero de parte
        try:
            # 1) Exportar desde SAP: ejecutar bom.vbs para este numero de parte (genera .txt)
            result = await asyncio.to_thread(
                subprocess.run,
                args_bom,
                capture_output=True,
                timeout=timeout_sec,
                cwd=str(vbs_path.parent),
                text=True,
            )
            # Si el script falló (p. ej. material no encontrado en SAP, o SAP no abierto), marcar parte como no válida y continuar
            if result.returncode != 0:
                err = (result.stderr or "").strip() or (result.stdout or "").strip()
                estado = "error"
                mensaje = f"El script VBS falló (código {result.returncode}). Asegúrate de tener SAP GUI abierto en esta misma sesión. Detalle: {err[:300]}"
                try:
                    await crud.set_parte_valido(db, numero_parte, False)
                    await db.commit()
                    mensaje += " Se marcó la parte como no válida."
                except Exception as mark_err:
                    await db.rollback()
                    mensaje += f" No se pudo actualizar 'valido': {str(mark_err)[:160]}"
                errores += 1
                detalle.append({"parte_no": numero_parte, "estado": estado, "mensaje": mensaje})
                continue
            # Esperar a que aparezca el archivo (SAP puede tardar un poco)
            archivo = Path(export_dir) / f"{numero_parte}.txt"
            for _ in range(15):
                await asyncio.sleep(1)
                if archivo.is_file():
                    break
            if not archivo.is_file():
                estado = "error"
                mensaje = "No se generó el archivo .txt. ¿SAP GUI está abierto en esta PC y la transacción CS13 terminó de exportar?"
                try:
                    await crud.set_parte_valido(db, numero_parte, False)
                    await db.commit()
                    mensaje += " Se marcó la parte como no válida."
                except Exception as mark_err:
                    await db.rollback()
                    mensaje += f" No se pudo actualizar 'valido': {str(mark_err)[:160]}"
                errores += 1
                detalle.append({"parte_no": numero_parte, "estado": estado, "mensaje": mensaje})
                continue
            # 2) Leer y parsear el .txt exportado por SAP
            contenido = archivo.read_text(encoding="utf-8", errors="replace")
            payload = parse_sap_bom_txt(contenido)
            if not payload:
                estado = "error"
                mensaje = "No se pudo parsear el archivo"
                errores += 1
                detalle.append({"parte_no": numero_parte, "estado": estado, "mensaje": mensaje})
                continue
            # 3) Insertar/actualizar en BD (partes, bom, bom_revision, bom_item) ANTES de pasar al siguiente numero de parte
            resp = await load_bom(db, payload)
            if resp.ok:
                await crud.set_parte_valido(db, numero_parte, True)
            await db.commit()
            procesados += 1
            if resp.sin_cambios:
                sin_cambios += 1
                estado = "sin_cambios"
                mensaje = resp.mensaje
            elif resp.ok:
                con_cambios += 1
                estado = "actualizado"
                mensaje = resp.mensaje
            else:
                errores += 1
                estado = "error"
                mensaje = resp.mensaje
            detalle.append({"parte_no": numero_parte, "estado": estado, "mensaje": mensaje})
        except subprocess.TimeoutExpired:
            await db.rollback()
            errores += 1
            estado = "error"
            mensaje = f"Timeout ({timeout_sec}s) ejecutando SAP"
            detalle.append({"parte_no": numero_parte, "estado": estado, "mensaje": mensaje})
        except Exception as e:
            await db.rollback()
            errores += 1
            estado = "error"
            mensaje = str(e)
            detalle.append({"parte_no": numero_parte, "estado": estado, "mensaje": mensaje})
    try:
        await crud.create_bom_historial(
            db=db,
            user_id=getattr(current_user, "id", None),
            estado="SUCCESS",
            detalle=f"Procesados {procesados}/{total}; con cambios: {con_cambios}, sin cambios: {sin_cambios}, errores: {errores}",
            total=total,
            procesados=procesados,
            con_cambios=con_cambios,
            sin_cambios=sin_cambios,
            errores=errores,
            detalle_json=detalle[:200] if detalle else None,
        )
    except Exception:
        await db.rollback()
        logger.exception("[actualizar-boms] No se pudo registrar historial de movimiento")

    payload = {
        "ok": True,
        "mensaje": f"Procesados {procesados}/{total}; con cambios: {con_cambios}, sin cambios: {sin_cambios}, errores: {errores}",
        "total": total,
        "procesados": procesados,
        "con_cambios": con_cambios,
        "sin_cambios": sin_cambios,
        "errores": errores,
        "detalle": detalle,
    }
    if reset_info is not None:
        payload["reset"] = reset_info
    return payload


@app.post("/api/actualizar-boms/ejecutar-actualizacion-stream")
async def api_ejecutar_actualizacion_boms_stream(
    limit: Optional[int] = None,
    reset_all: bool = False,
    only_invalid: bool = False,
    only_diff_gt0: bool = False,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Igual que ejecutar-actualizacion pero devuelve un stream NDJSON: un evento por cada
    numero de parte activo con BOM procesado, para ver en tiempo real si se inserto u omitio.
    """
    vbs_path = Path(settings.BOM_VBS_PATH or str(Path(__file__).resolve().parent / "bom.vbs")).resolve()
    export_dir_raw = settings.BOM_EXPORT_DIR or ""
    export_dir = str(Path(export_dir_raw).resolve()) if export_dir_raw else ""
    timeout_sec = settings.BOM_VBS_TIMEOUT_SEC

    async def generate():
        if not export_dir_raw or not Path(export_dir).is_dir():
            yield json.dumps({"tipo": "error", "mensaje": "Configura BOM_EXPORT_DIR y asegúrate de que la carpeta exista."}) + "\n"
            return
        if platform.system() != "Windows":
            yield json.dumps({"tipo": "error", "mensaje": "Solo soportado en Windows."}) + "\n"
            return
        if reset_all:
            reset_info = await crud.reset_bom_para_reproceso(db)
            await db.commit()
            logger.info("BOM stream reproceso desde cero: %s", reset_info)
            yield json.dumps({
                "tipo": "reset",
                "mensaje": "Reproceso desde cero aplicado (se limpió BOM/revisiones/items y se reactivaron partes).",
                "detalle": reset_info,
            }) + "\n"

        if only_invalid and only_diff_gt0:
            yield json.dumps({"tipo": "error", "mensaje": "Solo puedes usar uno de estos filtros: only_invalid u only_diff_gt0."}) + "\n"
            return
        if only_invalid:
            part_numbers = await crud.list_partes_numeros_no_validos(db, limit=limit)
        elif only_diff_gt0:
            part_numbers = await crud.list_partes_numeros_diferencia_gt0(db, limit=limit)
        else:
            part_numbers = await crud.list_partes_numeros(db, limit=limit)
        total = len(part_numbers)
        logger.info("BOM stream iniciado. Total partes=%s", total)
        procesados = 0
        con_cambios = 0
        sin_cambios = 0
        errores = 0
        detalle = []
        yield json.dumps({"tipo": "inicio", "total": total}) + "\n"
        for idx_parte, numero_parte in enumerate(part_numbers):
            estado = "error"
            mensaje = ""
            items_insertados = None
            es_ultima_parte = (idx_parte == len(part_numbers) - 1)
            args_bom = ["cscript", "//nologo", str(vbs_path), numero_parte, export_dir]
            if es_ultima_parte:
                args_bom.append("1")  # cerrar SAP al final del ultimo numero de parte
            try:
                logger.info("BOM parte inicio: %s", numero_parte)
                yield json.dumps({"tipo": "parte_inicio", "parte_no": numero_parte}) + "\n"
                result = await asyncio.to_thread(
                    subprocess.run,
                    args_bom,
                    capture_output=True,
                    timeout=timeout_sec,
                    cwd=str(vbs_path.parent),
                    text=True,
                )
                if result.returncode != 0:
                    err = (result.stderr or "").strip() or (result.stdout or "").strip()
                    estado = "error"
                    mensaje = f"Script VBS falló: {err[:200]}"
                    try:
                        await crud.set_parte_valido(db, numero_parte, False)
                        await db.commit()
                        mensaje += " Se marcó la parte como no válida."
                    except Exception as mark_err:
                        await db.rollback()
                        mensaje += f" No se pudo actualizar 'valido': {str(mark_err)[:160]}"
                    errores += 1
                    detalle.append({"parte_no": numero_parte, "estado": estado, "mensaje": mensaje})
                    logger.warning("BOM parte error VBS: %s | %s", numero_parte, mensaje)
                    yield json.dumps({"tipo": "parte", "parte_no": numero_parte, "estado": estado, "mensaje": mensaje, "items_insertados": None}) + "\n"
                    continue
                archivo = Path(export_dir) / f"{numero_parte}.txt"
                for _ in range(15):
                    await asyncio.sleep(1)
                    if archivo.is_file():
                        break
                if not archivo.is_file():
                    estado = "error"
                    mensaje = "No se generó el archivo .txt"
                    try:
                        await crud.set_parte_valido(db, numero_parte, False)
                        await db.commit()
                        mensaje += " Se marcó la parte como no válida."
                    except Exception as mark_err:
                        await db.rollback()
                        mensaje += f" No se pudo actualizar 'valido': {str(mark_err)[:160]}"
                    errores += 1
                    detalle.append({"parte_no": numero_parte, "estado": estado, "mensaje": mensaje})
                    logger.warning("BOM parte sin archivo: %s | %s", numero_parte, mensaje)
                    yield json.dumps({"tipo": "parte", "parte_no": numero_parte, "estado": estado, "mensaje": mensaje, "items_insertados": None}) + "\n"
                    continue
                contenido = archivo.read_text(encoding="utf-8", errors="replace")
                payload = parse_sap_bom_txt(contenido)
                if not payload:
                    estado = "error"
                    mensaje = "No se pudo parsear el archivo"
                    errores += 1
                    detalle.append({"parte_no": numero_parte, "estado": estado, "mensaje": mensaje})
                    logger.warning("BOM parte parse error: %s", numero_parte)
                    yield json.dumps({"tipo": "parte", "parte_no": numero_parte, "estado": estado, "mensaje": mensaje, "items_insertados": None}) + "\n"
                    continue
                resp = await load_bom(db, payload)
                if resp.ok:
                    await crud.set_parte_valido(db, numero_parte, True)
                await db.commit()
                procesados += 1
                items_insertados = getattr(resp, "items_insertados", None) or 0
                if resp.sin_cambios:
                    sin_cambios += 1
                    estado = "sin_cambios"
                    mensaje = resp.mensaje
                elif resp.ok:
                    con_cambios += 1
                    estado = "actualizado"
                    mensaje = resp.mensaje
                else:
                    errores += 1
                    estado = "error"
                    mensaje = resp.mensaje
                detalle.append({"parte_no": numero_parte, "estado": estado, "mensaje": mensaje})
                logger.info("BOM parte fin: %s | estado=%s", numero_parte, estado)
                yield json.dumps({"tipo": "parte", "parte_no": numero_parte, "estado": estado, "mensaje": mensaje, "items_insertados": items_insertados}) + "\n"
            except subprocess.TimeoutExpired:
                await db.rollback()
                errores += 1
                estado = "error"
                mensaje = f"Timeout ({timeout_sec}s)"
                detalle.append({"parte_no": numero_parte, "estado": estado, "mensaje": mensaje})
                logger.warning("BOM parte timeout: %s", numero_parte)
                yield json.dumps({"tipo": "parte", "parte_no": numero_parte, "estado": estado, "mensaje": mensaje, "items_insertados": None}) + "\n"
            except Exception as e:
                await db.rollback()
                errores += 1
                estado = "error"
                mensaje = str(e)
                detalle.append({"parte_no": numero_parte, "estado": estado, "mensaje": mensaje})
                logger.exception("BOM parte excepción: %s", numero_parte)
                yield json.dumps({"tipo": "parte", "parte_no": numero_parte, "estado": estado, "mensaje": mensaje, "items_insertados": None}) + "\n"
        logger.info(
            "BOM stream finalizado. total=%s procesados=%s con_cambios=%s sin_cambios=%s errores=%s",
            total, procesados, con_cambios, sin_cambios, errores
        )
        try:
            await crud.create_bom_historial(
                db=db,
                user_id=getattr(current_user, "id", None),
                estado="SUCCESS",
                detalle=f"Procesados {procesados}/{total}; con cambios: {con_cambios}, sin cambios: {sin_cambios}, errores: {errores}",
                total=total,
                procesados=procesados,
                con_cambios=con_cambios,
                sin_cambios=sin_cambios,
                errores=errores,
                detalle_json=detalle[:200] if detalle else None,
            )
        except Exception:
            await db.rollback()
            logger.exception("[actualizar-boms stream] No se pudo registrar historial de movimiento")
        yield json.dumps({
            "tipo": "fin",
            "total": total,
            "procesados": procesados,
            "con_cambios": con_cambios,
            "sin_cambios": sin_cambios,
            "errores": errores,
            "detalle": detalle,
        }) + "\n"

    return StreamingResponse(
        generate(),
        media_type="application/x-ndjson",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        },
    )


@app.get("/api/actualizar-boms/movimientos")
async def api_actualizar_boms_movimientos(
    limit: int = 5,
    current_user: User = Depends(require_roles(["admin"])),
    db: AsyncSession = Depends(get_db),
):
    """API para listar movimientos recientes de actualización de BOMs. Solo administradores."""
    _ = current_user
    limit = max(1, min(limit, 20))
    rows_db = await crud.list_bom_historial(db, limit=limit, offset=0)
    rows = [
        {
            "id": row.id,
            "fecha": row.created_at.isoformat() if row.created_at else None,
            "usuario": (row.user.nombre or row.user.email) if row.user else "Sistema",
            "accion": row.accion,
            "estado": row.estado,
            "total": row.total,
            "procesados": row.procesados,
            "con_cambios": row.con_cambios,
            "sin_cambios": row.sin_cambios,
            "errores": row.errores,
            "detalle": row.detalle,
        }
        for row in rows_db
    ]
    return {"ok": True, "rows": rows}


@app.get("/api/actualizar-boms/tablas")
async def api_actualizar_boms_tablas(
    limit: int = 20,
    parte_no: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Devuelve registros recientes de partes, bom, bom_revision y bom_item, con filtro opcional por numero_parte."""
    from sqlalchemy import select
    from app.db.models import Parte, Bom, BomRevision, BomItem

    limit = max(1, min(limit, 200))
    parte_no = (parte_no or "").strip()

    if parte_no:
        pattern = f"%{parte_no}%"
        res_partes = await db.execute(
            select(Parte)
            .where(Parte.numero_parte.ilike(pattern))
            .order_by(Parte.id.desc())
            .limit(limit)
        )
        partes = [
            {
                "id": p.id,
                "numero_parte": p.numero_parte,
                "descripcion": p.descripcion,
                "valido": p.valido,
                "created_at": p.created_at.isoformat() if p.created_at else None,
            }
            for p in res_partes.scalars().all()
        ]

        res_bom = await db.execute(
            select(Bom, Parte.numero_parte)
            .join(Parte, Parte.id == Bom.parte_id)
            .where(Parte.numero_parte.ilike(pattern))
            .order_by(Bom.id.desc())
            .limit(limit)
        )
        boms = [
            {
                "id": b.id,
                "parte_id": b.parte_id,
                "parte_no": parte_numero,
                "plant": b.plant,
                "usage": b.usage,
                "alternative": b.alternative,
                "base_qty": float(b.base_qty) if b.base_qty is not None else None,
                "reqd_qty": float(b.reqd_qty) if b.reqd_qty is not None else None,
                "base_unit": b.base_unit,
                "created_at": b.created_at.isoformat() if b.created_at else None,
                "updated_at": b.updated_at.isoformat() if b.updated_at else None,
            }
            for b, parte_numero in res_bom.all()
        ]

        res_rev = await db.execute(
            select(BomRevision, Parte.numero_parte)
            .join(Bom, Bom.id == BomRevision.bom_id)
            .join(Parte, Parte.id == Bom.parte_id)
            .where(Parte.numero_parte.ilike(pattern))
            .order_by(BomRevision.id.desc())
            .limit(limit)
        )
        revisiones = [
            {
                "id": r.id,
                "bom_id": r.bom_id,
                "parte_no": parte_numero,
                "revision_no": r.revision_no,
                "effective_from": str(r.effective_from) if r.effective_from else None,
                "effective_to": str(r.effective_to) if r.effective_to else None,
                "source": r.source,
                "hash": r.hash,
                "created_at": r.created_at.isoformat() if r.created_at else None,
            }
            for r, parte_numero in res_rev.all()
        ]

        res_item = await db.execute(
            select(BomItem, Parte.numero_parte)
            .join(BomRevision, BomRevision.id == BomItem.bom_revision_id)
            .join(Bom, Bom.id == BomRevision.bom_id)
            .join(Parte, Parte.id == Bom.parte_id)
            .where(Parte.numero_parte.ilike(pattern))
            .order_by(BomItem.id.desc())
            .limit(limit)
        )
        items = [
            {
                "id": i.id,
                "bom_revision_id": i.bom_revision_id,
                "parte_no": parte_numero,
                "componente_id": i.componente_id,
                "item_no": i.item_no,
                "qty": float(i.qty) if i.qty is not None else None,
                "measure": i.measure,
                "comm_code": i.comm_code,
                "origin": i.origin,
                "created_at": i.created_at.isoformat() if i.created_at else None,
            }
            for i, parte_numero in res_item.all()
        ]
    else:
        res_partes = await db.execute(select(Parte).order_by(Parte.id.desc()).limit(limit))
        res_bom = await db.execute(select(Bom).order_by(Bom.id.desc()).limit(limit))
        res_rev = await db.execute(select(BomRevision).order_by(BomRevision.id.desc()).limit(limit))
        res_item = await db.execute(select(BomItem).order_by(BomItem.id.desc()).limit(limit))

        partes = [
            {
                "id": p.id,
                "numero_parte": p.numero_parte,
                "descripcion": p.descripcion,
                "valido": p.valido,
                "created_at": p.created_at.isoformat() if p.created_at else None,
            }
            for p in res_partes.scalars().all()
        ]
        boms = [
            {
                "id": b.id,
                "parte_id": b.parte_id,
                "parte_no": None,
                "plant": b.plant,
                "usage": b.usage,
                "alternative": b.alternative,
                "base_qty": float(b.base_qty) if b.base_qty is not None else None,
                "reqd_qty": float(b.reqd_qty) if b.reqd_qty is not None else None,
                "base_unit": b.base_unit,
                "created_at": b.created_at.isoformat() if b.created_at else None,
                "updated_at": b.updated_at.isoformat() if b.updated_at else None,
            }
            for b in res_bom.scalars().all()
        ]
        revisiones = [
            {
                "id": r.id,
                "bom_id": r.bom_id,
                "parte_no": None,
                "revision_no": r.revision_no,
                "effective_from": str(r.effective_from) if r.effective_from else None,
                "effective_to": str(r.effective_to) if r.effective_to else None,
                "source": r.source,
                "hash": r.hash,
                "created_at": r.created_at.isoformat() if r.created_at else None,
            }
            for r in res_rev.scalars().all()
        ]
        items = [
            {
                "id": i.id,
                "bom_revision_id": i.bom_revision_id,
                "parte_no": None,
                "componente_id": i.componente_id,
                "item_no": i.item_no,
                "qty": float(i.qty) if i.qty is not None else None,
                "measure": i.measure,
                "comm_code": i.comm_code,
                "origin": i.origin,
                "created_at": i.created_at.isoformat() if i.created_at else None,
            }
            for i in res_item.scalars().all()
        ]

    return {
        "ok": True,
        "limit": limit,
        "parte_no": parte_no or None,
        "partes": partes,
        "bom": boms,
        "bom_revision": revisiones,
        "bom_item": items,
    }


@app.get("/api/actualizar-boms/partes")
async def api_actualizar_boms_partes(
    q: Optional[str] = None,
    limit: int = 50,
    offset: int = 0,
    sort_key: Optional[str] = None,
    sort_direction: str = "desc",
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Lista números de parte para vista resumen con contador de BOMs y revisiones."""
    from sqlalchemy import select, func, asc, desc
    from app.db.models import Parte, Bom, BomRevision

    limit = max(1, min(limit, 200))
    offset = max(0, offset)
    q = (q or "").strip()

    bom_count_sq = (
        select(func.count(Bom.id))
        .where(Bom.parte_id == Parte.id)
        .correlate(Parte)
        .scalar_subquery()
    )
    rev_count_sq = (
        select(func.count(BomRevision.id))
        .select_from(BomRevision)
        .join(Bom, Bom.id == BomRevision.bom_id)
        .where(Bom.parte_id == Parte.id)
        .correlate(Parte)
        .scalar_subquery()
    )
    query = select(
        Parte.numero_parte,
        Parte.descripcion,
        Parte.valido,
        Parte.created_at,
        Parte.qty_total,
        Parte.diferencia,
        bom_count_sq.label("total_boms"),
        rev_count_sq.label("total_revisiones"),
    )
    count_query = select(func.count()).select_from(Parte)
    query = query.where(bom_count_sq > 0)
    count_query = count_query.where(bom_count_sq > 0)
    if q:
        like_q = f"%{q}%"
        query = query.where(
            Parte.numero_parte.ilike(like_q) | Parte.descripcion.ilike(like_q)
        )
        count_query = count_query.where(
            Parte.numero_parte.ilike(like_q) | Parte.descripcion.ilike(like_q)
        )

    sort_key = (sort_key or "").strip().lower()
    sort_direction = (sort_direction or "desc").strip().lower()
    if sort_key == "diferencia":
        direction_fn = asc if sort_direction == "asc" else desc
        query = query.order_by(direction_fn(Parte.diferencia).nulls_last(), Parte.numero_parte)
    else:
        query = query.order_by(Parte.numero_parte)

    query = query.offset(offset).limit(limit)
    result = await db.execute(query)
    total_result = await db.execute(count_query)
    total = total_result.scalar_one()

    rows = []
    for numero_parte, descripcion, valido, created_at, qty_total, diferencia, total_boms, total_revisiones in result.all():
        qty_total_value = qty_total or 0
        rows.append({
            "numero_parte": numero_parte,
            "descripcion": descripcion,
            "valido": bool(valido),
            "created_at": created_at.isoformat() if created_at else None,
            "qty_total": float(qty_total_value),
            "diferencia": None if diferencia is None else float(diferencia),
            "total_boms": int(total_boms or 0),
            "total_revisiones": int(total_revisiones or 0),
        })

    return {
        "ok": True,
        "q": q or None,
        "limit": limit,
        "offset": offset,
        "sort_key": sort_key or None,
        "sort_direction": sort_direction,
        "total": int(total),
        "rows": rows,
    }


@app.get("/api/actualizar-boms/parte-detalle")
async def api_actualizar_boms_parte_detalle(
    parte_no: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Devuelve BOM vigente e historial de revisiones para un número de parte."""
    from sqlalchemy import select
    from app.db.models import Parte, Bom, BomRevision, BomItem

    parte_no = (parte_no or "").strip()
    if not parte_no:
        return JSONResponse(status_code=400, content={"ok": False, "mensaje": "parte_no es requerido"})

    parte = await crud.get_parte_by_numero(db, parte_no)
    if not parte:
        return JSONResponse(status_code=404, content={"ok": False, "mensaje": "Número de parte no encontrado"})

    boms_result = await db.execute(
        select(Bom).where(Bom.parte_id == parte.id).order_by(Bom.updated_at.desc(), Bom.id.desc())
    )
    boms = list(boms_result.scalars().all())

    elementos = []
    historial = []

    for bom in boms:
        rev_vigente = await crud.get_current_bom_revision(db, bom.id)
        if rev_vigente:
            items_result = await db.execute(
                select(BomItem, Parte)
                .outerjoin(Parte, Parte.id == BomItem.componente_id)
                .where(BomItem.bom_revision_id == rev_vigente.id)
                .order_by(BomItem.item_no, BomItem.id)
            )
            for item, componente in items_result.all():
                componente_no = componente.numero_parte if componente else f"ID:{item.componente_id}"
                elementos.append({
                    "bom_id": bom.id,
                    "plant": bom.plant,
                    "usage": bom.usage,
                    "alternative": bom.alternative,
                    "revision_no": rev_vigente.revision_no,
                    "item_no": item.item_no,
                    "componente_no": componente_no,
                    "qty": float(item.qty) if item.qty is not None else None,
                    "measure": item.measure,
                    "comm_code": item.comm_code,
                    "origin": item.origin,
                })

        revisiones = await crud.list_bom_revisiones(db, bom.id, limit=200)
        for rev in revisiones:
            historial.append({
                "bom_id": bom.id,
                "plant": bom.plant,
                "usage": bom.usage,
                "alternative": bom.alternative,
                "revision_no": rev.revision_no,
                "effective_from": str(rev.effective_from) if rev.effective_from else None,
                "effective_to": str(rev.effective_to) if rev.effective_to else None,
                "vigente": rev.effective_to is None,
                "source": rev.source,
                "hash": rev.hash,
                "created_at": rev.created_at.isoformat() if rev.created_at else None,
            })

    historial.sort(key=lambda x: ((x.get("created_at") or ""), x.get("revision_no") or 0), reverse=True)

    return {
        "ok": True,
        "parte": {
            "id": parte.id,
            "numero_parte": parte.numero_parte,
            "descripcion": parte.descripcion,
            "valido": bool(parte.valido),
            "created_at": parte.created_at.isoformat() if parte.created_at else None,
        },
        "elementos_bom": elementos,
        "historial": historial,
    }


@app.get("/api/actualizar-boms/bom-vigente")
async def api_bom_vigente(
    parte_no: str,
    plant: str,
    usage: str,
    alternative: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Obtiene el BOM vigente de una parte (revisión con effective_to IS NULL) y sus items."""
    bom, rev, items = await crud.get_bom_vigente_by_parte(db, parte_no, plant, usage, alternative)
    if bom is None:
        return {"ok": False, "mensaje": "No existe BOM para esa parte/plant/usage/alternative", "bom": None, "revision": None, "items": []}
    items_out = [{"id": i.id, "componente_id": i.componente_id, "item_no": i.item_no, "qty": float(i.qty), "measure": i.measure, "comm_code": i.comm_code, "origin": i.origin} for i in items]
    rev_out = {"id": rev.id, "bom_id": rev.bom_id, "revision_no": rev.revision_no, "effective_from": str(rev.effective_from), "effective_to": str(rev.effective_to) if rev.effective_to else None, "hash": rev.hash} if rev else None
    return {"ok": True, "bom": {"id": bom.id, "parte_id": bom.parte_id, "plant": bom.plant, "usage": bom.usage, "alternative": bom.alternative}, "revision": rev_out, "items": items_out}


@app.get("/api/actualizar-boms/historial-revisiones")
async def api_historial_revisiones(
    bom_id: int,
    limit: int = 100,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Lista el historial de revisiones de un BOM."""
    revisiones = await crud.list_bom_revisiones(db, bom_id, limit=limit)
    return {
        "ok": True,
        "bom_id": bom_id,
        "revisiones": [
            {"id": r.id, "revision_no": r.revision_no, "effective_from": str(r.effective_from), "effective_to": str(r.effective_to) if r.effective_to else None, "hash": r.hash}
            for r in revisiones
        ],
    }


@app.get("/api/actualizar-boms/boms-por-componente")
async def api_boms_por_componente(
    componente_no: str,
    limit: int = 200,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Lista en qué BOMs aparece un componente (por número de parte)."""
    rows = await crud.list_boms_where_componente(db, componente_no, limit=limit)
    return {"ok": True, "componente_no": componente_no, "boms": rows}


@app.get("/api/actualizar-boms/compare-revisiones")
async def api_compare_revisiones(
    revision_id_a: int,
    revision_id_b: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Compara dos revisiones y devuelve diferencias (items de cada una)."""
    rev_a, items_a = await crud.get_bom_revision_with_items(db, revision_id_a)
    rev_b, items_b = await crud.get_bom_revision_with_items(db, revision_id_b)
    if not rev_a or not rev_b:
        return {"ok": False, "mensaje": "Una o ambas revisiones no existen"}
    ids_a = {i.componente_id: (float(i.qty), i.item_no) for i in items_a}
    ids_b = {i.componente_id: (float(i.qty), i.item_no) for i in items_b}
    solo_en_a = [c for c in ids_a if c not in ids_b]
    solo_en_b = [c for c in ids_b if c not in ids_a]
    cambiados = [c for c in ids_a if c in ids_b and (ids_a[c] != ids_b[c])]
    return {
        "ok": True,
        "revision_a": {"id": rev_a.id, "revision_no": rev_a.revision_no},
        "revision_b": {"id": rev_b.id, "revision_no": rev_b.revision_no},
        "solo_en_a": solo_en_a,
        "solo_en_b": solo_en_b,
        "cambiados_qty_o_item_no": cambiados,
    }


@app.get("/proveedores")
async def proveedores(request: Request, current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    """Página de proveedores - requiere autenticación."""
    # Cargar los proveedores desde la base de datos
    proveedores_data = await crud.list_proveedores(db, limit=1000)
    
    # Obtener la fecha de la última compra para cada proveedor (optimizado con una sola consulta)
    from sqlalchemy import select, func, String, cast
    from app.db.models import Compra
    
    # Obtener todas las fechas de última compra agrupadas por proveedor en una sola consulta
    codigos_proveedores = [p.codigo_proveedor for p in proveedores_data]
    if codigos_proveedores:
        # En algunas BD legacy `compras.codigo_proveedor` puede estar como VARCHAR;
        # comparamos como texto para evitar error varchar = integer.
        codigos_proveedores_str = [str(c).strip() for c in codigos_proveedores if c is not None]
        codigo_proveedor_compra_txt = cast(Compra.codigo_proveedor, String)
        query_ultimas_compras = select(
            codigo_proveedor_compra_txt.label("codigo_proveedor"),
            func.max(Compra.posting_date).label('fecha_ultima_compra')
        ).where(
            codigo_proveedor_compra_txt.in_(codigos_proveedores_str),
            Compra.posting_date.isnot(None)
        ).group_by(codigo_proveedor_compra_txt)
        
        result = await db.execute(query_ultimas_compras)
        fechas_por_proveedor = {str(row.codigo_proveedor).strip(): row.fecha_ultima_compra for row in result.all()}
    else:
        fechas_por_proveedor = {}
    
    # Agregar la fecha de la última compra a cada proveedor
    for proveedor in proveedores_data:
        proveedor.fecha_ultima_compra = fechas_por_proveedor.get(str(proveedor.codigo_proveedor).strip())
    
    # Calcular estadísticas
    total_proveedores = await crud.count_proveedores(db)
    proveedores_activos = await crud.count_proveedores(db, estatus=True)
    proveedores_inactivos = await crud.count_proveedores(db, estatus=False)
    
    # Historial de movimientos solo para administradores
    historial_reciente = await crud.list_proveedor_historial(db, limit=10, offset=0) if current_user.rol == "admin" else []
    
    return templates.TemplateResponse(
        "proveedores.html",
        {
            "request": request,
            "active_page": "proveedores",
            "current_user": current_user,
            "proveedores": proveedores_data,
            "total_proveedores": total_proveedores,
            "proveedores_activos": proveedores_activos,
            "proveedores_inactivos": proveedores_inactivos,
            "historial_reciente": historial_reciente
        }
    )


@app.get("/materiales")
async def materiales(request: Request, current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    """Página de materiales - requiere autenticación."""
    # Cargar los materiales desde la base de datos
    materiales_data = await crud.list_materiales(db, limit=1000)
    
    # Calcular estadísticas
    total_materiales = await crud.count_materiales(db)
    
    # Historial de movimientos solo para administradores
    historial_reciente = await crud.list_material_historial(db, limit=10, offset=0) if current_user.rol == "admin" else []
    
    return templates.TemplateResponse(
        "materiales.html",
        {
            "request": request,
            "active_page": "materiales",
            "current_user": current_user,
            "materiales": materiales_data,
            "total_materiales": total_materiales,
            "historial_reciente": historial_reciente
        }
    )


@app.get("/pesos-netos")
async def pesos_netos(request: Request, current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    """Página de pesos netos - requiere autenticación."""
    total_pesos_netos = await crud.count_pesos_netos(db)
    return templates.TemplateResponse(
        "pesos_netos.html",
        {
            "request": request,
            "active_page": "pesos_netos",
            "current_user": current_user,
            "total_pesos_netos": total_pesos_netos,
        },
    )


@app.get("/api/pesos-netos")
async def api_pesos_netos(
    q: Optional[str] = None,
    limit: int = 50,
    offset: int = 0,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """API para listar pesos netos con búsqueda y paginación."""
    limit = max(1, min(limit, 200))
    offset = max(0, offset)
    q = (q or "").strip()

    rows_db = await crud.list_pesos_netos(db, search=q or None, limit=limit, offset=offset)
    total = await crud.count_pesos_netos(db, search=q or None)

    rows = [
        {
            "numero_parte": row.numero_parte,
            "descripcion": row.descripcion,
            "gross": float(row.gross) if row.gross is not None else None,
            "net": float(row.net) if row.net is not None else None,
            "kgm": float(row.kgm) if row.kgm is not None else None,
            "created_at": row.created_at.isoformat() if row.created_at else None,
            "updated_at": row.updated_at.isoformat() if row.updated_at else None,
        }
        for row in rows_db
    ]

    return {
        "ok": True,
        "q": q or None,
        "limit": limit,
        "offset": offset,
        "total": int(total),
        "rows": rows,
    }


@app.get("/api/pesos-netos/movimientos")
async def api_pesos_netos_movimientos(
    limit: int = 5,
    current_user: User = Depends(require_roles(["admin"])),
    db: AsyncSession = Depends(get_db),
):
    """API para listar movimientos recientes de actualización de pesos netos. Solo administradores."""
    _ = current_user
    limit = max(1, min(limit, 20))
    rows_db = await crud.list_peso_neto_historial(db, limit=limit, offset=0)
    rows = [
        {
            "id": row.id,
            "fecha": row.created_at.isoformat() if row.created_at else None,
            "usuario": (row.user.nombre or row.user.email) if row.user else "Sistema",
            "accion": row.accion,
            "estado": row.estado,
            "archivo_nombre": row.archivo_nombre,
            "filas_archivo": row.filas_archivo,
            "filas_invalidas": row.filas_invalidas,
            "duplicados_archivo": row.duplicados_archivo,
            "candidatos_unicos": row.candidatos_unicos,
            "upserts": row.upserts,
            "insertados": row.insertados,
            "actualizados": row.actualizados,
            "detalle": row.detalle,
        }
        for row in rows_db
    ]
    return {"ok": True, "rows": rows}


def _peso_neto_normalize_col(name) -> str:
    import unicodedata

    if name is None:
        return ""
    s = str(name).strip().lower().replace("\ufeff", "")
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    return " ".join(s.split())


def _peso_neto_to_decimal(value):
    import math
    from decimal import Decimal, InvalidOperation

    if value is None:
        return None
    if isinstance(value, Decimal):
        return value
    if isinstance(value, int):
        return Decimal(value)
    if isinstance(value, float):
        if math.isnan(value):
            return None
        return Decimal(str(value))

    s = str(value).strip()
    if not s or s.lower() == "nan":
        return None
    s = s.replace(" ", "")
    if "," in s and "." in s:
        s = s.replace(",", "")
    elif "," in s:
        s = s.replace(",", ".")

    try:
        return Decimal(s)
    except (InvalidOperation, ValueError):
        return None


def _peso_neto_load_source_dataframe(file_path: Path):
    import pandas as pd

    errors: list[str] = []

    for enc in ("utf-16", "utf-16le", "utf-8-sig", "cp1252", "latin1"):
        try:
            df = pd.read_csv(file_path, sep="\t", dtype=object, encoding=enc)
            if not df.empty and len(df.columns) >= 3:
                return df
        except Exception as exc:
            errors.append(f"read_csv({enc}): {exc}")

    try:
        df = pd.read_excel(file_path, dtype=object)
        if not df.empty:
            return df
    except Exception as exc:
        errors.append(f"read_excel: {exc}")

    try:
        tables = pd.read_html(file_path)
        if tables:
            return sorted(tables, key=lambda t: len(t), reverse=True)[0]
    except Exception as exc:
        errors.append(f"read_html: {exc}")

    raise ValueError(
        "No se pudo leer el archivo de pesos netos como TSV, Excel o HTML.\n"
        + "\n".join(errors[-5:])
    )


def _peso_neto_realign_header(df):
    """Realinea encabezado cuando el export de SAP trae filas en blanco o header corrido."""
    if df is None or df.empty:
        return df

    normalized_cols = {_peso_neto_normalize_col(c) for c in df.columns}
    if "material" in normalized_cols and ("net weight" in normalized_cols or "peso neto" in normalized_cols):
        return df

    raw = df.copy()
    try:
        raw.columns = [f"c{i}" for i in range(len(raw.columns))]
    except Exception:
        return df

    max_scan = min(len(raw), 15)
    header_idx = None
    for i in range(max_scan):
        row_vals = [str(v).strip() for v in raw.iloc[i].tolist() if str(v).strip() and str(v).strip().lower() != "nan"]
        row_norm = {_peso_neto_normalize_col(v) for v in row_vals}
        if "material" in row_norm and ("net weight" in row_norm or "peso neto" in row_norm):
            header_idx = i
            break

    if header_idx is None:
        return df

    new_cols = []
    for idx, v in enumerate(raw.iloc[header_idx].tolist()):
        col_name = str(v).strip()
        if not col_name or col_name.lower() == "nan":
            col_name = f"unnamed_{idx}"
        new_cols.append(col_name)

    data = raw.iloc[header_idx + 1 :].copy()
    data.columns = new_cols
    data = data.reset_index(drop=True)

    # Limpia filas completamente vacías.
    data = data.dropna(how="all")
    return data


def _peso_neto_parse_records(file_path: Path) -> tuple[list[dict], dict]:
    from decimal import Decimal

    df = _peso_neto_realign_header(_peso_neto_load_source_dataframe(file_path))
    if df.empty:
        raise ValueError("El archivo exportado no contiene filas.")

    col_map = {_peso_neto_normalize_col(c): c for c in df.columns}
    col_numero = (
        col_map.get("material")
        or col_map.get("numeros de parte")
        or col_map.get("numero de parte")
        or col_map.get("numero_parte")
        or col_map.get("part_no")
    )
    col_descripcion = (
        col_map.get("descripcion del material")
        or col_map.get("material description")
        or col_map.get("descripcion")
    )
    col_gross = col_map.get("peso bruto") or col_map.get("gross weight") or col_map.get("gross")
    col_net = col_map.get("peso neto") or col_map.get("net weight") or col_map.get("net")

    if not col_numero:
        raise ValueError("No se encontró la columna de número de parte (ej. 'Material').")
    if not col_net:
        raise ValueError("No se encontró la columna de peso neto (ej. 'Peso neto').")

    total_filas = len(df)
    filas_invalidas = 0
    duplicados_archivo = 0
    candidatos: list[dict] = []
    idx_by_numero: dict[str, int] = {}

    for _, row in df.iterrows():
        numero_raw = row.get(col_numero)
        numero = str(numero_raw).strip() if numero_raw is not None else ""
        if not numero or numero.lower() == "nan":
            filas_invalidas += 1
            continue
        if numero.endswith(".0"):
            numero = numero[:-2]

        descripcion = (
            str(row.get(col_descripcion)).strip()
            if col_descripcion and row.get(col_descripcion) is not None
            else None
        )
        gross = _peso_neto_to_decimal(row.get(col_gross)) if col_gross else None
        net = _peso_neto_to_decimal(row.get(col_net))

        if net is None:
            kgm = None
        elif net == 0:
            kgm = Decimal("0")
        else:
            kgm = net / Decimal("1000")

        rec = {
            "numero_parte": numero,
            "descripcion": descripcion,
            "gross": gross,
            "net": net,
            "kgm": kgm,
        }

        if numero in idx_by_numero:
            duplicados_archivo += 1
            candidatos[idx_by_numero[numero]] = rec
        else:
            idx_by_numero[numero] = len(candidatos)
            candidatos.append(rec)

    resumen = {
        "filas_archivo": total_filas,
        "filas_invalidas": filas_invalidas,
        "duplicados_archivo": duplicados_archivo,
        "candidatos_unicos": len(candidatos),
    }
    return candidatos, resumen


@app.post("/api/pesos-netos/actualizar")
async def api_actualizar_pesos_netos_desde_sap(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Ejecuta MM17 (peso_neto.vbs), procesa el archivo exportado y hace upsert en peso_neto.
    """
    from sqlalchemy import select, text
    from app.db.models import PesoNeto

    started_at = datetime.now()
    vbs_path = Path(
        settings.PESO_NETO_VBS_PATH or str(Path(__file__).resolve().parent / "peso_neto.vbs")
    ).resolve()
    export_dir = Path(settings.PESO_NETO_EXPORT_DIR or r"C:\Users\anad5004\Documents\Leoni_RPA\peso_neto").resolve()
    export_file = export_dir / (settings.PESO_NETO_EXPORT_FILENAME or "pesos_netos.xls")
    timeout_sec = settings.PESO_NETO_VBS_TIMEOUT_SEC

    logger.info(
        "[peso_neto] Inicio actualización usuario=%s vbs=%s export=%s",
        getattr(current_user, "email", None),
        vbs_path,
        export_file,
    )

    async def registrar_movimiento(
        estado: str,
        detalle: str,
        resumen: Optional[dict] = None,
    ) -> None:
        resumen = resumen or {}
        try:
            await crud.create_peso_neto_historial(
                db=db,
                user_id=getattr(current_user, "id", None),
                estado=estado,
                detalle=detalle,
                archivo_nombre=export_file.name,
                filas_archivo=resumen.get("filas_archivo"),
                filas_invalidas=resumen.get("filas_invalidas"),
                duplicados_archivo=resumen.get("duplicados_archivo"),
                candidatos_unicos=resumen.get("candidatos_unicos"),
                upserts=resumen.get("upserts"),
                insertados=resumen.get("insertados"),
                actualizados=resumen.get("actualizados"),
            )
        except Exception:
            await db.rollback()
            logger.exception("[peso_neto] No se pudo registrar historial de movimiento")

    if platform.system() != "Windows":
        logger.warning("[peso_neto] Cancelado: plataforma no Windows (%s)", platform.system())
        await registrar_movimiento(
            estado="CANCELLED",
            detalle=f"Cancelado: plataforma no Windows ({platform.system()}).",
        )
        return JSONResponse(
            status_code=501,
            content={"ok": False, "mensaje": "La ejecución de peso_neto.vbs solo está soportada en Windows."},
        )
    if not vbs_path.is_file():
        logger.warning("[peso_neto] Cancelado: no existe VBS (%s)", vbs_path)
        await registrar_movimiento(
            estado="FAILED",
            detalle=f"No se encontró el script VBS: {vbs_path}",
        )
        return JSONResponse(
            status_code=400,
            content={"ok": False, "mensaje": f"No se encontró el script VBS: {vbs_path}"},
        )
    try:
        export_dir.mkdir(parents=True, exist_ok=True)
    except Exception as mk_err:
        logger.exception("[peso_neto] Error creando carpeta de exportación")
        await registrar_movimiento(
            estado="FAILED",
            detalle=f"No se pudo preparar la carpeta de exportación: {mk_err}",
        )
        return JSONResponse(
            status_code=400,
            content={"ok": False, "mensaje": f"No se pudo preparar la carpeta de exportación: {mk_err}"},
        )

    mtime_anterior = export_file.stat().st_mtime if export_file.exists() else None
    try:
        result = await asyncio.to_thread(
            subprocess.run,
            ["cscript", "//nologo", str(vbs_path), str(export_dir)],
            capture_output=True,
            timeout=timeout_sec,
            cwd=str(vbs_path.parent),
            text=True,
        )
    except subprocess.TimeoutExpired:
        logger.warning("[peso_neto] Timeout ejecutando VBS (%ss)", timeout_sec)
        await registrar_movimiento(
            estado="FAILED",
            detalle=f"Timeout ejecutando MM17 ({timeout_sec}s).",
        )
        return JSONResponse(
            status_code=504,
            content={"ok": False, "mensaje": f"Timeout ejecutando MM17 ({timeout_sec}s)."},
        )
    except Exception as e:
        logger.exception("[peso_neto] Error ejecutando VBS")
        await registrar_movimiento(
            estado="FAILED",
            detalle=f"Error al ejecutar peso_neto.vbs: {str(e)}",
        )
        return JSONResponse(
            status_code=500,
            content={"ok": False, "mensaje": f"Error al ejecutar peso_neto.vbs: {str(e)}"},
        )

    if result.returncode != 0:
        err = (result.stderr or "").strip() or (result.stdout or "").strip()
        logger.warning("[peso_neto] VBS returncode=%s detalle=%s", result.returncode, err[:300])
        await registrar_movimiento(
            estado="FAILED",
            detalle=f"VBS falló (code {result.returncode}): {err[:400]}",
        )
        return JSONResponse(
            status_code=500,
            content={"ok": False, "mensaje": f"VBS falló: {err[:400]}"},
        )

    for _ in range(45):
        await asyncio.sleep(1)
        if not export_file.is_file():
            continue
        if mtime_anterior is None:
            break
        if export_file.stat().st_mtime > mtime_anterior:
            break
    else:
        logger.warning("[peso_neto] No se detectó archivo actualizado: %s", export_file)
        await registrar_movimiento(
            estado="FAILED",
            detalle=f"No se detectó archivo de salida actualizado: {export_file.name}",
        )
        return JSONResponse(
            status_code=500,
            content={
                "ok": False,
                "mensaje": f"No se detectó archivo de salida actualizado: {export_file.name}",
            },
        )

    try:
        candidatos, resumen_parseo = await asyncio.to_thread(_peso_neto_parse_records, export_file)
        logger.info(
            "[peso_neto] Parseo OK archivo=%s filas=%s invalidas=%s duplicados=%s unicos=%s",
            export_file.name,
            resumen_parseo.get("filas_archivo"),
            resumen_parseo.get("filas_invalidas"),
            resumen_parseo.get("duplicados_archivo"),
            resumen_parseo.get("candidatos_unicos"),
        )
    except Exception as parse_err:
        logger.exception("[peso_neto] Error parseando archivo")
        await registrar_movimiento(
            estado="FAILED",
            detalle=f"Error procesando archivo de pesos netos: {parse_err}",
        )
        return JSONResponse(
            status_code=500,
            content={"ok": False, "mensaje": f"Error procesando archivo de pesos netos: {parse_err}"},
        )

    if not candidatos:
        await registrar_movimiento(
            estado="FAILED",
            detalle="El archivo no contiene registros válidos para importar.",
            resumen=resumen_parseo,
        )
        return JSONResponse(
            status_code=400,
            content={"ok": False, "mensaje": "El archivo no contiene registros válidos para importar."},
        )

    upsert_sql = text(
        """
        INSERT INTO peso_neto (numero_parte, descripcion, gross, net, kgm, updated_at)
        VALUES (:numero_parte, :descripcion, :gross, :net, :kgm, now())
        ON CONFLICT (numero_parte) DO UPDATE
        SET
            descripcion = EXCLUDED.descripcion,
            gross = EXCLUDED.gross,
            net = EXCLUDED.net,
            kgm = EXCLUDED.kgm,
            updated_at = now()
        """
    )

    try:
        chunk_size = 2000
        existentes: set[str] = set()
        numeros = [str(r["numero_parte"]) for r in candidatos]
        for i in range(0, len(numeros), chunk_size):
            chunk_numeros = numeros[i : i + chunk_size]
            existing_rows = await db.execute(
                select(PesoNeto.numero_parte).where(PesoNeto.numero_parte.in_(chunk_numeros))
            )
            existentes.update(str(v) for v in existing_rows.scalars().all())

        insertados = len(candidatos) - len(existentes)
        actualizados = len(existentes)

        chunk_size = 2000
        upserts = 0
        for i in range(0, len(candidatos), chunk_size):
            chunk = candidatos[i : i + chunk_size]
            await db.execute(upsert_sql, chunk)
            upserts += len(chunk)
        await db.commit()
        logger.info(
            "[peso_neto] Fin OK archivo=%s upserts=%s insertados=%s actualizados=%s",
            export_file.name,
            upserts,
            insertados,
            actualizados,
        )
    except Exception as db_err:
        await db.rollback()
        logger.exception("[peso_neto] Error guardando en BD")
        await registrar_movimiento(
            estado="FAILED",
            detalle=f"Error guardando pesos netos en BD: {db_err}",
            resumen=resumen_parseo,
        )
        return JSONResponse(
            status_code=500,
            content={"ok": False, "mensaje": f"Error guardando pesos netos en BD: {db_err}"},
        )

    resumen_final = {
        **resumen_parseo,
        "upserts": upserts,
        "insertados": insertados,
        "actualizados": actualizados,
    }
    duracion_segundos = max(0, int((datetime.now() - started_at).total_seconds()))
    detalle_historial = (
        f"Proceso completado. Archivo: {export_file.name}. "
        f"Filas archivo: {resumen_final.get('filas_archivo', 'N/A')}. "
        f"Filas inválidas: {resumen_final.get('filas_invalidas', 'N/A')}. "
        f"Duplicados en archivo: {resumen_final.get('duplicados_archivo', 'N/A')}. "
        f"Candidatos únicos: {resumen_final.get('candidatos_unicos', 'N/A')}. "
        f"Upserts: {resumen_final.get('upserts', 'N/A')}. "
        f"Insertados: {resumen_final.get('insertados', 'N/A')}. "
        f"Actualizados: {resumen_final.get('actualizados', 'N/A')}. "
        f"Duración: {duracion_segundos}s."
    )
    await registrar_movimiento(
        estado="SUCCESS",
        detalle=detalle_historial,
        resumen=resumen_final,
    )

    return {
        "ok": True,
        "mensaje": (
            f"Pesos netos actualizados correctamente. "
            f"Archivo procesado: {export_file.name}. "
            f"Insertados: {insertados}. Actualizados: {actualizados}."
        ),
        "resumen": {
            **resumen_final,
            "archivo": str(export_file),
        },
    }


@app.get("/cross-reference")
async def cross_reference(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Página de cross reference - requiere autenticación."""
    total_cross_reference = await crud.count_cross_reference(db)
    return templates.TemplateResponse(
        "cross_reference.html",
        {
            "request": request,
            "active_page": "cross_reference",
            "current_user": current_user,
            "total_cross_reference": total_cross_reference,
        },
    )


def _parse_cross_reference_txt(content: str) -> list[dict]:
    """
    Parsea export TXT de VD59 y extrae:
    - customer
    - material
    - customer_material (Customer Material Number)
    """
    import re

    lines = [line.rstrip("\r\n") for line in content.splitlines()]
    header_idx = None
    header_parts = []
    col_customer = None
    col_material = None
    col_customer_material = None

    for i, line in enumerate(lines):
        if "|" not in line:
            continue
        if "Customer Material Number" in line and "Material" in line and "Customer" in line:
            header_idx = i
            header_parts = [p.strip() for p in line.split("|")]
            break

    if header_idx is None:
        return []

    for i, h in enumerate(header_parts):
        key = " ".join(h.strip().lower().split())
        key = re.sub(r"[^a-z0-9 ]+", "", key)
        if key == "customer":
            col_customer = i
        elif key == "material":
            col_material = i
        elif key in ("customer material number", "customer material"):
            col_customer_material = i

    if col_customer is None or col_material is None or col_customer_material is None:
        return []

    rows = []
    for line in lines[header_idx + 1:]:
        if not line.startswith("|") or "---" in line:
            continue
        parts = [p.strip() for p in line.split("|")]
        max_idx = max(col_customer, col_material, col_customer_material)
        if len(parts) <= max_idx:
            continue
        customer = parts[col_customer].strip()
        material = parts[col_material].strip()
        customer_material = parts[col_customer_material].strip()
        if not customer or not material or not customer_material:
            continue
        rows.append(
            {
                "customer": customer,
                "material": material,
                "customer_material": customer_material,
            }
        )
    return rows


@app.post("/api/cross-reference/actualizar")
async def api_actualizar_cross_reference_desde_sap(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Ejecuta VD59 por cada cliente y hace upsert en cross_reference
    (customer, material, customer_material).
    """
    from sqlalchemy import select, text
    from app.db.models import Cliente

    vbs_path = Path(
        settings.CROSS_REFERENCE_VBS_PATH or str(Path(__file__).resolve().parent / "cross_reference.vbs")
    ).resolve()
    export_dir_raw = settings.CROSS_REFERENCE_EXPORT_DIR or ""
    export_dir = str(Path(export_dir_raw).resolve()) if export_dir_raw else ""
    timeout_sec = settings.CROSS_REFERENCE_VBS_TIMEOUT_SEC

    if not export_dir_raw or not Path(export_dir).is_dir():
        return JSONResponse(
            status_code=400,
            content={"ok": False, "mensaje": "Configura CROSS_REFERENCE_EXPORT_DIR y asegúrate de que exista."},
        )
    if platform.system() != "Windows":
        return JSONResponse(
            status_code=501,
            content={"ok": False, "mensaje": "La ejecución de cross_reference.vbs solo está soportada en Windows."},
        )

    clientes_result = await db.execute(select(Cliente.codigo_cliente).order_by(Cliente.codigo_cliente))
    clientes = [str(c) for c in clientes_result.scalars().all() if c is not None]
    if not clientes:
        try:
            await crud.create_cross_reference_historial(
                db=db,
                user_id=getattr(current_user, "id", None),
                estado="SUCCESS",
                detalle="No hay clientes para procesar.",
                clientes_total=0,
                clientes_ok=0,
                upserts=0,
                errores=0,
            )
        except Exception:
            await db.rollback()
        return {"ok": True, "mensaje": "No hay clientes para procesar.", "resumen": {"clientes": 0, "upserts": 0, "errores": 0}}

    upsert_sql = text(
        """
        INSERT INTO cross_reference (customer, material, customer_material, updated_at)
        VALUES (:customer, :material, :customer_material, now())
        ON CONFLICT (customer, material, customer_material) DO UPDATE
        SET updated_at = now()
        RETURNING (xmax = 0) AS inserted
        """
    )

    total_upserts = 0
    errores = []
    procesados_ok = 0

    for idx_cliente, codigo_cliente in enumerate(clientes):
        es_ultimo_cliente = (idx_cliente == len(clientes) - 1)
        args_vbs = ["cscript", "//nologo", str(vbs_path), codigo_cliente, export_dir]
        if es_ultimo_cliente:
            args_vbs.append("1")  # cerrar SAP al final del ultimo cliente
        try:
            logger.info("[cross_reference] Cliente %s: inicio procesamiento", codigo_cliente)
            result = await asyncio.to_thread(
                subprocess.run,
                args_vbs,
                capture_output=True,
                timeout=timeout_sec,
                cwd=str(vbs_path.parent),
                text=True,
            )
            if result.returncode != 0:
                err = (result.stderr or "").strip() or (result.stdout or "").strip()
                errores.append({"cliente": codigo_cliente, "mensaje": f"VBS falló: {err[:200]}"})
                logger.warning("[cross_reference] Cliente %s: ERROR VBS - %s", codigo_cliente, err[:200])
                continue

            archivo = Path(export_dir) / f"{codigo_cliente}.txt"
            for _ in range(10):
                await asyncio.sleep(1)
                if archivo.is_file():
                    break
            if not archivo.is_file():
                errores.append({"cliente": codigo_cliente, "mensaje": "No se generó TXT de salida."})
                logger.warning("[cross_reference] Cliente %s: ERROR no se generó TXT", codigo_cliente)
                continue

            contenido = archivo.read_text(encoding="utf-8", errors="replace")
            rows = _parse_cross_reference_txt(contenido)
            if not rows:
                errores.append({"cliente": codigo_cliente, "mensaje": "Sin filas parseables en el TXT."})
                logger.warning("[cross_reference] Cliente %s: ERROR sin filas parseables", codigo_cliente)
                continue

            inserted_count = 0
            updated_count = 0
            for row in rows:
                r = await db.execute(upsert_sql, row)
                was_inserted = bool(r.scalar_one())
                if was_inserted:
                    inserted_count += 1
                else:
                    updated_count += 1
                total_upserts += 1
            await db.commit()
            procesados_ok += 1
            logger.info(
                "[cross_reference] Cliente %s: OK filas=%s insertados=%s actualizados=%s",
                codigo_cliente,
                len(rows),
                inserted_count,
                updated_count,
            )
        except subprocess.TimeoutExpired:
            await db.rollback()
            errores.append({"cliente": codigo_cliente, "mensaje": f"Timeout ({timeout_sec}s)."})
            logger.warning("[cross_reference] Cliente %s: ERROR timeout (%ss)", codigo_cliente, timeout_sec)
        except Exception as e:
            await db.rollback()
            errores.append({"cliente": codigo_cliente, "mensaje": str(e)[:240]})
            logger.exception("[cross_reference] Cliente %s: ERROR excepción", codigo_cliente)

    try:
        await crud.create_cross_reference_historial(
            db=db,
            user_id=getattr(current_user, "id", None),
            estado="SUCCESS",
            detalle=f"Clientes OK: {procesados_ok}/{len(clientes)}. Upserts: {total_upserts}. Errores: {len(errores)}.",
            clientes_total=len(clientes),
            clientes_ok=procesados_ok,
            upserts=total_upserts,
            errores=len(errores),
            detalle_errores=errores[:50] if errores else None,
        )
    except Exception:
        await db.rollback()
        logger.exception("[cross_reference] No se pudo registrar historial de movimiento")

    return {
        "ok": True,
        "mensaje": f"Cross Reference actualizado. Clientes OK: {procesados_ok}/{len(clientes)}. Upserts: {total_upserts}. Errores: {len(errores)}.",
        "resumen": {
            "clientes_total": len(clientes),
            "clientes_ok": procesados_ok,
            "upserts": total_upserts,
            "errores": len(errores),
            "detalle_errores": errores[:50],
        },
    }


@app.get("/api/cross-reference")
async def api_cross_reference(
    q: Optional[str] = None,
    limit: int = 50,
    offset: int = 0,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """API para listar cross_reference con búsqueda y paginación."""
    limit = max(1, min(limit, 200))
    offset = max(0, offset)
    q = (q or "").strip()

    rows_db = await crud.list_cross_reference(db, search=q or None, limit=limit, offset=offset)
    total = await crud.count_cross_reference(db, search=q or None)

    rows = [
        {
            "customer": row.customer,
            "material": row.material,
            "customer_material": row.customer_material,
            "created_at": row.created_at.isoformat() if row.created_at else None,
            "updated_at": row.updated_at.isoformat() if row.updated_at else None,
        }
        for row in rows_db
    ]

    return {
        "ok": True,
        "q": q or None,
        "limit": limit,
        "offset": offset,
        "total": int(total),
        "rows": rows,
    }


@app.get("/api/cross-reference/movimientos")
async def api_cross_reference_movimientos(
    limit: int = 5,
    current_user: User = Depends(require_roles(["admin"])),
    db: AsyncSession = Depends(get_db),
):
    """API para listar movimientos recientes de actualización de Cross Reference. Solo administradores."""
    _ = current_user
    limit = max(1, min(limit, 20))
    rows_db = await crud.list_cross_reference_historial(db, limit=limit, offset=0)
    rows = [
        {
            "id": row.id,
            "fecha": row.created_at.isoformat() if row.created_at else None,
            "usuario": (row.user.nombre or row.user.email) if row.user else "Sistema",
            "accion": row.accion,
            "estado": row.estado,
            "clientes_total": row.clientes_total,
            "clientes_ok": row.clientes_ok,
            "upserts": row.upserts,
            "errores": row.errores,
            "detalle": row.detalle,
        }
        for row in rows_db
    ]
    return {"ok": True, "rows": rows}


@app.get("/precios-compra")
async def precios_compra(request: Request, current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    """Página de precios de compra - requiere autenticación."""
    # Cargar los precios de materiales desde la base de datos con relaciones
    precios_data = await crud.list_precios_materiales(db, limit=1000)
    
    # Calcular estadísticas
    total_precios = await crud.count_precios_materiales(db)
    
    # Historial de movimientos solo para administradores
    historial_reciente = await crud.list_precio_material_historial(db, limit=10, offset=0) if current_user.rol == "admin" else []
    
    return templates.TemplateResponse(
        "precios_compra.html",
        {
            "request": request,
            "active_page": "precios_compra",
            "current_user": current_user,
            "precios": precios_data,
            "total_precios": total_precios,
            "historial_reciente": historial_reciente
        }
    )


@app.post("/api/precios-compra/actualizar")
async def actualizar_precios_compra_desde_compras(
    request: Request,
    current_user: User = Depends(require_roles(["admin"])),
    db: AsyncSession = Depends(get_db)
):
    """Sincroniza precios de materiales desde la tabla compras.
    
    Busca todas las compras con numero_material, codigo_proveedor y price válidos
    y crea o actualiza los precios en precios_materiales.
    """
    try:
        resultado = await crud.sincronizar_precios_materiales_desde_compras(
            db=db,
            user_id=current_user.id
        )
        
        # Determinar el mensaje según el resultado
        if resultado["nuevos_creados"] > 0 or resultado["actualizados"] > 0:
            mensaje = (
                f"✓ Sincronización completada. Se crearon {resultado['nuevos_creados']} "
                f"precio(s) nuevo(s) y se actualizaron {resultado['actualizados']} precio(s) "
                f"de {resultado['total_encontrados']} encontrados en compras."
            )
        elif resultado["total_encontrados"] > 0:
            mensaje = (
                f"✓ Sincronización completada. Todos los precios "
                f"({resultado['total_encontrados']}) ya están actualizados."
            )
        else:
            mensaje = "✓ Sincronización completada. No se encontraron datos válidos en compras."
        
        if resultado["errores"]:
            mensaje += f" Se encontraron {len(resultado['errores'])} error(es)."
        
        # Considerar exitoso si no hay errores críticos
        success = len(resultado["errores"]) == 0 or resultado["nuevos_creados"] > 0 or resultado["actualizados"] > 0
        
        return JSONResponse({
            "success": success,
            "total_encontrados": resultado["total_encontrados"],
            "nuevos_creados": resultado["nuevos_creados"],
            "actualizados": resultado["actualizados"],
            "errores": resultado["errores"],
            "mensaje": mensaje
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "error": str(e),
                "mensaje": f"Error al sincronizar precios desde compras: {str(e)}"
            }
        )


@app.get("/api/precios-compra/{precio_id}/ultimas-compras")
async def api_precio_ultimas_compras(
    precio_id: int,
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Devuelve las últimas compras para el material+proveedor de ese precio."""
    from sqlalchemy import select, desc
    from app.db.models import Compra, PrecioMaterial

    precio = await crud.get_precio_material_by_id(db, precio_id)
    if not precio:
        return JSONResponse({"compras": [], "error": "Precio no encontrado"}, status_code=404)

    q = (
        select(
            Compra.posting_date,
            Compra.purchasing_document,
            Compra.numero_material,
            Compra.nombre_proveedor,
            Compra.price,
        )
        .where(
            Compra.codigo_proveedor == precio.codigo_proveedor,
            Compra.numero_material == precio.numero_material,
        )
        .order_by(desc(Compra.posting_date), desc(Compra.id))
        .limit(100)
    )
    result = await db.execute(q)
    rows = result.all()

    items = []
    for r in rows:
        items.append({
            "fecha": r.posting_date.isoformat() if r.posting_date else None,
            "orden_compra": r.purchasing_document,
            "material": r.numero_material,
            "proveedor": r.nombre_proveedor,
            "precio": float(r.price) if r.price is not None else None,
        })
    return JSONResponse({"compras": items})


@app.get("/api/precios-compra/historial")
async def api_precios_compra_historial(
    request: Request,
    current_user: User = Depends(require_roles(["admin"])),
    db: AsyncSession = Depends(get_db)
):
    """Devuelve los últimos 15 movimientos del historial de precios de materiales. Solo administradores."""
    historial = await crud.list_precio_material_historial(
        db,
        limit=15,
        offset=0
    )
    items = []
    for h in historial:
        items.append({
            "id": h.id,
            "precio_material_id": h.precio_material_id,
            "codigo_proveedor": h.codigo_proveedor,
            "numero_material": h.numero_material,
            "operacion": h.operacion.value,
            "user_email": h.user.email if h.user else None,
            "user_nombre": h.user.nombre if h.user else None,
            "datos_antes": h.datos_antes,
            "datos_despues": h.datos_despues,
            "campos_modificados": h.campos_modificados,
            "created_at": h.created_at.isoformat() if h.created_at else None,
        })
    return JSONResponse({"movimientos": items})


@app.post("/api/precios-compra/actualizar-pais-origen")
async def actualizar_pais_origen_precios(
    request: Request,
    current_user: User = Depends(require_roles(["admin"])),
    db: AsyncSession = Depends(get_db)
):
    """Actualiza el país de origen en precios_materiales usando datos de pais_origen_material.
    
    Busca todos los registros en precios_materiales donde country_origin esté vacío o sea null,
    y para cada uno busca en pais_origen_material usando codigo_proveedor y numero_material
    para obtener el pais_origen correspondiente.
    """
    try:
        resultado = await crud.actualizar_pais_origen_en_precios_materiales(
            db=db,
            user_id=current_user.id
        )
        
        # Determinar el mensaje según el resultado
        if resultado["actualizados"] > 0:
            mensaje = (
                f"✓ Actualización completada. Se actualizaron {resultado['actualizados']} "
                f"país(es) de origen de {resultado['total_procesados']} registros procesados."
            )
            if resultado["no_encontrados"] > 0:
                mensaje += f" {resultado['no_encontrados']} registro(s) no tienen país de origen definido en la tabla de países."
        elif resultado["total_procesados"] > 0:
            mensaje = (
                f"✓ Actualización completada. No se requirieron cambios en los "
                f"{resultado['total_procesados']} registros procesados."
            )
            if resultado["no_encontrados"] > 0:
                mensaje += f" {resultado['no_encontrados']} registro(s) no tienen país de origen definido en la tabla de países."
        else:
            mensaje = "✓ No hay registros sin país de origen para actualizar."
        
        if resultado["errores"]:
            mensaje += f" Se encontraron {len(resultado['errores'])} error(es)."
        
        return JSONResponse({
            "success": resultado["exitoso"],
            "mensaje": mensaje,
            "total_procesados": resultado["total_procesados"],
            "actualizados": resultado["actualizados"],
            "no_encontrados": resultado["no_encontrados"],
            "errores": resultado["errores"][:10] if resultado["errores"] else []
        })
        
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "error": str(e),
                "mensaje": f"Error al actualizar países de origen: {str(e)}"
            }
        )


@app.get("/paises-origen")
async def paises_origen(request: Request, current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    """Página de países de origen de materiales - requiere autenticación."""
    # Cargar los países de origen desde la base de datos con relaciones
    paises_data = await crud.list_paises_origen_material(db, limit=1000)
    
    # Calcular estadísticas
    total_paises = await crud.count_paises_origen_material(db)
    
    # Calcular número de partes únicos (numero_material únicos)
    from sqlalchemy import select, func, distinct
    from app.db.models import PaisOrigenMaterial
    
    result = await db.execute(
        select(func.count(distinct(PaisOrigenMaterial.numero_material)))
    )
    total_partes_unicos = result.scalar() or 0
    
    # Calcular materiales con país de origen "Pendiente"
    result_pendientes = await db.execute(
        select(func.count(distinct(PaisOrigenMaterial.numero_material))).where(
            PaisOrigenMaterial.pais_origen == "Pendiente"
        )
    )
    materiales_pendientes = result_pendientes.scalar() or 0
    
    # Historial de movimientos solo para administradores
    historial_reciente = await crud.list_pais_origen_material_historial(db, limit=10, offset=0) if current_user.rol == "admin" else []
    
    return templates.TemplateResponse(
        "paises_origen.html",
        {
            "request": request,
            "active_page": "paises_origen",
            "current_user": current_user,
            "paises": paises_data,
            "total_paises": total_paises,
            "total_partes_unicos": total_partes_unicos,
            "materiales_pendientes": materiales_pendientes,
            "historial_reciente": historial_reciente
        }
    )


@app.get("/carga-proveedor")
async def carga_proveedor(
    request: Request, 
    current_user: User = Depends(get_current_user), 
    db: AsyncSession = Depends(get_db),
    page: int = 1,
    search: str = "",
    estatus: str = ""
):
    """Página de carga de cliente/proveedor - requiere autenticación."""
    limit = 50
    offset = (page - 1) * limit
    
    # Obtener proveedores
    proveedores = await crud.list_carga_proveedores(
        db, 
        limit=limit, 
        offset=offset,
        codigo_proveedor=search if search else None,
        estatus=estatus if estatus else None
    )
    total_proveedores = await crud.count_carga_proveedores(
        db,
        codigo_proveedor=search if search else None,
        estatus=estatus if estatus else None
    )
    
    total_pages = (total_proveedores + limit - 1) // limit
    
    # Historial de movimientos solo para administradores
    if current_user.rol == "admin":
        historial = await crud.list_carga_proveedor_historial(db, limit=20)
        total_historial = await crud.count_carga_proveedor_historial(db)
    else:
        historial = []
        total_historial = 0
    
    # Obtener conteos por estatus para las tarjetas
    count_alta = await crud.count_carga_proveedores(db, estatus="Alta")
    count_baja = await crud.count_carga_proveedores(db, estatus="Baja")
    count_sin_modificacion = await crud.count_carga_proveedores(db, estatus="Sin modificacion")
    count_total = await crud.count_carga_proveedores(db)
    
    return templates.TemplateResponse(
        "carga_cliente_proveedor.html",
        {
            "request": request,
            "active_page": "carga_proveedor",
            "current_user": current_user,
            "proveedores": proveedores,
            "total_proveedores": total_proveedores,
            "page": page,
            "total_pages": total_pages,
            "search": search,
            "estatus": estatus,
            "historial": historial,
            "total_historial": total_historial,
            "count_alta": count_alta,
            "count_baja": count_baja,
            "count_sin_modificacion": count_sin_modificacion,
            "count_total": count_total,
        }
    )


@app.post("/carga-proveedor/actualizar")
async def actualizar_carga_proveedores(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """
    Actualiza los estatus de proveedores según las compras de los últimos 6 meses.
    Solo permite la actualización si no hay registros creados en el historial este mes.
    """
    mes_actual = MESES_VIRTUALES_ES[datetime.now().month - 1]
    count_este_mes = await crud.count_carga_proveedor_historial_este_mes(db)
    if count_este_mes > 0:
        return JSONResponse(
            status_code=200,
            content={
                "success": False,
                "ya_actualizado": True,
                "message": f"Ese mes ({mes_actual}) ya se actualizaron los registros de Carga Proveedor.",
                "mes": mes_actual,
            }
        )
    resultado = await crud.actualizar_estatus_carga_proveedores_por_compras(db)
    
    if resultado["exitoso"]:
        return JSONResponse(
            content={
                "success": True,
                "message": f"Actualización completada. Nuevos: {resultado['proveedores_nuevos']}, Sin modificación: {resultado['proveedores_sin_modificacion']}, Baja: {resultado['proveedores_marcados_baja']}, Omitidos (MX): {resultado.get('proveedores_omitidos_mx', 0)}",
                "data": resultado
            },
            status_code=200
        )
    else:
        return JSONResponse(
            content={
                "success": False,
                "message": f"Error durante la actualización: {resultado.get('error', 'Error desconocido')}",
                "data": resultado
            },
            status_code=500
        )


@app.get("/carga-proveedores-nacional")
async def carga_proveedores_nacional(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    page: int = 1,
    search: str = "",
    estatus: str = "",
    operacion: str = ""
):
    """Vista de carga de proveedores nacionales (Aduanas) - requiere autenticación."""
    limit = 50
    offset = (page - 1) * limit
    registros = await crud.list_carga_proveedores_nacional(
        db,
        limit=limit,
        offset=offset,
        codigo_proveedor=search if search else None,
        estatus=estatus if estatus else None,
        operacion=operacion if operacion else None,
    )
    total = await crud.count_carga_proveedores_nacional(
        db,
        codigo_proveedor=search if search else None,
        estatus=estatus if estatus else None,
        operacion=operacion if operacion else None,
    )
    total_pages = max(1, (total + limit - 1) // limit)
    count_alta = await crud.count_carga_proveedores_nacional(db, estatus="Alta")
    count_baja = await crud.count_carga_proveedores_nacional(db, estatus="Baja")
    count_sin_modificacion = await crud.count_carga_proveedores_nacional(db, estatus="Sin modificacion")
    count_total = await crud.count_carga_proveedores_nacional(db)
    # Historial solo para administradores
    historial = await crud.list_carga_proveedores_nacional_historial(db, limit=50) if current_user.rol == "admin" else []
    total_historial = await crud.count_carga_proveedores_nacional_historial(db) if current_user.rol == "admin" else 0
    return templates.TemplateResponse(
        "carga_proveedores_nacional.html",
        {
            "request": request,
            "active_page": "carga_proveedores_nacional",
            "current_user": current_user,
            "registros": registros,
            "total": total,
            "page": page,
            "total_pages": total_pages,
            "search": search,
            "estatus": estatus,
            "operacion": operacion,
            "count_alta": count_alta,
            "count_baja": count_baja,
            "count_sin_modificacion": count_sin_modificacion,
            "count_total": count_total,
            "historial": historial,
            "total_historial": total_historial,
        }
    )


@app.post("/carga-proveedores-nacional/actualizar")
async def actualizar_carga_proveedores_nacional(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """
    Actualiza los estatus de proveedores nacionales (solo MX) según compras de los últimos 6 meses.
    Solo permite la actualización si no hay registros creados en el historial este mes.
    """
    mes_actual = MESES_VIRTUALES_ES[datetime.now().month - 1]
    count_este_mes = await crud.count_carga_proveedores_nacional_historial_este_mes(db)
    if count_este_mes > 0:
        return JSONResponse(
            status_code=200,
            content={
                "success": False,
                "ya_actualizado": True,
                "message": f"Ese mes ({mes_actual}) ya se actualizaron los registros de Carga Proveedores Nacional.",
                "mes": mes_actual,
            }
        )
    resultado = await crud.actualizar_estatus_carga_proveedores_nacional_por_compras(db)
    if resultado.get("exitoso"):
        return JSONResponse(
            content={
                "success": True,
                "message": f"Actualización completada. Nuevos: {resultado['proveedores_nuevos']}, Sin modificación: {resultado['proveedores_sin_modificacion']}, Baja: {resultado['proveedores_marcados_baja']}, Eliminados: {resultado['proveedores_eliminados']}",
                "data": resultado
            },
            status_code=200
        )
    return JSONResponse(
        content={
            "success": False,
            "message": resultado.get("error", "Error desconocido"),
            "data": resultado
        },
        status_code=500
    )


@app.get("/carga-cliente")
async def carga_cliente(
    request: Request, 
    current_user: User = Depends(get_current_user), 
    db: AsyncSession = Depends(get_db),
    page: int = 1,
    search: str = "",
    estatus: str = ""
):
    """Página de carga de cliente - requiere autenticación."""
    limit = 50
    offset = (page - 1) * limit

    search_param = search.strip() if search else None

    # Obtener clientes (filtro por código o nombre)
    clientes = await crud.list_carga_clientes(
        db,
        limit=limit,
        offset=offset,
        search=search_param,
        estatus=estatus if estatus else None
    )
    total_clientes = await crud.count_carga_clientes(
        db,
        search=search_param,
        estatus=estatus if estatus else None
    )
    
    total_pages = (total_clientes + limit - 1) // limit
    
    # Obtener conteos por estatus para las tarjetas
    count_alta = await crud.count_carga_clientes(db, estatus="Alta")
    count_baja = await crud.count_carga_clientes(db, estatus="Baja")
    count_sin_modificacion = await crud.count_carga_clientes(db, estatus="Sin modificacion")
    count_total = await crud.count_carga_clientes(db)
    
    # Historial de movimientos solo para administradores
    historial = await crud.list_carga_cliente_historial(db, limit=50) if current_user.rol == "admin" else []
    
    return templates.TemplateResponse(
        "carga_cliente.html",
        {
            "request": request,
            "active_page": "carga_cliente",
            "current_user": current_user,
            "clientes": clientes,
            "total_clientes": total_clientes,
            "page": page,
            "total_pages": total_pages,
            "search": search,
            "estatus": estatus,
            "count_alta": count_alta,
            "count_baja": count_baja,
            "count_sin_modificacion": count_sin_modificacion,
            "count_total": count_total,
            "historial": historial,
        }
    )


@app.post("/carga-cliente/actualizar")
async def actualizar_carga_clientes(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """
    Actualiza los estatus de carga_clientes basándose en las ventas.
    Solo permite la actualización si no hay registros creados en el historial este mes.
    """
    mes_actual = MESES_VIRTUALES_ES[datetime.now().month - 1]
    count_este_mes = await crud.count_carga_cliente_historial_este_mes(db)
    if count_este_mes > 0:
        return {
            "success": False,
            "ya_actualizado": True,
            "message": f"Ese mes ({mes_actual}) ya se actualizaron los registros de Carga Cliente.",
            "mes": mes_actual,
        }
    try:
        resumen = await crud.actualizar_carga_clientes_desde_ventas(db)
        
        return {
            "success": True,
            "message": f"Actualización completada: {resumen['altas']} altas, {resumen['bajas']} bajas, {resumen['sin_modificacion']} sin modificación, {resumen['eliminados']} eliminados, {resumen['sin_cambios']} ya actualizados este mes",
            "resumen": resumen
        }
    except Exception as e:
        return {
            "success": False,
            "message": f"Error al actualizar: {str(e)}",
            "resumen": None
        }


@app.post("/carga-cliente/actualizar-domicilio")
async def actualizar_domicilio_carga_clientes(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """
    Completa domicilio y país en carga_clientes usando la tabla clientes.
    """
    try:
        resultado = await crud.actualizar_domicilio_pais_carga_clientes_desde_clientes(db)
        return {
            "success": True,
            "message": (
                f"Actualización de domicilio/pais completada. "
                f"Pendientes: {resultado['total_pendientes']}, "
                f"actualizados: {resultado['actualizados']}, "
                f"sin actualizar: {resultado['sin_actualizar']}."
            ),
            "resultado": resultado,
        }
    except Exception as e:
        return {
            "success": False,
            "message": f"Error al actualizar domicilio/pais: {str(e)}",
            "resultado": None,
        }


@app.get("/carga-proveedor-cliente/descargar-excel")
async def descargar_carga_proveedor_cliente_excel(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Descarga un Excel con una sola hoja: primero proveedores (carga_proveedores) y luego clientes (carga_clientes)."""
    import io
    import pandas as pd

    COLUMNAS = [
        "Nombre o Razón social",
        "Apellido Paterno",
        "Apellido Materno",
        "País",
        "Domicilio",
        "Cliente/Proveedor",
        "Estatus",
    ]

    def _valor(v):
        if v is None:
            return ""
        if hasattr(v, "strftime"):
            return v.strftime("%Y-%m-%d") if hasattr(v, "hour") and v.hour == 0 and v.minute == 0 else v.strftime("%Y-%m-%d %H:%M")
        if isinstance(v, bool):
            return "Sí" if v else "No"
        return str(v)

    # Obtener datos
    proveedores = await crud.list_carga_proveedores(db, limit=50000, offset=0)
    clientes = await crud.list_carga_clientes(db, limit=50000, offset=0)

    filas = []

    # Primero proveedores (carga_proveedores)
    for p in proveedores:
        filas.append({
            "Nombre o Razón social": _valor(p.nombre),
            "Apellido Paterno": _valor(p.apellido_paterno),
            "Apellido Materno": _valor(p.apellido_materno),
            "País": _valor(p.pais),
            "Domicilio": _valor(p.domicilio),
            "Cliente/Proveedor": _valor(p.cliente_proveedor) or "Proveedor",
            "Estatus": _valor(p.estatus),
        })

    # Luego clientes (carga_clientes)
    for c in clientes:
        filas.append({
            "Nombre o Razón social": _valor(c.nombre),
            "Apellido Paterno": "",
            "Apellido Materno": "",
            "País": _valor(c.pais),
            "Domicilio": _valor(c.domicilio),
            "Cliente/Proveedor": _valor(c.cliente_proveedor) or "Cliente",
            "Estatus": _valor(c.estatus),
        })

    df = pd.DataFrame(filas, columns=COLUMNAS)
    buffer = io.BytesIO()
    df.to_excel(buffer, index=False, engine="openpyxl", sheet_name="Carga Proveedores y Clientes")
    buffer.seek(0)

    nombre_archivo = "carga_proveedores_clientes.xlsx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{nombre_archivo}"'},
    )


@app.get("/carga-proveedores-nacional/descargar-excel")
async def descargar_carga_proveedores_nacional_excel(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Descarga un Excel de carga_proveedores_nacional con columnas: RFC, Operaciones Virtuales, Estatus."""
    import io
    import pandas as pd

    COLUMNAS = ["RFC", "Operaciones Virtuales", "Estatus"]

    def _valor(v):
        if v is None:
            return ""
        return str(v)

    registros = await crud.list_carga_proveedores_nacional(db, limit=50000, offset=0)
    filas = [
        {
            "RFC": _valor(p.rfc),
            "Operaciones Virtuales": _valor(p.operacion),
            "Estatus": _valor(p.estatus),
        }
        for p in registros
    ]

    df = pd.DataFrame(filas, columns=COLUMNAS)
    buffer = io.BytesIO()
    df.to_excel(buffer, index=False, engine="openpyxl", sheet_name="Carga Proveedores Nacionales")
    buffer.seek(0)

    nombre_archivo = "carga_proveedores_nacionales.xlsx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{nombre_archivo}"'},
    )


@app.get("/virtuales/guia")
async def virtuales_guia(current_user: User = Depends(get_current_user)):
    """Sirve la guía de materialidad operaciones virtuales (PPTX) para abrir en nueva pestaña."""
    base_dir = os.path.dirname(os.path.abspath(__file__))
    path = os.path.join(base_dir, "MATERIALIDAD OPERACIONES VIRTUALES 2026 (1) 2.pptx")
    if not os.path.isfile(path):
        from fastapi import HTTPException
        raise HTTPException(status_code=404, detail="Guía no encontrada")
    return FileResponse(
        path,
        filename="Materialidad_Operaciones_Virtuales_2026.pptx",
        media_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
    )


@app.get("/virtuales")
async def virtuales(request: Request, current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    """Página Virtuales (master unificado) dentro de Aduanas - requiere autenticación."""
    mes_actual = MESES_VIRTUALES_ES[datetime.now().month - 1]

    # Tabla: todos los registros (sin filtrar por mes)
    virtuales_data = await crud.list_master_unificado_virtuales(db, limit=2000, offset=0)

    # Tarjetas de estadísticas: solo datos del mes actual
    total_virtuales = await crud.count_master_unificado_virtuales(db, mes=mes_actual)
    virtuales_mes_actual = await crud.list_master_unificado_virtuales(db, limit=5000, offset=0, mes=mes_actual)
    from collections import Counter
    estatus_counts = Counter(v.estatus or "Sin estatus" for v in virtuales_mes_actual)
    estatus_counts = dict(sorted(estatus_counts.items(), key=lambda x: x[1], reverse=True))

    # Valores distintos de tipo_exportacion para el filtro (desde toda la tabla)
    tipo_exportacion_opciones = await crud.get_tipo_exportacion_distintos_master_virtuales(db)

    # Historial: administrador ve todos los movimientos; operador solo los suyos
    if current_user.rol == "admin":
        historial_reciente = await crud.list_master_unificado_virtuales_historial(db, limit=10, offset=0)
    elif current_user.rol == "operador":
        historial_reciente = await crud.list_master_unificado_virtuales_historial(db, user_id=current_user.id, limit=10, offset=0)
    else:
        historial_reciente = []

    # Años para el selector de descarga Excel (actual y 5 anteriores)
    año_actual = datetime.now().year
    años_disponibles = list(range(año_actual, año_actual - 6, -1))
    
    return templates.TemplateResponse(
        "virtuales.html",
        {
            "request": request,
            "active_page": "virtuales",
            "current_user": current_user,
            "virtuales": virtuales_data,
            "total_virtuales": total_virtuales,
            "estatus_counts": estatus_counts,
            "tipo_exportacion_opciones": tipo_exportacion_opciones,
            "historial_reciente": historial_reciente,
            "años_disponibles": años_disponibles,
            "mes_actual_nombre": mes_actual,
        }
    )


@app.get("/virtuales/descargar-excel")
async def descargar_virtuales_excel(
    request: Request,
    mes: str = "",
    anio: Optional[int] = None,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Descarga los registros de virtuales del mes y año seleccionados en Excel."""
    import io
    import pandas as pd

    if not mes or not mes.strip():
        return JSONResponse(
            status_code=400,
            content={"error": "Seleccione un mes para descargar."},
        )
    año_actual = datetime.now().year
    if anio is None or anio < 2000 or anio > año_actual + 1:
        anio = año_actual

    mes = mes.strip()
    registros = await crud.list_master_unificado_virtuales(
        db, limit=50000, offset=0, mes=mes, año=anio
    )

    def _valor(v):
        if v is None:
            return ""
        if hasattr(v, "strftime"):
            return v.strftime("%Y-%m-%d") if hasattr(v, "hour") and v.hour == 0 and v.minute == 0 else v.strftime("%Y-%m-%d %H:%M")
        if isinstance(v, bool):
            return "Sí" if v else "No"
        return str(v)

    columnas = [
        "numero", "proveedor_cliente", "impo_expo", "agente", "mes", "estatus", "tipo",
        "incoterm", "tipo_exportacion", "escenario", "materialidad", "plazo", "pedimento", "aduana", "patente",
        "destino", "cliente_space", "complemento", "tipo_immex", "factura", "fecha_pago",
        "informacion", "servicio_cliente", "firma", "solicitud_previo", "op_regular", "carretes",
        "created_at",
    ]
    filas = []
    for r in registros:
        filas.append({
            "numero": _valor(r.numero),
            "proveedor_cliente": _valor(r.proveedor_cliente),
            "impo_expo": _valor(r.impo_expo),
            "agente": _valor(r.agente),
            "mes": _valor(r.mes),
            "estatus": _valor(r.estatus),
            "tipo": _valor(r.tipo),
            "incoterm": _valor(r.incoterm),
            "tipo_exportacion": _valor(r.tipo_exportacion),
            "escenario": _valor(r.escenario),
            "materialidad": _valor(r.materialidad),
            "plazo": _valor(r.plazo),
            "pedimento": _valor(r.pedimento),
            "aduana": _valor(r.aduana),
            "patente": _valor(r.patente),
            "destino": _valor(r.destino),
            "cliente_space": _valor(r.cliente_space),
            "complemento": _valor(r.complemento),
            "tipo_immex": _valor(r.tipo_immex),
            "factura": _valor(r.factura),
            "fecha_pago": _valor(r.fecha_pago),
            "informacion": _valor(r.informacion),
            "servicio_cliente": _valor(r.servicio_cliente),
            "firma": _valor(r.firma),
            "solicitud_previo": _valor(r.solicitud_previo),
            "op_regular": _valor(r.op_regular),
            "carretes": _valor(r.carretes),
            "created_at": _valor(r.created_at),
        })

    df = pd.DataFrame(filas, columns=columnas)
    buffer = io.BytesIO()
    df.to_excel(buffer, index=False, engine="openpyxl")
    buffer.seek(0)
    nombre_archivo = f"virtuales_{mes}_{anio}.xlsx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{nombre_archivo}"'},
    )


@app.get("/virtuales/descargar-plantilla-externa")
async def descargar_plantilla_externa_virtuales(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Descarga plantilla externa de virtuales con columnas base y campos vacíos para captura."""
    import io
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    registros = await crud.list_master_unificado_virtuales(db, limit=50000, offset=0)

    headers = [
        "Numero Cliente/Proveedor",
        "Proveedor_cliente",
        "Impo_Expo",
        "Mes",
        "Pedimento",
        "Aduana",
        "Patente",
        "Firma",
    ]

    wb = Workbook()
    ws = wb.active
    ws.title = "Plantilla Externa"

    azul_header = PatternFill(fill_type="solid", start_color="FF003D89", end_color="FF003D89")
    fuente_header = Font(bold=True, color="FFFFFFFF")
    alineacion_header = Alignment(horizontal="center", vertical="center")

    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.fill = azul_header
        cell.font = fuente_header
        cell.alignment = alineacion_header

    # Dedupe por numero + impo_expo:
    # para cada numero puede existir como maximo un IMPO y un EXPO.
    deduplicados = {}
    registros_ordenados = sorted(
        registros,
        key=lambda x: (
            x.numero if x.numero is not None else 0,
            (x.impo_expo or "").strip().upper(),
            x.created_at or datetime.min,
        ),
    )
    for r in registros_ordenados:
        if r.numero is None:
            continue
        numero = int(r.numero)
        impo_expo = (r.impo_expo or "").strip().upper()
        clave = (numero, impo_expo)
        deduplicados[clave] = r

    filas_plantilla = sorted(
        deduplicados.values(),
        key=lambda x: (
            x.numero if x.numero is not None else 0,
            (x.impo_expo or "").strip().upper(),
        ),
    )

    for r in filas_plantilla:
        ws.append([
            r.numero if r.numero is not None else "",
            r.proveedor_cliente or "",
            r.impo_expo or "",
            "",
            "",
            "",
            "",
            "",
        ])

    ws.freeze_panes = "A2"
    anchos = [28, 34, 14, 14, 26, 22, 22, 20]
    for i, width in enumerate(anchos, start=1):
        ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = width

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    nombre_archivo = f"plantilla_externa_virtuales_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{nombre_archivo}"'},
    )


@app.post("/api/virtuales")
async def crear_virtual(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Crea un nuevo registro de virtuales."""
    try:
        data = await request.json()
    except Exception:
        return JSONResponse(
            status_code=400,
            content={"error": "No se pudo leer el cuerpo de la solicitud"}
        )
    
    numero_raw = data.get("numero")
    if numero_raw in (None, "", "null"):
        return JSONResponse(
            status_code=400,
            content={"error": "El campo 'numero' es requerido"}
        )
    
    try:
        numero = int(numero_raw)
    except (ValueError, TypeError):
        return JSONResponse(
            status_code=400,
            content={"error": "El campo 'numero' debe ser numérico"}
        )
    
    existente = await crud.get_master_unificado_virtuales_by_numero(db, numero)
    if existente:
        return JSONResponse(
            status_code=409,
            content={"error": f"Ya existe un registro con el número {numero}"}
        )

    # Campos obligatorios al crear: numero (ya validado), nombre (proveedor_cliente), impo_expo, tipo
    proveedor_cliente = (data.get("proveedor_cliente") or "").strip()
    if not proveedor_cliente:
        return JSONResponse(
            status_code=400,
            content={"error": "El campo 'Nombre' (Proveedor/Cliente) es obligatorio"}
        )
    impo_expo = (data.get("impo_expo") or "").strip()
    if not impo_expo:
        return JSONResponse(
            status_code=400,
            content={"error": "El campo 'Impo/Expo' es obligatorio"}
        )
    tipo = (data.get("tipo") or "").strip()
    if not tipo:
        return JSONResponse(
            status_code=400,
            content={"error": "El campo 'Tipo' es obligatorio"}
        )
    
    def parse_bool(value):
        if value is None or str(value).strip() == "":
            return None
        if isinstance(value, bool):
            return value
        valor_normalizado = str(value).strip().lower()
        if valor_normalizado in {"si", "sí", "true", "1"}:
            return True
        if valor_normalizado in {"no", "false", "0"}:
            return False
        return None
    
    def parse_int(value):
        if value is None or str(value).strip() == "":
            return None
        try:
            return int(value)
        except (ValueError, TypeError):
            return None
    
    def parse_str(value):
        if value is None:
            return None
        valor = str(value).strip()
        return valor or None
    
    fecha_pago = None
    fecha_pago_raw = data.get("fecha_pago")
    if isinstance(fecha_pago_raw, str) and fecha_pago_raw.strip():
        for formato in ("%d/%m/%Y", "%Y-%m-%d"):
            try:
                fecha_pago = datetime.strptime(fecha_pago_raw.strip(), formato).date()
                break
            except ValueError:
                continue
    
    try:
        nuevo = await crud.create_master_unificado_virtuales(
            db=db,
            solicitud_previo=parse_bool(data.get("solicitud_previo")),
            agente=parse_str(data.get("agente")),
            pedimento=parse_int(data.get("pedimento")),
            aduana=parse_int(data.get("aduana")),
            patente=parse_int(data.get("patente")),
            destino=parse_int(data.get("destino")),
            cliente_space=parse_str(data.get("cliente_space")),
            impo_expo=parse_str(data.get("impo_expo")),
            proveedor_cliente=parse_str(data.get("proveedor_cliente")),
            mes=parse_str(data.get("mes")),
            complemento=parse_str(data.get("complemento")),
            tipo_immex=parse_str(data.get("tipo_immex")),
            factura=parse_str(data.get("factura")),
            fecha_pago=fecha_pago,
            informacion=parse_str(data.get("informacion")),
            estatus=parse_str(data.get("estatus")),
            op_regular=parse_bool(data.get("op_regular")),
            tipo=parse_str(data.get("tipo")),
            numero=numero,
            carretes=parse_bool(data.get("carretes")),
            servicio_cliente=parse_str(data.get("servicio_cliente")),
            plazo=parse_str(data.get("plazo")),
            firma=parse_str(data.get("firma")),
            incoterm=parse_str(data.get("incoterm")),
            tipo_exportacion=parse_str(data.get("tipo_exportacion")),
            escenario=parse_str(data.get("escenario")),
            materialidad=parse_bool(data.get("materialidad")),
            user_id=current_user.id
        )
    except Exception as exc:
        return JSONResponse(
            status_code=500,
            content={"error": f"Error al crear el registro: {str(exc)}"}
        )
    
    return JSONResponse(
        status_code=201,
        content={
            "success": True,
            "message": "Registro creado correctamente",
            "numero": nuevo.numero
        }
    )


@app.delete("/api/virtuales/{numero}")
async def eliminar_virtual(
    numero: int,
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Elimina un registro de virtuales. Solo admin y solicitando detalle."""
    if (current_user.rol or "").lower() != "admin":
        return JSONResponse(
            status_code=403,
            content={"error": "Solo los usuarios administradores pueden eliminar registros"}
        )
    
    try:
        payload = await request.json()
    except Exception:
        payload = {}
    
    detalle = (payload.get("detalle") or payload.get("comentario") or payload.get("motivo") or "").strip()
    if not detalle:
        return JSONResponse(
            status_code=400,
            content={"error": "El detalle de eliminación es obligatorio"}
        )
    
    registro = await crud.get_master_unificado_virtuales_by_numero(db, numero)
    if not registro:
        return JSONResponse(
            status_code=404,
            content={"error": f"No se encontró el registro con número {numero}"}
        )
    
    datos_antes = crud.master_unificado_virtual_to_dict(registro)
    eliminado = await crud.delete_master_unificado_virtuales(db, numero)
    if not eliminado:
        return JSONResponse(
            status_code=404,
            content={"error": f"No se encontró el registro con número {numero}"}
        )
    
    await crud.create_master_unificado_virtual_historial(
        db,
        numero=numero,
        operacion=MasterUnificadoVirtualOperacion.DELETE,
        user_id=current_user.id,
        datos_antes=datos_antes,
        datos_despues=None,
        campos_modificados=None,
        comentario=detalle
    )
    
    return JSONResponse(
        status_code=200,
        content={
            "success": True,
            "message": "Registro eliminado correctamente",
            "numero": numero
        }
    )


@app.post("/api/virtuales/actualizar")
async def actualizar_master_virtuales(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """
    Ejecuta la actualización del master de virtuales basándose en compras
    del mes anterior. Crea registros para proveedores existentes y nuevos.
    Solo permite la actualización si no hay registros creados para el mes actual.
    """
    try:
        mes_actual = MESES_VIRTUALES_ES[datetime.now().month - 1]
        count_mes = await crud.count_master_unificado_virtuales(db, mes=mes_actual)
        if count_mes > 0:
            return JSONResponse(
                status_code=200,
                content={
                    "success": False,
                    "ya_actualizado": True,
                    "message": f"Ese mes ({mes_actual}) ya se actualizaron los registros.",
                    "mes": mes_actual,
                }
            )
        resumen = await crud.actualizar_master_virtuales_desde_compras(
            db=db,
            user_id=current_user.id
        )
        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "message": "Actualización completada",
                "resumen": resumen
            }
        )
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "error": str(e),
                "message": "Error al actualizar el master de virtuales"
            }
        )


@app.get("/api/virtuales/historial")
async def api_virtuales_historial(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    numero: Optional[int] = None,
    operacion: Optional[str] = None,
    limit: int = 50,
    offset: int = 0
):
    """API para consultar el historial de movimientos del master virtuales. Admin ve todos; operador solo los suyos."""
    from app.db.models import MasterUnificadoVirtualOperacion
    operacion_enum = None
    if operacion:
        try:
            operacion_enum = MasterUnificadoVirtualOperacion(operacion.upper())
        except ValueError:
            pass
    # Administrador ve todos los movimientos; operador solo los propios
    user_id_filter = None if current_user.rol == "admin" else current_user.id
    historial = await crud.list_master_unificado_virtuales_historial(
        db, numero=numero, operacion=operacion_enum, user_id=user_id_filter, limit=limit, offset=offset
    )
    total = await crud.count_master_unificado_virtuales_historial(
        db, numero=numero, operacion=operacion_enum, user_id=user_id_filter
    )
    items = []
    for h in historial:
        items.append({
            "id": h.id,
            "numero": h.numero,
            "operacion": h.operacion.value if h.operacion else None,
            "user_id": h.user_id,
            "user_email": h.user.email if h.user else None,
            "datos_antes": h.datos_antes,
            "datos_despues": h.datos_despues,
            "campos_modificados": h.campos_modificados,
            "comentario": h.comentario,
            "created_at": h.created_at.isoformat() if h.created_at else None,
        })
    return JSONResponse(
        status_code=200,
        content={"historial": items, "total": total}
    )


@app.get("/api/virtuales/movimientos")
async def api_virtuales_movimientos(
    request: Request,
    numero: Optional[str] = None,
    tipo: str = "",
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Devuelve las últimas 5 compras (si tipo=Proveedor) o ventas (si tipo=Cliente) para el numero dado."""
    if not numero or not tipo:
        return JSONResponse(
            status_code=400,
            content={"error": "Se requieren numero y tipo (Proveedor o Cliente)."},
        )
    tipo_lower = tipo.strip().lower()
    if "proveedor" in tipo_lower:
        compras = await crud.get_ultimas_compras_proveedor(db, str(numero).strip(), limit=5)
        items = []
        for c in compras:
            items.append({
                "posting_date": c.posting_date.isoformat() if c.posting_date else None,
                "purchasing_document": c.purchasing_document,
                "material_document": c.material_document,
                "numero_material": c.numero_material,
                "descripcion_material": (c.descripcion_material or "")[:80],
                "quantity": float(c.quantity) if c.quantity is not None else None,
                "amount": float(c.amount) if c.amount is not None else None,
                "currency": c.currency,
                "nombre_proveedor": c.nombre_proveedor,
            })
        return JSONResponse(content={"tipo": "compras", "movimientos": items})
    if "cliente" in tipo_lower:
        try:
            cod_cliente = int(str(numero).strip())
        except (TypeError, ValueError):
            cod_cliente = None
        ventas = await crud.get_ultimas_ventas_cliente(db, cod_cliente, limit=5)
        items = []
        for v in ventas:
            items.append({
                "periodo": v.periodo.isoformat() if v.periodo else None,
                "cliente": v.cliente,
                "codigo_cliente": v.codigo_cliente,
                "producto": v.producto,
                "descripcion_producto": (v.descripcion_producto or "")[:80],
                "sales_total_mts": float(v.sales_total_mts) if v.sales_total_mts is not None else None,
                "turnover_wo_metal": float(v.turnover_wo_metal) if v.turnover_wo_metal is not None else None,
                "planta": v.planta,
            })
        return JSONResponse(content={"tipo": "ventas", "movimientos": items})
    return JSONResponse(
        status_code=400,
        content={"error": "tipo debe ser Proveedor o Cliente."},
    )


@app.put("/api/virtuales/{numero}")
async def actualizar_virtual(
    numero: int,
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Actualiza un registro de virtuales."""
    try:
        data = await request.json()
        
        # Convertir valores booleanos
        solicitud_previo = None
        if data.get("solicitud_previo"):
            solicitud_previo = data.get("solicitud_previo") == "Sí"
        
        op_regular = None
        if data.get("op_regular"):
            op_regular = data.get("op_regular") == "Sí"
        
        carretes = None
        if data.get("carretes"):
            carretes = data.get("carretes") == "Sí"
        
        materialidad = None
        if "materialidad" in data:
            val = data.get("materialidad")
            if val is True or (isinstance(val, str) and str(val).strip().lower() in ("si", "sí", "true", "1")):
                materialidad = True
            elif val is False or (isinstance(val, str) and str(val).strip().lower() in ("no", "false", "0")):
                materialidad = False
        
        # Convertir valores numéricos
        pedimento = None
        if data.get("pedimento"):
            try:
                pedimento = int(data.get("pedimento"))
            except (ValueError, TypeError):
                pass
        
        aduana = None
        if data.get("aduana"):
            try:
                aduana = int(data.get("aduana"))
            except (ValueError, TypeError):
                pass
        
        patente = None
        if data.get("patente"):
            try:
                patente = int(data.get("patente"))
            except (ValueError, TypeError):
                pass
        
        destino = None
        if data.get("destino"):
            try:
                destino = int(data.get("destino"))
            except (ValueError, TypeError):
                pass
        
        # Si el frontend envía id del registro, usarlo para identificar (evita ambigüedad cuando hay varios con el mismo numero)
        master_id = None
        if data.get("id") is not None and data.get("id") != "":
            try:
                master_id = int(data.get("id"))
            except (ValueError, TypeError):
                pass
        
        virtual_actualizado = await crud.update_master_unificado_virtuales(
            db=db,
            numero=numero,
            master_id=master_id,
            solicitud_previo=solicitud_previo,
            agente=data.get("agente") or None,
            pedimento=pedimento,
            aduana=aduana,
            patente=patente,
            destino=destino,
            cliente_space=data.get("cliente_space") or None,
            impo_expo=data.get("impo_expo") or None,
            proveedor_cliente=data.get("proveedor_cliente") or None,
            mes=data.get("mes") or None,
            complemento=data.get("complemento") or None,
            tipo_immex=data.get("tipo_immex") or None,
            factura=data.get("factura") or None,
            informacion=data.get("informacion") or None,
            estatus=data.get("estatus") or None,
            op_regular=op_regular,
            tipo=data.get("tipo") or None,
            carretes=carretes,
            servicio_cliente=data.get("servicio_cliente") or None,
            plazo=data.get("plazo") or None,
            firma=data.get("firma") or None,
            incoterm=data.get("incoterm") or None,
            tipo_exportacion=data.get("tipo_exportacion") or None,
            escenario=data.get("escenario") or None,
            materialidad=materialidad,
            user_id=current_user.id
        )
        
        if not virtual_actualizado:
            return JSONResponse(
                status_code=404,
                content={"error": f"No se encontró el registro con número {numero}"}
            )
        
        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "message": "Registro actualizado correctamente",
                "numero": numero
            }
        )
        
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Error al actualizar el registro: {str(e)}"}
        )


@app.post("/api/virtuales/{numero}/duplicar-impo")
async def duplicar_virtual_expo_a_impo(
    numero: int,
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Duplica un registro EXPO como IMPO con la misma información."""
    try:
        # Obtener el registro EXPO más reciente para este numero
        registro_expo = await crud.get_master_unificado_virtuales_by_numero_impoexpo(
            db=db,
            numero=numero,
            impo_expo="EXPO"
        )
        
        if not registro_expo:
            return JSONResponse(
                status_code=404,
                content={"error": f"No se encontró un registro EXPO con número {numero}"}
            )
        
        # Verificar si ya existe un registro IMPO para este numero y mes
        if await crud.existe_virtual_numero_mes(
            db=db,
            numero=numero,
            mes_captura=registro_expo.mes or "",
            tipo=registro_expo.tipo or "",
            impo_expo="IMPO"
        ):
            return JSONResponse(
                status_code=409,
                content={"error": f"Ya existe un registro IMPO para el número {numero} en el mes {registro_expo.mes}"}
            )
        
        # Sincronizar secuencia de id para evitar UniqueViolationError (id ya existente)
        from sqlalchemy import text
        try:
            await db.execute(
                text(
                    "SELECT setval(pg_get_serial_sequence('master_unificado_virtuales', 'id'), "
                    "COALESCE((SELECT MAX(id) FROM master_unificado_virtuales), 1))"
                )
            )
        except Exception:
            pass  # Si la tabla no usa secuencia o falla, se intenta el create igualmente
        
        # Crear nuevo registro copiando todos los campos pero con impo_expo="IMPO"
        nuevo_impo = await crud.create_master_unificado_virtuales(
            db=db,
            solicitud_previo=registro_expo.solicitud_previo,
            agente=registro_expo.agente,
            pedimento=registro_expo.pedimento,
            aduana=registro_expo.aduana,
            patente=registro_expo.patente,
            destino=registro_expo.destino,
            cliente_space=registro_expo.cliente_space,
            impo_expo="IMPO",
            proveedor_cliente=registro_expo.proveedor_cliente,
            mes=registro_expo.mes,
            complemento=registro_expo.complemento,
            tipo_immex=registro_expo.tipo_immex,
            factura=registro_expo.factura,
            fecha_pago=registro_expo.fecha_pago,
            informacion=registro_expo.informacion,
            estatus=registro_expo.estatus,
            op_regular=registro_expo.op_regular,
            tipo=registro_expo.tipo,
            numero=registro_expo.numero,
            carretes=registro_expo.carretes,
            servicio_cliente=registro_expo.servicio_cliente,
            plazo=registro_expo.plazo,
            firma=registro_expo.firma,
            incoterm=registro_expo.incoterm,
            tipo_exportacion=registro_expo.tipo_exportacion,
            escenario=registro_expo.escenario,
            materialidad=registro_expo.materialidad,
            user_id=current_user.id
        )
        
        # Actualizar el comentario del historial para indicar que es una duplicación
        from app.db.models import MasterUnificadoVirtualHistorial
        from sqlalchemy import select, desc
        historial_mas_reciente = await db.execute(
            select(MasterUnificadoVirtualHistorial)
            .where(MasterUnificadoVirtualHistorial.numero == nuevo_impo.numero)
            .order_by(desc(MasterUnificadoVirtualHistorial.created_at))
            .limit(1)
        )
        historial = historial_mas_reciente.scalar_one_or_none()
        if historial:
            historial.comentario = f"Duplicado desde registro EXPO (número {numero})"
            await db.commit()
        
        return JSONResponse(
            status_code=201,
            content={
                "success": True,
                "message": "Registro duplicado como IMPO correctamente",
                "numero": nuevo_impo.numero
            }
        )
        
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Error al duplicar el registro: {str(e)}"}
        )


@app.post("/api/virtuales/actualizar-desde-excel")
async def actualizar_virtuales_desde_excel(
    archivo: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Lee un Excel con columnas Codigo cliente, Impo/Expo, Mes y opcionales patente, aduana, complemento, firma, pedimento.
    Para cada fila busca el registro en master_unificado_virtuales por (numero, impo_expo, mes) y actualiza
    patente, aduana, complemento, firma y pedimento. Los cambios se registran en el historial.
    """
    import tempfile
    import shutil
    from pathlib import Path
    import pandas as pd

    if not archivo.filename:
        return JSONResponse(status_code=400, content={"error": "No se envió ningún archivo"})
    ext = Path(archivo.filename).suffix.lower()
    if ext not in (".xlsx", ".xls"):
        return JSONResponse(
            status_code=400,
            content={"error": "El archivo debe ser Excel (.xlsx o .xls)"},
        )

    def _norm(s):
        if s is None or (isinstance(s, float) and pd.isna(s)):
            return ""
        return str(s).strip().lower().replace(" ", "_").replace("/", "_")

    def _find_col(df, opciones):
        cols_lower = {_norm(c): c for c in df.columns}
        for op in opciones:
            if _norm(op) in cols_lower:
                return cols_lower[_norm(op)]
        return None

    def _int_or_none(val):
        if val is None or (isinstance(val, float) and pd.isna(val)):
            return None
        try:
            return int(float(val))
        except (TypeError, ValueError):
            return None

    def _str_or_none(val):
        if val is None or (isinstance(val, float) and pd.isna(val)):
            return None
        s = str(val).strip()
        return s if s else None

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            shutil.copyfileobj(archivo.file, tmp)
            tmp_path = tmp.name
        try:
            engine = "openpyxl" if ext == ".xlsx" else "xlrd"
            df = pd.read_excel(tmp_path, engine=engine)
        finally:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass

        col_codigo = _find_col(df, ["Codigo cliente", "codigo cliente", "numero", "codigo_cliente", "número", "numero de cliente"])
        col_impo = _find_col(df, ["Impo/Expo", "impo/expo", "impo_expo", "impo expo", "tipo"])
        col_mes = _find_col(df, ["Mes", "mes"])
        col_patente = _find_col(df, ["patente", "Patente"])
        col_aduana = _find_col(df, ["aduana", "Aduana"])
        col_complemento = _find_col(df, ["complemento", "Complemento"])
        col_firma = _find_col(df, ["firma", "Firma"])
        col_pedimento = _find_col(df, ["pedimento", "Pedimento"])

        if not col_codigo or not col_impo or not col_mes:
            return JSONResponse(
                status_code=400,
                content={
                    "error": "El Excel debe tener las columnas: Codigo cliente, Impo/Expo y Mes. No se encontraron todas."
                },
            )

        actualizados = 0
        no_encontrados = []

        for idx, row in df.iterrows():
            codigo = _int_or_none(row.get(col_codigo))
            if codigo is None:
                continue
            impo_expo_raw = _str_or_none(row.get(col_impo))
            if not impo_expo_raw:
                continue
            impo_expo = impo_expo_raw.strip().upper()
            if impo_expo not in ("IMPO", "EXPO"):
                continue
            mes_raw = _str_or_none(row.get(col_mes))
            if not mes_raw:
                continue
            mes = mes_raw.strip()

            patente = _int_or_none(row.get(col_patente)) if col_patente else None
            aduana = _int_or_none(row.get(col_aduana)) if col_aduana else None
            complemento = _str_or_none(row.get(col_complemento)) if col_complemento else None
            firma = _str_or_none(row.get(col_firma)) if col_firma else None
            pedimento = _int_or_none(row.get(col_pedimento)) if col_pedimento else None

            if patente is None and aduana is None and complemento is None and firma is None and pedimento is None:
                continue

            master = await crud.update_master_unificado_virtuales_campos_desde_excel(
                db,
                numero=codigo,
                impo_expo=impo_expo,
                mes=mes,
                user_id=current_user.id,
                patente=patente,
                aduana=aduana,
                complemento=complemento,
                firma=firma,
                pedimento=pedimento,
            )
            if master:
                actualizados += 1
            else:
                no_encontrados.append({"fila": int(idx) + 2, "codigo": codigo, "impo_expo": impo_expo, "mes": mes})

        await db.commit()

        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "message": f"Se actualizaron {actualizados} registro(s) desde el Excel.",
                "actualizados": actualizados,
                "no_encontrados": no_encontrados,
            },
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={"error": f"Error al procesar el Excel: {str(e)}"},
        )


@app.post("/api/virtuales/crear-carpeta")
async def crear_carpeta_virtual(
    request: Request,
    current_user: User = Depends(get_current_user),
):
    """Crea una carpeta en el Desktop con nombre [codigo] [nombre] [impo/expo] [mes]."""
    import re
    try:
        data = await request.json()
        numero = (data.get("numero") or "")
        nombre = (data.get("proveedor_cliente") or "").strip()
        impo_expo = (data.get("impo_expo") or "").strip()
        mes = (data.get("mes") or "").strip()
        # Nombre de carpeta: cu_[codigo] [nombre] [impo/expo] [mes]_erob1001
        parte_media = f"{numero} {nombre} {impo_expo} {mes}".strip()
        # Sanitizar para filesystem: quitar caracteres no válidos \ / : * ? " < > |
        parte_media = re.sub(r'[\\/:*?"<>|]', " ", parte_media)
        parte_media = re.sub(r"\s+", " ", parte_media).strip()
        nombre_carpeta = f"cu_{parte_media}_erob1001" if parte_media else ""
        if not nombre_carpeta:
            return JSONResponse(
                status_code=400,
                content={"error": "No se pudo generar un nombre de carpeta válido"}
            )
        desktop = os.path.expanduser("~/Desktop")
        ruta_carpeta = os.path.join(desktop, nombre_carpeta)
        if os.path.isdir(ruta_carpeta):
            return JSONResponse(
                status_code=200,
                content={
                    "success": True,
                    "message": "La carpeta ya existe",
                    "ruta": ruta_carpeta,
                    "nombre_carpeta": nombre_carpeta,
                    "ya_existia": True,
                }
            )
        os.makedirs(ruta_carpeta, exist_ok=False)
        return JSONResponse(
            status_code=201,
            content={
                "success": True,
                "message": "Carpeta creada correctamente",
                "ruta": ruta_carpeta,
                "nombre_carpeta": nombre_carpeta,
            }
        )
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Error al crear la carpeta: {str(e)}"}
        )


@app.put("/api/paises-origen/{pais_id}")
async def actualizar_pais_origen(
    pais_id: int,
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Actualiza un país de origen de material."""
    try:
        data = await request.json()
        pais_origen = data.get("pais_origen")
        porcentaje_compra = data.get("porcentaje_compra")
        comentario = data.get("comentario")
        if porcentaje_compra is not None and porcentaje_compra != "":
            try:
                porcentaje_compra = float(porcentaje_compra)
            except (TypeError, ValueError):
                porcentaje_compra = None
        else:
            porcentaje_compra = None
        if comentario is not None and isinstance(comentario, str):
            comentario = comentario.strip() or None
        else:
            comentario = None
        
        if not pais_origen:
            return JSONResponse(
                status_code=400,
                content={"error": "El campo 'pais_origen' es requerido"}
            )
        
        pais_actualizado = await crud.update_pais_origen_material(
            db=db,
            pais_id=pais_id,
            pais_origen=pais_origen,
            porcentaje_compra=porcentaje_compra,
            comentario=comentario,
            user_id=current_user.id
        )
        
        if not pais_actualizado:
            return JSONResponse(
                status_code=404,
                content={"error": "País de origen no encontrado"}
            )
        
        return JSONResponse({
            "success": True,
            "message": "País de origen actualizado exitosamente",
            "pais": {
                "id": pais_actualizado.id,
                "codigo_proveedor": pais_actualizado.codigo_proveedor,
                "numero_material": pais_actualizado.numero_material,
                "pais_origen": pais_actualizado.pais_origen,
                "porcentaje_compra": float(pais_actualizado.porcentaje_compra) if pais_actualizado.porcentaje_compra is not None else None,
                "comentario": pais_actualizado.comentario,
            }
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "error": str(e),
                "message": f"Error al actualizar país de origen: {str(e)}"
            }
        )


@app.get("/api/paises-origen/historial")
async def api_paises_origen_historial(
    request: Request,
    current_user: User = Depends(require_roles(["admin"])),
    db: AsyncSession = Depends(get_db)
):
    """Devuelve los últimos 15 movimientos del historial de países de origen. Solo administradores."""
    historial = await crud.list_pais_origen_material_historial(
        db,
        limit=15,
        offset=0
    )
    items = []
    for h in historial:
        items.append({
            "id": h.id,
            "pais_origen_id": h.pais_origen_id,
            "codigo_proveedor": h.codigo_proveedor,
            "numero_material": h.numero_material,
            "operacion": h.operacion.value,
            "user_email": h.user.email if h.user else None,
            "user_nombre": h.user.nombre if h.user else None,
            "datos_antes": h.datos_antes,
            "datos_despues": h.datos_despues,
            "campos_modificados": h.campos_modificados,
            "created_at": h.created_at.isoformat() if h.created_at else None,
        })
    return JSONResponse({"movimientos": items})


@app.get("/api/proveedores")
async def api_proveedores(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    search: Optional[str] = None,
    estatus: Optional[bool] = None,
    pais: Optional[str] = None,
    limit: int = 100,
    offset: int = 0
):
    """API para obtener proveedores con filtros y paginación."""
    proveedores = await crud.list_proveedores(
        db,
        estatus=estatus,
        pais=pais,
        search=search,
        limit=limit,
        offset=offset
    )
    
    total = await crud.count_proveedores(
        db,
        estatus=estatus,
        pais=pais,
        search=search
    )
    
    return JSONResponse({
        "proveedores": [
            {
                "codigo_proveedor": p.codigo_proveedor,
                "nombre": p.nombre,
                "pais": p.pais,
                "domicilio": p.domicilio,
                "poblacion": p.poblacion,
                "cp": p.cp,
                "estatus": p.estatus,
                "estatus_compras": p.estatus_compras,
                "created_at": p.created_at.isoformat() if p.created_at else None,
                "updated_at": p.updated_at.isoformat() if p.updated_at else None
            }
            for p in proveedores
        ],
        "total": total,
        "limit": limit,
        "offset": offset
    })


@app.post("/api/materiales/normalizar-cientificos")
async def normalizar_materiales_cientificos(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Reemplaza todos los numero_material en notación científica por el número entero
    (ej. 3.4241E+11 -> 342410000000) en materiales, compras, precios y países de origen.
    No elimina datos; solo normaliza el formato."""
    try:
        resultado = await crud.normalizar_numero_material_cientifico_en_bd(db=db)
        total = (
            resultado["materiales_actualizados"]
            + resultado["materiales_eliminados_duplicados"]
            + resultado["compras_actualizadas"]
            + resultado["precios_materiales_actualizados"]
            + resultado["pais_origen_material_actualizados"]
            + resultado["historiales_actualizados"]
        )
        mensaje = (
            f"✓ Normalización completada: {resultado['materiales_actualizados']} material(es) actualizado(s), "
            f"{resultado['materiales_eliminados_duplicados']} duplicado(s) eliminado(s), "
            f"{resultado['compras_actualizadas']} compra(s), "
            f"{resultado['precios_materiales_actualizados']} precio(s), "
            f"{resultado['pais_origen_material_actualizados']} país(es) de origen, "
            f"{resultado['historiales_actualizados']} historial(es)."
        )
        if resultado["errores"]:
            mensaje += f" Errores: {len(resultado['errores'])}."
        return JSONResponse({
            "success": len(resultado["errores"]) == 0,
            "mensaje": mensaje,
            **resultado,
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "error": str(e),
                "mensaje": f"Error al normalizar notación científica: {str(e)}",
            }
        )


@app.post("/api/materiales/actualizar")
async def actualizar_materiales_desde_compras(
    request: Request,
    current_user: User = Depends(require_roles(["admin"])),
    db: AsyncSession = Depends(get_db)
):
    """Sincroniza materiales desde la tabla compras.
    
    Busca todos los numero_material únicos en compras que no estén registrados
    en materiales y los crea automáticamente usando la descripcion_material de compras.
    """
    try:
        resultado = await crud.sincronizar_materiales_desde_compras(
            db=db,
            user_id=current_user.id
        )
        
        # Determinar el mensaje según el resultado
        if resultado["nuevos_creados"] > 0:
            mensaje = f"✓ Sincronización completada. Se crearon {resultado['nuevos_creados']} nuevo(s) material(es) de {resultado['total_encontrados']} encontrados en compras."
        elif resultado["total_encontrados"] > 0:
            mensaje = f"✓ Sincronización completada. Todos los materiales ({resultado['total_encontrados']}) ya están registrados."
        else:
            mensaje = "✓ Sincronización completada. No se encontraron materiales en la tabla de compras."
        
        if resultado["errores"]:
            mensaje += f" Se encontraron {len(resultado['errores'])} error(es)."
        
        # Considerar exitoso si no hay errores críticos
        success = len(resultado["errores"]) == 0 or resultado["nuevos_creados"] > 0
        
        return JSONResponse({
            "success": success,
            "total_encontrados": resultado["total_encontrados"],
            "nuevos_creados": resultado["nuevos_creados"],
            "errores": resultado["errores"],
            "mensaje": mensaje
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "error": str(e),
                "mensaje": f"Error al sincronizar materiales desde compras: {str(e)}"
            }
        )


@app.post("/api/proveedores/actualizar")
async def actualizar_proveedores_desde_compras(
    request: Request,
    current_user: User = Depends(require_roles(["admin"])),
    db: AsyncSession = Depends(get_db)
):
    """Sincroniza proveedores desde la tabla compras.
    
    Busca todos los codigo_proveedor únicos en compras que no estén registrados
    en proveedores y los crea automáticamente usando el nombre_proveedor de compras.
    """
    try:
        resultado = await crud.sincronizar_proveedores_desde_compras(
            db=db,
            user_id=current_user.id
        )
        
        # Determinar el mensaje según el resultado
        if resultado["nuevos_creados"] > 0:
            mensaje = f"✓ Sincronización completada. Se crearon {resultado['nuevos_creados']} nuevo(s) proveedor(es) de {resultado['total_encontrados']} encontrados en compras."
        elif resultado["total_encontrados"] > 0:
            mensaje = f"✓ Sincronización completada. Todos los proveedores ({resultado['total_encontrados']}) ya están registrados."
        else:
            mensaje = "✓ Sincronización completada. No se encontraron proveedores en la tabla de compras."
        
        if resultado["errores"]:
            mensaje += f" Se encontraron {len(resultado['errores'])} error(es)."
        
        # Considerar exitoso si no hay errores críticos
        success = len(resultado["errores"]) == 0 or resultado["nuevos_creados"] > 0
        
        return JSONResponse({
            "success": success,
            "total_encontrados": resultado["total_encontrados"],
            "nuevos_creados": resultado["nuevos_creados"],
            "errores": resultado["errores"],
            "mensaje": mensaje
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "error": str(e),
                "mensaje": f"Error al sincronizar proveedores desde compras: {str(e)}"
            }
        )


@app.post("/api/proveedores/actualizar-estatus")
async def actualizar_estatus_proveedores(
    request: Request,
    current_user: User = Depends(require_roles(["admin"])),
    db: AsyncSession = Depends(get_db)
):
    """Actualiza el estatus_compras de proveedores basándose en compras y fecha de creación.
    
    Reglas:
    - Cliente existente con compras en los últimos 6 meses → estatus_compras = "activo"
    - Cliente existente sin compras en los últimos 6 meses → estatus_compras = "baja"
    - Cliente nuevo creado este mes → estatus_compras = "alta"
    """
    try:
        resultado = await crud.actualizar_estatus_proveedores_por_compras(
            db=db,
            user_id=current_user.id
        )
        
        # Construir mensaje detallado
        partes_mensaje = []
        if resultado["proveedores_marcados_baja"] > 0:
            partes_mensaje.append(f"{resultado['proveedores_marcados_baja']} como baja")
        if resultado["proveedores_marcados_activo"] > 0:
            partes_mensaje.append(f"{resultado['proveedores_marcados_activo']} como activo")
        if resultado["proveedores_marcados_alta"] > 0:
            partes_mensaje.append(f"{resultado['proveedores_marcados_alta']} como alta")
        
        if partes_mensaje:
            mensaje = f"✓ Actualización completada. Se actualizaron {resultado['proveedores_actualizados']} proveedor(es): {', '.join(partes_mensaje)} de {resultado['total_proveedores']} totales."
        else:
            mensaje = f"✓ Actualización completada. No se requirieron cambios en los {resultado['total_proveedores']} proveedores."
        
        if resultado["errores"]:
            mensaje += f" Se encontraron {len(resultado['errores'])} error(es)."
        
        # Considerar exitoso si no hay errores críticos
        success = len(resultado["errores"]) == 0
        
        return JSONResponse({
            "success": success,
            "total_proveedores": resultado["total_proveedores"],
            "proveedores_actualizados": resultado["proveedores_actualizados"],
            "proveedores_marcados_baja": resultado["proveedores_marcados_baja"],
            "proveedores_marcados_activo": resultado["proveedores_marcados_activo"],
            "proveedores_marcados_alta": resultado["proveedores_marcados_alta"],
            "errores": resultado["errores"],
            "mensaje": mensaje
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "error": str(e),
                "mensaje": f"Error al actualizar estatus de proveedores: {str(e)}"
            }
        )


@app.post("/api/paises-origen/actualizar")
async def actualizar_paises_origen_desde_compras(
    request: Request,
    current_user: User = Depends(require_roles(["admin"])),
    db: AsyncSession = Depends(get_db)
):
    """Sincroniza países de origen desde la tabla compras.
    
    Busca combinaciones únicas de codigo_proveedor y numero_material en compras
    que no estén registradas en pais_origen_material y las crea con "Pendiente".
    """
    try:
        resultado = await crud.sincronizar_paises_origen_desde_compras(
            db=db,
            user_id=current_user.id
        )
        
        # Determinar el mensaje según el resultado
        if resultado["nuevos_creados"] > 0:
            mensaje = (
                f"✓ Sincronización completada. Se crearon {resultado['nuevos_creados']} "
                f"registro(s) de {resultado['total_encontrados']} encontrados en compras."
            )
        elif resultado["total_encontrados"] > 0:
            mensaje = (
                f"✓ Sincronización completada. Todos los registros "
                f"({resultado['total_encontrados']}) ya están registrados."
            )
        else:
            mensaje = "✓ Sincronización completada. No se encontraron datos válidos en compras."
        
        if resultado["errores"]:
            mensaje += f" Se encontraron {len(resultado['errores'])} error(es)."
        
        # Considerar exitoso si no hay errores críticos
        success = len(resultado["errores"]) == 0 or resultado["nuevos_creados"] > 0
        
        return JSONResponse({
            "success": success,
            "total_encontrados": resultado["total_encontrados"],
            "nuevos_creados": resultado["nuevos_creados"],
            "errores": resultado["errores"],
            "mensaje": mensaje
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "error": str(e),
                "mensaje": f"Error al sincronizar países de origen desde compras: {str(e)}"
            }
        )


@app.get("/api/compras")
async def api_compras(
    request: Request,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    search: Optional[str] = None,
    fecha_inicio: Optional[str] = None,
    fecha_fin: Optional[str] = None,
    codigo_proveedor: Optional[str] = None,
    numero_material: Optional[str] = None,
    purchasing_document: Optional[int] = None,
    material_document: Optional[int] = None,
    limit: int = 100,
    offset: int = 0
):
    """API para obtener compras con filtros y paginación."""
    # Convertir fechas de string a datetime si están presentes
    fecha_inicio_dt = None
    fecha_fin_dt = None
    
    if fecha_inicio:
        try:
            fecha_inicio_dt = datetime.strptime(fecha_inicio, "%Y-%m-%d")
            # Agregar zona horaria si no la tiene
            if fecha_inicio_dt.tzinfo is None:
                fecha_inicio_dt = fecha_inicio_dt.replace(tzinfo=ZoneInfo('America/Mexico_City'))
        except ValueError:
            return JSONResponse(
                status_code=400,
                content={"error": "Formato de fecha_inicio inválido. Use formato YYYY-MM-DD"}
            )
    
    if fecha_fin:
        try:
            fecha_fin_dt = datetime.strptime(fecha_fin, "%Y-%m-%d")
            # Agregar zona horaria si no la tiene
            if fecha_fin_dt.tzinfo is None:
                fecha_fin_dt = fecha_fin_dt.replace(tzinfo=ZoneInfo('America/Mexico_City'))
        except ValueError:
            return JSONResponse(
                status_code=400,
                content={"error": "Formato de fecha_fin inválido. Use formato YYYY-MM-DD"}
            )
    
    compras = await crud.list_compras(
        db,
        search=search,
        fecha_inicio=fecha_inicio_dt,
        fecha_fin=fecha_fin_dt,
        codigo_proveedor=codigo_proveedor,
        numero_material=numero_material,
        purchasing_document=purchasing_document,
        material_document=material_document,
        limit=limit,
        offset=offset
    )
    
    total = await crud.count_compras(
        db,
        search=search,
        fecha_inicio=fecha_inicio_dt,
        fecha_fin=fecha_fin_dt,
        codigo_proveedor=codigo_proveedor,
        numero_material=numero_material,
        purchasing_document=purchasing_document,
        material_document=material_document
    )
    
    return JSONResponse({
        "compras": [
            {
                "id": c.id,
                "purchasing_document": c.purchasing_document,
                "item": c.item,
                "material_doc_year": c.material_doc_year,
                "material_document": c.material_document,
                "material_doc_item": c.material_doc_item,
                "movement_type": c.movement_type,
                "posting_date": c.posting_date.isoformat() if c.posting_date else None,
                "quantity": c.quantity,
                "order_unit": c.order_unit,
                "quantity_in_opun": c.quantity_in_opun,
                "order_price_unit": c.order_price_unit,
                "amount_in_lc": float(c.amount_in_lc) if c.amount_in_lc else None,
                "local_currency": c.local_currency,
                "amount": float(c.amount) if c.amount else None,
                "currency": c.currency,
                "gr_ir_clearing_value_lc": float(c.gr_ir_clearing_value_lc) if c.gr_ir_clearing_value_lc else None,
                "invoice_value": float(c.invoice_value) if c.invoice_value else None,
                "numero_material": c.numero_material,
                "plant": c.plant,
                "descripcion_material": c.descripcion_material,
                "nombre_proveedor": c.nombre_proveedor,
                "codigo_proveedor": c.codigo_proveedor,
                "price": float(c.price) if c.price else None,
                "created_at": c.created_at.isoformat() if c.created_at else None,
                "updated_at": c.updated_at.isoformat() if c.updated_at else None
            }
            for c in compras
        ],
        "total": total,
        "limit": limit,
        "offset": offset
    })


@app.post("/api/compras/iniciar-descarga")
async def iniciar_descarga(
    request: Request,
    fecha_inicio: str = Form(...),
    fecha_fin: str = Form(...),
    carpeta_salida: str = Form(default="C:\\Users\\anad5004\\Documents\\Leoni_RPA"),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Inicia el proceso de descarga de compras ejecutando compras_local.vbs con SAP GUI y espera a que termine."""
    import subprocess
    import asyncio
    from pathlib import Path
    execution_id = None
    try:
        # Validar y convertir fechas
        # El input type="date" devuelve formato YYYY-MM-DD
        try:
            fecha_inicio_dt = datetime.strptime(fecha_inicio, "%Y-%m-%d")
            fecha_fin_dt = datetime.strptime(fecha_fin, "%Y-%m-%d")
            # Agregar hora al final del día para fecha_fin para incluir todo el día
            fecha_fin_dt = fecha_fin_dt.replace(hour=23, minute=59, second=59)
        except ValueError as e:
            return JSONResponse(
                status_code=400,
                content={"error": f"Formato de fecha inválido: {str(e)}. Use formato YYYY-MM-DD"}
            )
        
        # Validar que fecha fin sea posterior a fecha inicio
        if fecha_fin_dt < fecha_inicio_dt:
            return JSONResponse(
                status_code=400,
                content={"error": "La fecha fin debe ser posterior a la fecha inicio"}
            )
        
        # Carpeta de salida fija para el script VBS
        carpeta_salida = "C:\\Users\\anad5004\\Documents\\Leoni_RPA"
        
        # Obtener información de la máquina
        try:
            hostname = socket.gethostname()
        except:
            hostname = platform.node() or "unknown"
        
        # Crear el registro de ejecución (retorna id para evitar MissingGreenlet al acceder al objeto tras commit)
        execution_id = await crud.create_execution(
            db=db,
            user_id=current_user.id,
            fecha_inicio_periodo=fecha_inicio_dt,
            fecha_fin_periodo=fecha_fin_dt,
            sistema_sap="SAP ECC",  # Valor por defecto, puede actualizarse después
            transaccion="ME80FN",  # Transaccion del script compras_local.vbs
            maquina=hostname
        )

        # Construir el nombre del archivo esperado (el VBS exporta .xlsx)
        fecha_inicio_str = fecha_inicio_dt.strftime("%Y%m%d")
        fecha_fin_str = fecha_fin_dt.strftime("%Y%m%d")
        nombre_archivo = f"compras_local_{fecha_inicio_str}_{fecha_fin_str}.xlsx, historial_compras_{fecha_inicio_str}_{fecha_fin_str}.xlsx"
        ruta_completa = f"{carpeta_salida.rstrip('/').rstrip(chr(92))}{os.sep}(archivos .xlsx)"
        
        # Ejecutar compras_local.vbs con las fechas seleccionadas
        script_dir = Path(__file__).resolve().parent
        vbs_path = script_dir / "compras_local.vbs"
        if not vbs_path.exists():
            await crud.update_execution_status(
                db=db, execution_id=execution_id, estado=ExecutionStatus.FAILED,
                detalles=f"No se encontró el script compras_local.vbs en {script_dir}"
            )
            return JSONResponse(
                status_code=500,
                content={"error": f"No se encontró el script compras_local.vbs en {script_dir}"}
            )

        # Actualizar con la información del archivo esperado (estado PENDING mientras corre)
        await crud.update_execution_status(
            db=db,
            execution_id=execution_id,
            estado=ExecutionStatus.PENDING,
            archivo_ruta=ruta_completa,
            archivo_nombre=nombre_archivo
        )

        # Ejecutar el script en un thread para no bloquear el event loop (run_in_executor evita conflicto greenlet/SQLAlchemy)
        loop = asyncio.get_running_loop()
        def _run_compras_vbs():
            return subprocess.run(
                ["cscript", "//nologo", str(vbs_path), fecha_inicio_str, fecha_fin_str],
                cwd=str(script_dir),
                capture_output=True,
                timeout=300,
            )
        try:
            result = await loop.run_in_executor(None, _run_compras_vbs)
        except subprocess.TimeoutExpired:
            await crud.update_execution_status(
                db=db, execution_id=execution_id, estado=ExecutionStatus.FAILED,
                detalles="El proceso de descarga superó el tiempo máximo (5 minutos)."
            )
            return JSONResponse(
                status_code=500,
                content={"error": "El proceso de descarga superó el tiempo máximo (5 minutos)."}
            )
        except Exception as e:
            await crud.update_execution_status(
                db=db, execution_id=execution_id, estado=ExecutionStatus.FAILED,
                detalles=f"Error al ejecutar el script: {str(e)}"
            )
            return JSONResponse(
                status_code=500,
                content={"error": f"Error al ejecutar el script: {str(e)}"}
            )

        if result.returncode != 0:
            err_msg = (result.stderr or result.stdout or b"").decode("utf-8", errors="replace").strip() or f"Código de salida: {result.returncode}"
            await crud.update_execution_status(
                db=db, execution_id=execution_id, estado=ExecutionStatus.FAILED,
                detalles=f"El script SAP finalizó con error. {err_msg}"
            )
            return JSONResponse(
                status_code=500,
                content={"error": f"El script SAP finalizó con error. {err_msg}"}
            )

        # Procesamiento automático: leer los Excel descargados e insertar/actualizar compras
        path_base = Path(carpeta_salida)
        path_compras_xlsx = path_base / f"compras_local_{fecha_inicio_str}_{fecha_fin_str}.xlsx"
        path_historial_xlsx = path_base / f"historial_compras_{fecha_inicio_str}_{fecha_fin_str}.xlsx"
        detalles_str = None
        resultado_proc = None
        if path_compras_xlsx.exists() and path_historial_xlsx.exists():
            resultado_proc = await _procesar_compras_historial_desde_rutas(db, path_compras_xlsx, path_historial_xlsx)
            if resultado_proc.get("success"):
                ins = resultado_proc.get("insertados", 0)
                dup = resultado_proc.get("duplicados", 0)
                errs = resultado_proc.get("errores", [])
                detalles_str = f"Procesamiento automático: {ins} insertados, {dup} duplicados (ya existían)."
                if errs:
                    detalles_str += f" {len(errs)} error(es)."
                # El cierre de Excel se hace al final de compras_local.vbs (igual que ventas.vbs). Aquí solo eliminamos los archivos tras procesarlos.
                for p in (path_compras_xlsx, path_historial_xlsx):
                    for _ in range(5):
                        try:
                            if p.exists():
                                p.unlink()
                            break
                        except Exception:
                            await asyncio.sleep(2)
            else:
                detalles_str = f"Procesamiento automático falló: {resultado_proc.get('error', 'Error desconocido')}"

        await crud.update_execution_status(
            db=db,
            execution_id=execution_id,
            estado=ExecutionStatus.SUCCESS,
            archivo_nombre=nombre_archivo,
            detalles=detalles_str,
        )

        content_response = {
            "success": True,
            "message": "Proceso de descarga terminado",
            "execution_id": execution_id,
            "archivo_esperado": nombre_archivo,
            "ruta_completa": ruta_completa,
            "procesamiento_automatico": detalles_str is not None,
        }
        if detalles_str and "falló" not in detalles_str and resultado_proc:
            content_response["message"] = detalles_str
            content_response["insertados"] = resultado_proc.get("insertados", 0)
            content_response["duplicados"] = resultado_proc.get("duplicados", 0)
            content_response["errores"] = resultado_proc.get("errores", [])
        elif detalles_str:
            content_response["procesamiento_error"] = detalles_str

        return JSONResponse(status_code=200, content=content_response)
        
    except Exception as e:
        err_msg = str(e)
        if execution_id is not None:
            await crud.update_execution_status(
                db=db, execution_id=execution_id, estado=ExecutionStatus.FAILED,
                detalles=f"Error al iniciar la descarga: {err_msg}"
            )
        return JSONResponse(
            status_code=500,
            content={"error": f"Error al iniciar la descarga: {err_msg}"}
        )


@app.post("/api/compras/procesar-archivos")
async def procesar_archivos(
    request: Request,
    archivo_ventas: UploadFile = File(...),
    archivo_po: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Procesa dos archivos Excel: cruza información usando 'Purchasing Document' y sube los datos a la tabla compras."""
    import tempfile
    import shutil
    import os
    from pathlib import Path
    import pandas as pd
    import traceback
    import time
    
    # Variables para el registro de ejecución
    execution = None
    execution_id = None
    fecha_inicio_ejecucion = None
    
    try:
        # Obtener información de la máquina
        try:
            hostname = socket.gethostname()
        except:
            hostname = platform.node() or "unknown"
        
        # Fecha de hoy para el periodo (el procesamiento no tiene periodo específico)
        fecha_hoy = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
        
        # Crear el registro de ejecución inicial (retorna id para evitar MissingGreenlet)
        execution_id = await crud.create_execution(
            db=db,
            user_id=current_user.id,
            fecha_inicio_periodo=fecha_hoy,
            fecha_fin_periodo=fecha_hoy,
            sistema_sap="Procesamiento de Archivos",
            transaccion="Procesar Archivos Excel",
            maquina=hostname
        )
        fecha_inicio_ejecucion = datetime.now()
        
        # Actualizar estado a RUNNING
        await crud.update_execution_status(
            db=db,
            execution_id=execution_id,
            estado=ExecutionStatus.RUNNING,
            fecha_inicio_ejecucion=fecha_inicio_ejecucion
        )
        
        # Validar que sean archivos Excel
        extensiones_permitidas = ['.xlsx', '.xls']
        
        nombre_archivo_ventas = archivo_ventas.filename or ''
        nombre_archivo_po = archivo_po.filename or ''
        
        extension_ventas = Path(nombre_archivo_ventas).suffix.lower()
        extension_po = Path(nombre_archivo_po).suffix.lower()
        
        if extension_ventas not in extensiones_permitidas:
            error_msg = "El archivo de Reporte de ventas debe ser un archivo Excel (.xlsx o .xls)"
            if execution_id:
                await crud.update_execution_status(
                    db=db,
                    execution_id=execution_id,
                    estado=ExecutionStatus.FAILED,
                    mensaje_error=error_msg
                )
            return JSONResponse(
                status_code=400,
                content={"error": error_msg}
            )
        
        if extension_po not in extensiones_permitidas:
            error_msg = "El archivo de Purchase Order History debe ser un archivo Excel (.xlsx o .xls)"
            if execution_id:
                await crud.update_execution_status(
                    db=db,
                    execution_id=execution_id,
                    estado=ExecutionStatus.FAILED,
                    mensaje_error=error_msg
                )
            return JSONResponse(
                status_code=400,
                content={"error": error_msg}
            )
        
        # Crear directorio temporal para los archivos
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            
            # Guardar archivos temporalmente
            archivo_ventas_path = temp_path / f"ventas_{nombre_archivo_ventas}"
            archivo_po_path = temp_path / f"po_{nombre_archivo_po}"
            
            with open(archivo_ventas_path, "wb") as f:
                shutil.copyfileobj(archivo_ventas.file, f)
            
            with open(archivo_po_path, "wb") as f:
                shutil.copyfileobj(archivo_po.file, f)
            
            # Leer archivos Excel
            try:
                # Leer Archivo 1 (base/destino) - donde se escribirá Proveedor
                if extension_ventas == '.xlsx':
                    df_archivo1 = pd.read_excel(archivo_ventas_path, engine='openpyxl')
                else:
                    df_archivo1 = pd.read_excel(archivo_ventas_path, engine='xlrd')
                
                # Leer Archivo 2 (referencia) - del que se obtendrá Name 1
                if extension_po == '.xlsx':
                    df_archivo2 = pd.read_excel(archivo_po_path, engine='openpyxl')
                else:
                    df_archivo2 = pd.read_excel(archivo_po_path, engine='xlrd')
            except Exception as e:
                error_msg = f"Error al leer los archivos Excel: {str(e)}"
                if execution_id:
                    await crud.update_execution_status(
                        db=db,
                        execution_id=execution_id,
                        estado=ExecutionStatus.FAILED,
                        mensaje_error=error_msg,
                        stack_trace=traceback.format_exc()
                    )
                return JSONResponse(
                    status_code=400,
                    content={"error": error_msg}
                )
            
            # Buscar la columna "Purchasing Document" en ambos archivos (tal cual esté escrita)
            columna_purchasing_doc = None
            for col in df_archivo1.columns:
                if str(col).strip() == "Purchasing Document":
                    columna_purchasing_doc = col
                    break
            
            if columna_purchasing_doc is None:
                error_msg = "No se encontró la columna 'Purchasing Document' en el Archivo 1 (base/destino)"
                if execution_id:
                    await crud.update_execution_status(
                        db=db,
                        execution_id=execution_id,
                        estado=ExecutionStatus.FAILED,
                        mensaje_error=error_msg
                    )
                return JSONResponse(
                    status_code=400,
                    content={"error": error_msg}
                )
            
            # Verificar que también existe en el Archivo 2
            columna_purchasing_doc_archivo2 = None
            for col in df_archivo2.columns:
                if str(col).strip() == "Purchasing Document":
                    columna_purchasing_doc_archivo2 = col
                    break
            
            if columna_purchasing_doc_archivo2 is None:
                error_msg = "No se encontró la columna 'Purchasing Document' en el Archivo 2 (referencia)"
                if execution_id:
                    await crud.update_execution_status(
                        db=db,
                        execution_id=execution_id,
                        estado=ExecutionStatus.FAILED,
                        mensaje_error=error_msg
                    )
                return JSONResponse(
                    status_code=400,
                    content={"error": error_msg}
                )
            
            # Buscar la columna "Name 1" en el Archivo 2
            columna_name1 = None
            for col in df_archivo2.columns:
                if str(col).strip() == "Name 1":
                    columna_name1 = col
                    break
            
            if columna_name1 is None:
                error_msg = "No se encontró la columna 'Name 1' en el Archivo 2 (referencia)"
                if execution_id:
                    await crud.update_execution_status(
                        db=db,
                        execution_id=execution_id,
                        estado=ExecutionStatus.FAILED,
                        mensaje_error=error_msg
                    )
                return JSONResponse(
                    status_code=400,
                    content={"error": error_msg}
                )
            
            # Buscar la columna "Supplier" en el Archivo 2 para codigo_proveedor
            columna_supplier = None
            for col in df_archivo2.columns:
                if str(col).strip().lower() == "supplier":
                    columna_supplier = col
                    break
            
            if columna_supplier is None:
                error_msg = "No se encontró la columna 'Supplier' en el Archivo 2 (referencia)"
                if execution_id:
                    await crud.update_execution_status(
                        db=db,
                        execution_id=execution_id,
                        estado=ExecutionStatus.FAILED,
                        mensaje_error=error_msg
                    )
                return JSONResponse(
                    status_code=400,
                    content={"error": error_msg}
                )
            
            # Buscar columna "Proveedor" original en Archivo 1 para codigo_proveedor (antes del cruce)
            # Esta columna puede contener el número del proveedor
            columna_proveedor_numero = None
            for col in df_archivo1.columns:
                if str(col).strip().lower() == "proveedor":
                    columna_proveedor_numero = col
                    break
            
            # Guardar los valores originales de "Proveedor" para codigo_proveedor si existe
            # Esto debe hacerse ANTES de que se modifique la columna "Proveedor" con el cruce
            if columna_proveedor_numero is not None:
                # Guardar una copia del valor original antes del cruce
                df_archivo1["Proveedor_Numero_Original"] = df_archivo1[columna_proveedor_numero].copy()
            else:
                # Si no existe "Proveedor" originalmente, crear columna vacía
                df_archivo1["Proveedor_Numero_Original"] = None
            
            # Buscar o crear columna "Proveedor" en Archivo 1 para nombre_proveedor (búsqueda case-insensitive)
            columna_proveedor = None
            for col in df_archivo1.columns:
                if str(col).strip().lower() == "proveedor":
                    columna_proveedor = col
                    break
            
            # Si no existe, crear la columna "Proveedor" para el nombre
            if columna_proveedor is None:
                df_archivo1["Proveedor"] = ""
                columna_proveedor = "Proveedor"
            
            # Crear columna para codigo_proveedor (Supplier del archivo 2)
            df_archivo1["Supplier_Numero"] = None
            
            # Crear diccionarios de mapeo desde Archivo 2
            # Clave: Purchasing Document (normalizado), Valor: Name 1 (para nombre_proveedor)
                    # Clave: Purchasing Document (normalizado), Valor: Supplier (para codigo_proveedor)
            # Usar la primera coincidencia si hay múltiples
            mapa_proveedores = {}
            mapa_suppliers = {}
            for idx, row in df_archivo2.iterrows():
                purchasing_doc = row[columna_purchasing_doc_archivo2]
                name1 = row[columna_name1]
                supplier = row[columna_supplier]
                
                # Normalizar: convertir a string, trim, y asegurar tipo comparable
                if pd.notna(purchasing_doc):
                    purchasing_doc_normalizado = str(purchasing_doc).strip()
                    
                    # Mapeo para nombre_proveedor (Name 1)
                    if pd.notna(name1) and purchasing_doc_normalizado and purchasing_doc_normalizado not in mapa_proveedores:
                        name1_valor = str(name1).strip()
                        mapa_proveedores[purchasing_doc_normalizado] = name1_valor
                    
                    # Mapeo para codigo_proveedor (Supplier)
                    if pd.notna(supplier) and purchasing_doc_normalizado and purchasing_doc_normalizado not in mapa_suppliers:
                        supplier_valor = str(supplier).strip()
                        mapa_suppliers[purchasing_doc_normalizado] = supplier_valor
            
            # Procesar cada fila del Archivo 1
            for idx, row in df_archivo1.iterrows():
                purchasing_doc_valor = row[columna_purchasing_doc]
                
                # Normalizar el valor para comparar
                if pd.notna(purchasing_doc_valor):
                    purchasing_doc_normalizado = str(purchasing_doc_valor).strip()
                    
                    # Buscar coincidencia en el mapa para nombre_proveedor (Name 1)
                    if purchasing_doc_normalizado in mapa_proveedores:
                        # Asignar el valor de Name 1 a Proveedor
                        df_archivo1.at[idx, columna_proveedor] = mapa_proveedores[purchasing_doc_normalizado]
                    else:
                        # Si no hay coincidencia, dejar vacío (o conservar si ya existe)
                        if pd.isna(df_archivo1.at[idx, columna_proveedor]) or df_archivo1.at[idx, columna_proveedor] == "":
                            df_archivo1.at[idx, columna_proveedor] = ""
                    
                    # Buscar coincidencia en el mapa para codigo_proveedor (Supplier)
                    if purchasing_doc_normalizado in mapa_suppliers:
                        # Asignar el valor de Supplier a Supplier_Numero
                        df_archivo1.at[idx, "Supplier_Numero"] = mapa_suppliers[purchasing_doc_normalizado]
                    else:
                        # Si no hay coincidencia, dejar None
                        df_archivo1.at[idx, "Supplier_Numero"] = None
                else:
                    # Si el valor de Purchasing Document está vacío, dejar vacío
                    if pd.isna(df_archivo1.at[idx, columna_proveedor]) or df_archivo1.at[idx, columna_proveedor] == "":
                        df_archivo1.at[idx, columna_proveedor] = ""
                    df_archivo1.at[idx, "Supplier_Numero"] = None
            
            # Crear columna "U/P" = Invoice Value / Quantity in OPUn
            # Buscar columnas de manera case-insensitive
            columna_invoice_value = None
            columna_quantity_opun = None
            
            for col in df_archivo1.columns:
                col_normalizado = str(col).strip().lower()
                if col_normalizado == "invoice value":
                    columna_invoice_value = col
                elif col_normalizado == "quantity in opun":
                    columna_quantity_opun = col
            
            # Verificar que ambas columnas existan
            if columna_invoice_value is None:
                error_msg = "No se encontró la columna 'Invoice Value' en el Archivo 1 (base/destino)"
                if execution_id:
                    await crud.update_execution_status(
                        db=db,
                        execution_id=execution_id,
                        estado=ExecutionStatus.FAILED,
                        mensaje_error=error_msg
                    )
                return JSONResponse(
                    status_code=400,
                    content={"error": error_msg}
                )
            
            if columna_quantity_opun is None:
                error_msg = "No se encontró la columna 'Quantity in OPUn' en el Archivo 1 (base/destino)"
                if execution_id:
                    await crud.update_execution_status(
                        db=db,
                        execution_id=execution_id,
                        estado=ExecutionStatus.FAILED,
                        mensaje_error=error_msg
                    )
                return JSONResponse(
                    status_code=400,
                    content={"error": error_msg}
                )
            
            # Crear la columna "U/P" y calcular el valor
            df_archivo1["U/P"] = ""
            
            for idx, row in df_archivo1.iterrows():
                invoice_value = row[columna_invoice_value]
                quantity_opun = row[columna_quantity_opun]
                
                # Convertir a numérico si es posible
                try:
                    invoice_value_num = pd.to_numeric(invoice_value, errors='coerce')
                    quantity_opun_num = pd.to_numeric(quantity_opun, errors='coerce')
                    
                    # Realizar la división solo si ambos valores son válidos y quantity no es cero
                    if pd.notna(invoice_value_num) and pd.notna(quantity_opun_num) and quantity_opun_num != 0:
                        df_archivo1.at[idx, "U/P"] = invoice_value_num / quantity_opun_num
                    else:
                        # Si hay valores nulos o quantity es cero, dejar vacío
                        df_archivo1.at[idx, "U/P"] = ""
                except (ValueError, TypeError, ZeroDivisionError):
                    # Si hay error en la conversión o división, dejar vacío
                    df_archivo1.at[idx, "U/P"] = ""
            
            # Mapear columnas del DataFrame a la tabla compras
            # Función auxiliar para convertir valores
            def convertir_valor(valor, tipo='int'):
                """Convierte un valor de pandas a tipo Python apropiado."""
                if pd.isna(valor) or valor == '' or valor is None:
                    return None
                try:
                    if tipo == 'int':
                        # Para valores grandes, Python maneja automáticamente enteros de cualquier tamaño
                        # Convertir directamente sin pasar por float para evitar pérdida de precisión
                        if isinstance(valor, int):
                            return valor
                        elif isinstance(valor, float):
                            # Si es float, convertir a int (puede manejar valores grandes)
                            return int(valor)
                        else:
                            # Para strings u otros tipos, convertir directamente
                            # Usar pd.to_numeric para manejar mejor los valores grandes
                            valor_numerico = pd.to_numeric(valor, errors='coerce')
                            if pd.notna(valor_numerico):
                                return int(valor_numerico)
                            return None
                    elif tipo == 'float':
                        return float(valor) if pd.notna(valor) else None
                    elif tipo == 'date':
                        if isinstance(valor, pd.Timestamp):
                            return valor.to_pydatetime()
                        elif isinstance(valor, datetime):
                            return valor
                        elif isinstance(valor, str):
                            return pd.to_datetime(valor, errors='coerce').to_pydatetime() if pd.notna(pd.to_datetime(valor, errors='coerce')) else None
                        return None
                    else:
                        return str(valor).strip() if pd.notna(valor) else None
                except (ValueError, TypeError, OverflowError):
                    return None
            
            # Función para mapear una fila del DataFrame a un diccionario para Compra
            def mapear_fila_a_compra(row, df_columns):
                """Mapea una fila del DataFrame a un diccionario para insertar en Compra."""
                compra_data = {}
                
                # Mapeo de columnas (case-insensitive)
                columnas_mapeo = {
                    'purchasing_document': ['Purchasing Document'],
                    'item': ['Item'],
                    'material_doc_year': ['Material Doc. Year'],
                    'material_document': ['Material Document'],
                    'material_doc_item': ['Material Doc.Item'],
                    'movement_type': ['Movement Type'],
                    'posting_date': ['Posting Date'],
                    'quantity': ['Quantity'],
                    'order_unit': ['Order Unit'],
                    'quantity_in_opun': ['Quantity in OPUn'],
                    'order_price_unit': ['Order Price Unit'],
                    'amount_in_lc': ['Amount in LC'],
                    'local_currency': ['Local currency'],
                    'amount': ['Amount'],
                    'currency': ['Currency'],
                    'gr_ir_clearing_value_lc': ['GR/IR clearing value in local currency'],
                    'gr_blck_stock_oun': ['GR Blck.Stock in OUn'],
                    'gr_blocked_stck_opun': ['GR blocked stck.OPUn'],
                    'delivery_completed': ['Delivery Completed'],
                    'fisc_year_ref_doc': ['Fisc. Year Ref. Doc.'],
                    'reference_document': ['Reference Document'],
                    'reference_doc_item': ['Reference Doc. Item'],
                    'invoice_value': ['Invoice Value'],
                    'numero_material': ['numero_Material', 'Material'],
                    'plant': ['Plant'],
                    'descripcion_material': ['Short Text', 'descripcion_material', 'Material Description', 'Short text'],
                    'nombre_proveedor': ['Supplier', 'supplier', 'SUPPLIER', 'Proveedor', 'proveedor', 'PROVEEDOR', 'nombre_proveedor'],
                    'codigo_proveedor': ['Supplier_Numero', 'Proveedor_Numero_Original', 'Proveedor', 'proveedor', 'PROVEEDOR', 'codigo_proveedor'],
                    'price': ['U/P', 'u/p', 'U/p', 'u/P', 'precio', 'Precio', 'PRECIO', 'Price', 'price', 'PRICE'],
                }
                
                # Buscar y mapear cada columna
                for campo_db, nombres_posibles in columnas_mapeo.items():
                    valor = None
                    for nombre_col in nombres_posibles:
                        # Buscar columna case-insensitive
                        for col in df_columns:
                            if str(col).strip().lower() == str(nombre_col).strip().lower():
                                valor = row[col]
                                break
                        if valor is not None and not pd.isna(valor):
                            break
                    
                    # Convertir según el tipo de campo
                    if campo_db in ['purchasing_document', 'item', 'material_doc_year', 'material_document', 
                                    'material_doc_item', 'quantity', 'quantity_in_opun']:
                        compra_data[campo_db] = convertir_valor(valor, 'int')
                    elif campo_db == 'codigo_proveedor':
                        # codigo_proveedor es un string, convertir a string directamente
                        if valor is not None and not pd.isna(valor):
                            compra_data[campo_db] = str(valor).strip()
                        else:
                            compra_data[campo_db] = None
                    elif campo_db == 'posting_date':
                        compra_data[campo_db] = convertir_valor(valor, 'date')
                    elif campo_db in ['amount_in_lc', 'amount', 'gr_ir_clearing_value_lc', 'invoice_value', 'price']:
                        compra_data[campo_db] = convertir_valor(valor, 'float')
                    elif campo_db == 'numero_material':
                        # Evitar notación científica (ej. Excel 3.41301E+11 -> 341301000000)
                        raw = convertir_valor(valor, 'str')
                        compra_data[campo_db] = (crud._normalizar_numero_material(raw) if raw else None)
                    else:
                        # plant, descripcion_material, nombre_proveedor y otros string
                        compra_data[campo_db] = convertir_valor(valor, 'str')
                
                return compra_data
            
            # Preparar datos para inserción en la base de datos
            compras_data = []
            for idx, row in df_archivo1.iterrows():
                compra_data = mapear_fila_a_compra(row, df_archivo1.columns)
                compras_data.append(compra_data)
            
            # Insertar datos en la base de datos (duplicados no se insertan; se cuentan)
            try:
                resultado = await crud.bulk_create_or_update_compras(db, compras_data)
                registros_insertados = resultado["insertados"]
                registros_duplicados = resultado["duplicados"]
                errores_compra = resultado.get("errores", [])
            except Exception as e:
                error_msg = f"Error al insertar datos en la base de datos: {str(e)}"
                if execution_id:
                    await crud.update_execution_status(
                        db=db,
                        execution_id=execution_id,
                        estado=ExecutionStatus.FAILED,
                        mensaje_error=error_msg,
                        detalles=error_msg,
                        stack_trace=traceback.format_exc()
                    )
                return JSONResponse(
                    status_code=500,
                    content={"error": error_msg}
                )
            
            # Calcular duración
            fecha_fin_ejecucion = datetime.now()
            duracion_segundos = int((fecha_fin_ejecucion - fecha_inicio_ejecucion).total_seconds())
            
            # Construir mensaje/detalles: insertados, duplicados, errores
            mensaje = f"Datos procesados exitosamente. {registros_insertados} registros insertados."
            if registros_duplicados > 0:
                mensaje += f" {registros_duplicados} registros duplicados (ya existían, no se insertaron de nuevo)."
            if errores_compra:
                mensaje += f" {len(errores_compra)} error(es) al procesar algunos registros."
            
            # Actualizar ejecución como exitosa con detalles
            if execution_id:
                await crud.update_execution_status(
                    db=db,
                    execution_id=execution_id,
                    estado=ExecutionStatus.SUCCESS,
                    fecha_fin_ejecucion=fecha_fin_ejecucion,
                    duracion_segundos=duracion_segundos,
                    detalles=mensaje
                )
            
            return JSONResponse(
                status_code=200,
                content={
                    "success": True,
                    "message": mensaje,
                    "registros_insertados": registros_insertados,
                    "registros_duplicados": registros_duplicados,
                    "errores": errores_compra,
                    "archivos_recibidos": {
                        "archivo_base": nombre_archivo_ventas,
                        "archivo_referencia": nombre_archivo_po
                    },
                    "execution_id": execution_id
                }
            )
        
    except Exception as e:
        # Manejar cualquier error no capturado
        error_msg = f"Error al procesar los archivos: {str(e)}"
        stack_trace_str = traceback.format_exc()
        
        if execution_id:
            fecha_fin_ejecucion = datetime.now()
            duracion_segundos = int((fecha_fin_ejecucion - fecha_inicio_ejecucion).total_seconds()) if fecha_inicio_ejecucion else None
            
            await crud.update_execution_status(
                db=db,
                execution_id=execution_id,
                estado=ExecutionStatus.FAILED,
                fecha_fin_ejecucion=fecha_fin_ejecucion,
                duracion_segundos=duracion_segundos,
                mensaje_error=error_msg,
                stack_trace=stack_trace_str
            )
        
        return JSONResponse(
            status_code=500,
            content={"error": error_msg}
        )


@app.post("/api/ventas/iniciar-descarga")
async def iniciar_descarga_ventas(
    request: Request,
    periodo_mes_anio: str = Form(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Inicia el proceso de descarga de ventas ejecutando ventas.vbs con SAP GUI. periodo_mes_anio en formato YYYY-MM. El script recibe el periodo en formato PPP.YYYY (ej. 001.2026)."""
    import subprocess
    import asyncio
    from pathlib import Path
    execution_id = None
    try:
        # Validar y convertir periodo (YYYY-MM) a primer/último día y a formato SAP PPP.YYYY
        try:
            from calendar import monthrange
            anio, mes = int(periodo_mes_anio[:4]), int(periodo_mes_anio[5:7])
            if mes < 1 or mes > 12:
                raise ValueError("Mes inválido")
            ultimo_dia = monthrange(anio, mes)[1]
            fecha_inicio_dt = datetime(anio, mes, 1)
            fecha_fin_dt = datetime(anio, mes, ultimo_dia, 23, 59, 59)
            # Formato SAP: PPP.YYYY (001 = enero, 002 = febrero, ... 012 = diciembre)
            periodo_sap = f"{mes:03d}.{anio}"
        except (ValueError, IndexError) as e:
            return JSONResponse(
                status_code=400,
                content={"error": f"Formato de periodo inválido. Use formato YYYY-MM (mes y año). {str(e)}"}
            )
        
        # Obtener información de la máquina
        try:
            hostname = socket.gethostname()
        except:
            hostname = platform.node() or "unknown"
        
        # Crear el registro de ejecución (retorna id para evitar MissingGreenlet)
        execution_id = await crud.create_sales_execution(
            db=db,
            user_id=current_user.id,
            fecha_inicio_periodo=fecha_inicio_dt,
            fecha_fin_periodo=fecha_fin_dt,
            sistema_sap="SAP ECC",
            transaccion="KE30",
            maquina=hostname
        )
        
        carpeta_salida = "C:\\Users\\anad5004\\Documents\\Leoni_RPA"
        fecha_inicio_str = fecha_inicio_dt.strftime("%Y%m%d")
        fecha_fin_str = fecha_fin_dt.strftime("%Y%m%d")
        nombre_archivo = f"KE30_US10_{periodo_sap.replace('.', '_')}.xlsx"
        ruta_completa = f"{carpeta_salida.rstrip('/').rstrip(chr(92))}{os.sep}{nombre_archivo}"
        
        await crud.update_sales_execution_status(
            db=db,
            execution_id=execution_id,
            estado=ExecutionStatus.PENDING,
            archivo_ruta=ruta_completa,
            archivo_nombre=nombre_archivo
        )
        
        # Ejecutar ventas.vbs con el periodo en formato PPP.YYYY
        script_dir = Path(__file__).resolve().parent
        vbs_path = script_dir / "ventas.vbs"
        if not vbs_path.exists():
            err = f"No se encontró el script ventas.vbs en {script_dir}"
            await crud.update_sales_execution_status(
                db=db, execution_id=execution_id, estado=ExecutionStatus.FAILED,
                mensaje_error=err, detalles=err
            )
            return JSONResponse(
                status_code=500,
                content={"error": f"No se encontró el script ventas.vbs en {script_dir}"}
            )
        
        loop = asyncio.get_running_loop()
        def _run_ventas_vbs():
            return subprocess.run(
                ["cscript", "//nologo", str(vbs_path), periodo_sap],
                cwd=str(script_dir),
                capture_output=True,
                timeout=300,
            )
        try:
            result = await loop.run_in_executor(None, _run_ventas_vbs)
        except subprocess.TimeoutExpired:
            err = "El proceso de descarga superó el tiempo máximo (5 minutos)."
            await crud.update_sales_execution_status(
                db=db, execution_id=execution_id, estado=ExecutionStatus.FAILED,
                mensaje_error=err, detalles=err
            )
            return JSONResponse(
                status_code=500,
                content={"error": "El proceso de descarga superó el tiempo máximo (5 minutos)."}
            )
        except Exception as e:
            err = f"Error al ejecutar el script: {str(e)}"
            await crud.update_sales_execution_status(
                db=db, execution_id=execution_id, estado=ExecutionStatus.FAILED,
                mensaje_error=err, detalles=err
            )
            return JSONResponse(
                status_code=500,
                content={"error": err}
            )
        
        if result.returncode != 0:
            err_msg = (result.stderr or result.stdout or b"").decode("utf-8", errors="replace").strip() or f"Código de salida: {result.returncode}"
            err_full = f"El script SAP finalizó con error. {err_msg}"
            await crud.update_sales_execution_status(
                db=db, execution_id=execution_id, estado=ExecutionStatus.FAILED,
                mensaje_error=err_full, detalles=err_full
            )
            return JSONResponse(
                status_code=500,
                content={"error": f"El script SAP finalizó con error. {err_msg}"}
            )
        
        # Procesar automáticamente el Excel generado (mismo flujo que compras)
        path_ventas_xlsx = Path(carpeta_salida) / f"KE30_US10_{periodo_sap.replace('.', '_')}.xlsx"
        content_response = {
            "success": True,
            "message": "Proceso de descarga finalizado correctamente",
            "execution_id": execution_id,
            "archivo_esperado": nombre_archivo,
            "ruta_completa": ruta_completa,
            "periodo_sap": periodo_sap,
        }
        if path_ventas_xlsx.exists():
            try:
                resultado_proceso = await _procesar_archivo_ventas_desde_ruta(
                    db, path_ventas_xlsx, execution_id, current_user.id
                )
                content_response["procesamiento_automatico"] = True
                content_response["registros_insertados"] = resultado_proceso.get("registros_insertados", 0)
                content_response["registros_duplicados"] = resultado_proceso.get("registros_duplicados", 0)
                content_response["message_procesamiento"] = resultado_proceso.get("message", "")
                content_response["precios_venta_actualizados"] = resultado_proceso.get("precios_venta_actualizados")
                content_response["precios_venta_sin_cambios"] = resultado_proceso.get("precios_venta_sin_cambios")
                content_response["precios_venta_sin_venta"] = resultado_proceso.get("precios_venta_sin_venta")
                content_response["precios_venta_error"] = resultado_proceso.get("precios_venta_error")
                if resultado_proceso.get("success"):
                    try:
                        path_ventas_xlsx.unlink()
                    except Exception:
                        pass
            except Exception as e:
                content_response["procesamiento_automatico"] = True
                content_response["error_procesamiento"] = str(e)
        else:
            await crud.update_sales_execution_status(
                db=db, execution_id=execution_id, estado=ExecutionStatus.SUCCESS,
                detalles="Descarga finalizada correctamente. No se encontró archivo Excel para procesar."
            )
        
        return JSONResponse(status_code=200, content=content_response)
        
    except Exception as e:
        err_msg = str(e)
        err_full = f"Error al iniciar la descarga: {err_msg}"
        if execution_id is not None:
            await crud.update_sales_execution_status(
                db=db, execution_id=execution_id, estado=ExecutionStatus.FAILED,
                mensaje_error=err_full, detalles=err_full
            )
        return JSONResponse(
            status_code=500,
            content={"error": err_full}
        )


def _procesar_df_historial(df: "pd.DataFrame") -> "pd.DataFrame":
    """Aplica transformaciones al historial: eliminar columnas, filtros, precio unitario."""
    import pandas as pd

    columnas_a_eliminar = [
        'GR Blck.Stock in OUn', 'GR blocked stck.OPUn', 'Delivery Completed',
        'Fisc. Year Ref. Doc.', 'Reference Document', 'Reference Doc. Item'
    ]
    columnas_encontradas = []
    for col in df.columns:
        col_str = str(col).strip()
        for col_eliminar in columnas_a_eliminar:
            if col_str == col_eliminar or col_str.lower() == col_eliminar.lower():
                columnas_encontradas.append(col)
                break
    if columnas_encontradas:
        df = df.drop(columns=columnas_encontradas)

    if 'nombre proveedor' not in df.columns:
        df['nombre proveedor'] = ''
    if 'codigo proveedor' not in df.columns:
        df['codigo proveedor'] = ''

    col_plant = next((c for c in df.columns if str(c).strip().lower() == 'plant'), None)
    if col_plant is not None:
        df[col_plant] = df[col_plant].astype(str).str.strip()
        df = df[(~df[col_plant].isin(['MX10', 'MX11'])) & (df[col_plant] != '') & (df[col_plant] != 'nan') & (df[col_plant].notna())]

    col_material = next((c for c in df.columns if str(c).strip().lower() == 'material'), None)
    if col_material is not None:
        df[col_material] = df[col_material].astype(str).str.strip()
        df = df[(df[col_material] != '') & (df[col_material] != 'nan') & (df[col_material].notna())]

    col_inv = next((c for c in df.columns if str(c).strip().lower() == 'invoice value'), None)
    if col_inv is not None:
        vals = df[col_inv].astype(str).str.strip()
        df = df[~vals.isin(['0', '0.0', '0.00'])]

    col_inv = next((c for c in df.columns if str(c).strip().lower() == 'amount'), None)
    col_qty = next((c for c in df.columns if str(c).strip().lower() == 'quantity'), None)
    if col_inv is not None and col_qty is not None:
        inv_n = pd.to_numeric(df[col_inv], errors='coerce')
        qty_n = pd.to_numeric(df[col_qty], errors='coerce')
        df['precio unitario'] = inv_n / qty_n
        df.loc[(qty_n <= 0) | qty_n.isna(), 'precio unitario'] = pd.NA
        df['precio unitario'] = df['precio unitario'].replace([float('inf'), float('-inf')], pd.NA)
    else:
        df['precio unitario'] = pd.NA

    return df


async def _procesar_compras_historial_desde_rutas(db: AsyncSession, path_compras, path_historial):
    """
    Procesa los dos Excel de compras e historial desde rutas en disco (Path o str).
    Retorna dict: {success, insertados?, duplicados?, errores?, message?} o {success: False, error}.
    """
    from pathlib import Path
    import pandas as pd
    path_compras = Path(path_compras)
    path_historial = Path(path_historial)
    from decimal import Decimal
    from datetime import datetime as dt

    def col(df, name):
        name = name.strip().lower()
        for c in df.columns:
            if str(c).strip().lower() == name:
                return c
        return None

    ext_ok = ['.xlsx', '.xls']
    e_compras = path_compras.suffix.lower()
    e_hist = path_historial.suffix.lower()
    if e_compras not in ext_ok or e_hist not in ext_ok:
        return {"success": False, "error": "Ambos archivos deben ser Excel (.xlsx o .xls)"}

    try:
        df_compras = pd.read_excel(path_compras, engine='openpyxl' if e_compras == '.xlsx' else 'xlrd')
        df = pd.read_excel(path_historial, engine='openpyxl' if e_hist == '.xlsx' else 'xlrd')
    except Exception as e:
        return {"success": False, "error": f"Error al leer archivos Excel: {str(e)}"}

    pc = col(df_compras, 'Purchasing Document')
    sc = col(df_compras, 'Supplier')
    n1 = col(df_compras, 'Name 1')
    if not all([pc, sc, n1]):
        return {"success": False, "error": "El archivo Compras debe tener columnas: Purchasing Document, Supplier, Name 1"}

    def to_int_codigo(v):
        if pd.isna(v) or v == '' or v is None:
            return pd.NA
        try:
            return int(pd.to_numeric(v, errors='raise'))
        except (ValueError, TypeError):
            return pd.NA

    mapa = {}
    for _, row in df_compras.iterrows():
        k = str(row[pc]).strip() if pd.notna(row[pc]) else None
        if not k or k in mapa:
            continue
        mapa[k] = (
            to_int_codigo(row[sc]),
            str(row[n1]).strip() if pd.notna(row[n1]) else ''
        )

    df = _procesar_df_historial(df)

    col_pd = col(df, 'Purchasing Document')
    if col_pd is None:
        return {"success": False, "error": "El archivo Historial debe tener la columna Purchasing Document"}

    def lookup(key):
        k = str(key).strip() if pd.notna(key) else None
        return mapa.get(k, (pd.NA, ''))

    codigos = df[col_pd].map(lambda x: lookup(x)[0])
    df['codigo proveedor'] = codigos.astype('Int64')
    df['nombre proveedor'] = df[col_pd].map(lambda x: lookup(x)[1])

    col_pd_map = col(df, 'Purchasing Document')
    col_item = col(df, 'Item')
    col_mdy = col(df, 'Material Doc. Year')
    col_md = col(df, 'Material Document')
    col_mdi = col(df, 'Material Doc.Item')
    col_mt = col(df, 'Movement Type')
    col_post = col(df, 'Posting Date')
    col_qty = col(df, 'Quantity')
    col_ou = col(df, 'Order Unit')
    col_qop = col(df, 'Quantity in OPUn')
    col_opu = col(df, 'Order Price Unit')
    col_alc = col(df, 'Amount in LC')
    col_lc = col(df, 'Local currency')
    col_amt = col(df, 'Amount')
    col_curr = col(df, 'Currency')
    col_grir = col(df, 'GR/IR clearing value in local currency')
    col_inv = col(df, 'Invoice Value')
    col_mat = col(df, 'Material')
    col_plant = col(df, 'Plant')
    col_desc = col(df, 'Short Text') or col(df, 'descripcion_material') or col(df, 'Material Description') or col(df, 'Short text')
    col_nom_prov = col(df, 'nombre proveedor')
    col_cod_prov = col(df, 'codigo proveedor')
    col_price = col(df, 'precio unitario')

    def safe_int(val):
        if pd.isna(val) or val == '' or val is None:
            return None
        try:
            return int(pd.to_numeric(val, errors='raise'))
        except (ValueError, TypeError):
            return None

    def safe_decimal(val):
        if pd.isna(val) or val == '' or val is None:
            return None
        try:
            return Decimal(str(val))
        except (ValueError, TypeError):
            return None

    def safe_str(val):
        if pd.isna(val) or val == '' or val is None:
            return None
        return str(val).strip()

    def safe_date(val):
        if pd.isna(val) or val == '' or val is None:
            return None
        try:
            if isinstance(val, pd.Timestamp):
                dt_val = val.to_pydatetime()
            elif isinstance(val, dt):
                dt_val = val
            else:
                dt_val = pd.to_datetime(val).to_pydatetime()
            if dt_val.tzinfo is None:
                dt_val = dt_val.replace(tzinfo=ZoneInfo('America/Mexico_City'))
            return dt_val
        except Exception:
            return None

    compras_data = []
    for _, row in df.iterrows():
        compra_dict = {
            'purchasing_document': safe_int(row[col_pd_map]) if col_pd_map else None,
            'item': safe_int(row[col_item]) if col_item else None,
            'material_doc_year': safe_int(row[col_mdy]) if col_mdy else None,
            'material_document': safe_int(row[col_md]) if col_md else None,
            'material_doc_item': safe_int(row[col_mdi]) if col_mdi else None,
            'movement_type': safe_str(row[col_mt]) if col_mt else None,
            'posting_date': safe_date(row[col_post]) if col_post else None,
            'quantity': safe_int(row[col_qty]) if col_qty else None,
            'order_unit': safe_str(row[col_ou]) if col_ou else None,
            'quantity_in_opun': safe_int(row[col_qop]) if col_qop else None,
            'order_price_unit': safe_str(row[col_opu]) if col_opu else None,
            'amount_in_lc': safe_decimal(row[col_alc]) if col_alc else None,
            'local_currency': safe_str(row[col_lc]) if col_lc else None,
            'amount': safe_decimal(row[col_amt]) if col_amt else None,
            'currency': safe_str(row[col_curr]) if col_curr else None,
            'gr_ir_clearing_value_lc': safe_decimal(row[col_grir]) if col_grir else None,
            'invoice_value': safe_decimal(row[col_inv]) if col_inv else None,
            'numero_material': safe_str(row[col_mat]) if col_mat else None,
            'plant': safe_str(row[col_plant]) if col_plant else None,
            'descripcion_material': safe_str(row[col_desc]) if col_desc else None,
            'nombre_proveedor': safe_str(row[col_nom_prov]) if col_nom_prov else None,
            'codigo_proveedor': safe_str(row[col_cod_prov]) if col_cod_prov and pd.notna(row[col_cod_prov]) else None,
            'price': safe_decimal(row[col_price]) if col_price else None,
        }
        # Evitar notación científica en numero_material (ej. Excel 3.42E+11 -> 342000000000)
        if compra_dict.get('numero_material'):
            n_norm = crud._normalizar_numero_material(compra_dict['numero_material'])
            compra_dict['numero_material'] = n_norm if n_norm else None
        if compra_dict.get('purchasing_document') or compra_dict.get('numero_material'):
            compras_data.append(compra_dict)

    if not compras_data:
        return {"success": False, "error": "No se encontraron datos válidos para insertar"}

    resultado = await crud.bulk_create_or_update_compras(db, compras_data)
    ins = resultado['insertados']
    dup = resultado['duplicados']
    errs = resultado.get('errores', [])
    msg = f"Procesamiento completado: {ins} insertados, {dup} duplicados (ya existían, no se insertaron)."
    if errs:
        msg += f" {len(errs)} error(es) al procesar algunos registros."
    return {
        "success": True,
        "insertados": ins,
        "duplicados": dup,
        "errores": errs,
        "message": msg,
    }


@app.post("/api/compras/procesar-compras-historial")
async def procesar_compras_historial(
    request: Request,
    archivo_compras: UploadFile = File(...),
    archivo_historial: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Procesa historial de compras, enriquece con Compras (Supplier->codigo proveedor, Name 1->nombre proveedor) e inserta en la tabla compras."""
    import tempfile
    import shutil
    from pathlib import Path
    import pandas as pd

    ext_ok = ['.xlsx', '.xls']
    n_compras = archivo_compras.filename or ''
    n_hist = archivo_historial.filename or ''
    e_compras = Path(n_compras).suffix.lower()
    e_hist = Path(n_hist).suffix.lower()

    if e_compras not in ext_ok or e_hist not in ext_ok:
        return JSONResponse(status_code=400, content={"error": "Ambos archivos deben ser Excel (.xlsx o .xls)"})

    try:
        with tempfile.TemporaryDirectory() as tmp:
            tmp = Path(tmp)
            path_compras = tmp / f"compras_{n_compras}"
            path_hist = tmp / f"hist_{n_hist}"

            with open(path_compras, "wb") as f:
                shutil.copyfileobj(archivo_compras.file, f)
            with open(path_hist, "wb") as f:
                shutil.copyfileobj(archivo_historial.file, f)

            resultado = await _procesar_compras_historial_desde_rutas(db, path_compras, path_hist)

        if resultado["success"]:
            # Sincronizar proveedores, materiales, precios de compra y países de origen (como en descarga)
            msg_extra = []
            try:
                r_prov = await crud.sincronizar_proveedores_desde_compras(db, user_id=current_user.id)
                msg_extra.append(f"Proveedores: {r_prov.get('nuevos_creados', 0)} nuevo(s)")
            except Exception as e_prov:
                msg_extra.append(f"Proveedores (error): {str(e_prov)[:60]}")
            try:
                r_mat = await crud.sincronizar_materiales_desde_compras(db, user_id=current_user.id)
                msg_extra.append(f"Materiales: {r_mat.get('nuevos_creados', 0)} nuevo(s)")
            except Exception as e_mat:
                msg_extra.append(f"Materiales (error): {str(e_mat)[:60]}")
            try:
                r_precios = await crud.sincronizar_precios_materiales_desde_compras(db, user_id=current_user.id)
                msg_extra.append(f"Precios compra: {r_precios.get('nuevos_creados', 0)} nuevo(s), {r_precios.get('actualizados', 0)} actualizado(s)")
            except Exception as e_pre:
                msg_extra.append(f"Precios compra (error): {str(e_pre)[:60]}")
            try:
                r_paises = await crud.sincronizar_paises_origen_desde_compras(db, user_id=current_user.id)
                msg_extra.append(f"Países origen: {r_paises.get('nuevos_creados', 0)} nuevo(s)")
            except Exception as e_pai:
                msg_extra.append(f"Países origen (error): {str(e_pai)[:60]}")
            message = resultado["message"]
            if msg_extra:
                message += " " + "; ".join(msg_extra) + "."
            return JSONResponse(
                status_code=200,
                content={
                    "success": True,
                    "message": message,
                    "insertados": resultado["insertados"],
                    "duplicados": resultado["duplicados"],
                    "errores": resultado.get("errores", []),
                }
            )
        status_code = 400 if "columnas" in resultado.get("error", "") or "datos válidos" in resultado.get("error", "") else 500
        return JSONResponse(status_code=status_code, content={"error": resultado["error"]})

    except Exception as e:
        import traceback
        return JSONResponse(
            status_code=500,
            content={"error": f"Error al procesar: {str(e)}\n{traceback.format_exc()}"}
        )


@app.post("/api/compras/procesar-historial")
async def procesar_historial_compras(
    request: Request,
    archivo_po: UploadFile = File(...),
    carpeta_salida: str = Form(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Procesa el archivo de historial de compras: elimina columnas específicas y guarda el archivo procesado."""
    import tempfile
    import shutil
    import os
    from pathlib import Path
    import pandas as pd
    import traceback
    
    try:
        # Validar que sea un archivo Excel
        extensiones_permitidas = ['.xlsx', '.xls']
        nombre_archivo_po = archivo_po.filename or ''
        extension_po = Path(nombre_archivo_po).suffix.lower()
        
        if extension_po not in extensiones_permitidas:
            return JSONResponse(
                status_code=400,
                content={"error": "El archivo debe ser un archivo Excel (.xlsx o .xls)"}
            )
        
        # Validar carpeta de salida
        carpeta_salida_path = Path(carpeta_salida)
        if not carpeta_salida_path.exists():
            return JSONResponse(
                status_code=400,
                content={"error": f"La carpeta de salida no existe: {carpeta_salida}"}
            )
        
        if not carpeta_salida_path.is_dir():
            return JSONResponse(
                status_code=400,
                content={"error": f"La ruta especificada no es una carpeta: {carpeta_salida}"}
            )
        
        # Crear directorio temporal para el archivo
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            
            # Guardar archivo temporalmente
            archivo_po_path = temp_path / f"po_{nombre_archivo_po}"
            
            with open(archivo_po_path, "wb") as f:
                shutil.copyfileobj(archivo_po.file, f)
            
            # Leer archivo Excel
            try:
                if extension_po == '.xlsx':
                    df = pd.read_excel(archivo_po_path, engine='openpyxl')
                else:
                    df = pd.read_excel(archivo_po_path, engine='xlrd')
            except Exception as e:
                return JSONResponse(
                    status_code=400,
                    content={"error": f"Error al leer el archivo Excel: {str(e)}"}
                )
            
            # Columnas a eliminar (con posibles variaciones)
            columnas_a_eliminar = [
                'GR Blck.Stock in OUn',
                'GR blocked stck.OPUn',
                'Delivery Completed',
                'Fisc. Year Ref. Doc.',
                'Reference Document',
                'Reference Doc. Item'
            ]
            
            # Buscar y eliminar columnas (comparación exacta y case-insensitive)
            columnas_encontradas = []
            for col in df.columns:
                col_str = str(col).strip()
                for col_eliminar in columnas_a_eliminar:
                    # Comparación exacta
                    if col_str == col_eliminar:
                        columnas_encontradas.append(col)
                        break
                    # Comparación case-insensitive
                    elif col_str.lower() == col_eliminar.lower():
                        columnas_encontradas.append(col)
                        break
            
            # Eliminar las columnas encontradas
            if columnas_encontradas:
                df = df.drop(columns=columnas_encontradas)
            
            # Agregar nuevas columnas: nombre proveedor y codigo proveedor
            # Verificar si las columnas ya existen para no duplicarlas
            if 'nombre proveedor' not in df.columns:
                df['nombre proveedor'] = ''
            if 'codigo proveedor' not in df.columns:
                df['codigo proveedor'] = ''
            
            # Eliminar filas donde la columna "Plant" tenga "MX10", "MX11" o esté vacía
            # Buscar la columna "Plant" (case-insensitive)
            columna_plant = None
            for col in df.columns:
                if str(col).strip().lower() == 'plant':
                    columna_plant = col
                    break
            
            if columna_plant is not None:
                # Contar filas antes de eliminar
                filas_antes = len(df)
                # Convertir a string y limpiar espacios
                df[columna_plant] = df[columna_plant].astype(str).str.strip()
                # Eliminar filas donde Plant sea "MX10", "MX11" o esté vacía (nan, None, o string vacío)
                df = df[
                    (~df[columna_plant].isin(['MX10', 'MX11'])) & 
                    (df[columna_plant] != '') & 
                    (df[columna_plant] != 'nan') & 
                    (df[columna_plant].notna())
                ]
                filas_eliminadas = filas_antes - len(df)
            else:
                filas_eliminadas = 0
            
            # Eliminar filas donde la columna "Material" esté vacía
            # Buscar la columna "Material" (case-insensitive)
            columna_material = None
            for col in df.columns:
                if str(col).strip().lower() == 'material':
                    columna_material = col
                    break
            
            if columna_material is not None:
                # Contar filas antes de eliminar
                filas_antes_material = len(df)
                # Convertir a string y limpiar espacios
                df[columna_material] = df[columna_material].astype(str).str.strip()
                # Eliminar filas donde Material esté vacía (nan, None, o string vacío)
                df = df[
                    (df[columna_material] != '') & 
                    (df[columna_material] != 'nan') & 
                    (df[columna_material].notna())
                ]
                filas_eliminadas_material = filas_antes_material - len(df)
            else:
                filas_eliminadas_material = 0
            
            # Eliminar filas donde la columna "Invoice Value" sea "0", "0.0" o "0.00"
            # Buscar la columna "Invoice Value" (case-insensitive)
            columna_invoice = None
            for col in df.columns:
                if str(col).strip().lower() == 'invoice value':
                    columna_invoice = col
                    break
            
            if columna_invoice is not None:
                # Convertir a string y limpiar espacios para comparación
                valores_invoice = df[columna_invoice].astype(str).str.strip()
                # Eliminar filas donde Invoice Value sea "0", "0.0" o "0.00"
                df = df[~valores_invoice.isin(['0', '0.0', '0.00'])]
            
            # Crear columna "precio unitario" = Invoice Value / Quantity in OPUn
            col_invoice = None
            col_quantity = None
            for col in df.columns:
                c = str(col).strip().lower()
                if c == 'invoice value':
                    col_invoice = col
                elif c == 'quantity in opun':
                    col_quantity = col
            
            if col_invoice is not None and col_quantity is not None:
                invoice_num = pd.to_numeric(df[col_invoice], errors='coerce')
                quantity_num = pd.to_numeric(df[col_quantity], errors='coerce')
                df['precio unitario'] = invoice_num / quantity_num
                # Evitar división por cero: donde quantity sea 0 o NaN, precio unitario = NaN
                df.loc[(quantity_num <= 0) | quantity_num.isna(), 'precio unitario'] = pd.NA
                df['precio unitario'] = df['precio unitario'].replace([float('inf'), float('-inf')], pd.NA)
            else:
                df['precio unitario'] = pd.NA
            
            # Guardar el archivo procesado
            nombre_archivo_salida = "historial_compras_procesado.xlsx"
            ruta_archivo_salida = carpeta_salida_path / nombre_archivo_salida
            
            # Guardar como Excel
            try:
                df.to_excel(ruta_archivo_salida, index=False, engine='openpyxl')
            except Exception as e:
                return JSONResponse(
                    status_code=500,
                    content={"error": f"Error al guardar el archivo procesado: {str(e)}"}
                )
            
            return JSONResponse(
                status_code=200,
                content={
                    "success": True,
                    "message": f"Archivo procesado exitosamente. Se eliminaron {len(columnas_encontradas)} columnas.",
                    "archivo_guardado": str(ruta_archivo_salida),
                    "columnas_eliminadas": columnas_encontradas,
                    "total_columnas_eliminadas": len(columnas_encontradas)
                }
            )
        
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Error al procesar el historial: {str(e)}"}
        )


async def _procesar_ventas_core(db, df):
    import pandas as pd
    # Valores a buscar en el segundo renglón (índice 1)
    valores_a_buscar = ['OE', 'Budget', 'Act-Budget', 'incl. srap', '2005246']
    
    # Obtener el segundo renglón (índice 1)
    segundo_renglon = df.iloc[1]
    
    # Identificar columnas a eliminar
    columnas_a_eliminar = []
    for idx, valor in enumerate(segundo_renglon):
        # Convertir a string y limpiar espacios
        valor_str = str(valor).strip() if pd.notna(valor) else ''
        # Comparar con los valores a buscar (case-insensitive)
        for valor_buscar in valores_a_buscar:
            if valor_str.lower() == valor_buscar.lower():
                # Usar el índice de la columna
                columnas_a_eliminar.append(idx)
                break
    
    # Eliminar las columnas encontradas
    if columnas_a_eliminar:
        # Eliminar columnas por índice
        df = df.drop(df.columns[columnas_a_eliminar], axis=1)
    
    # Buscar columnas donde el segundo renglón (índice 1) tenga "Actual"
    # Obtener el segundo renglón actualizado después de eliminar columnas
    segundo_renglon_actualizado = df.iloc[1]
    columnas_actual = []
    
    for idx, valor in enumerate(segundo_renglon_actualizado):
        # Convertir a string y limpiar espacios
        valor_str = str(valor).strip() if pd.notna(valor) else ''
        # Comparar con "Actual" (case-insensitive)
        if valor_str.lower() == 'actual':
            columnas_actual.append(idx)
    
    # Copiar el valor del renglón 1 (índice 0) al renglón 3 (índice 2) para las columnas "Actual"
    columnas_modificadas = 0
    for col_idx in columnas_actual:
        # Obtener el valor del renglón 1 (índice 0)
        valor_renglon_1 = df.iloc[0, col_idx]
        # Copiar al renglón 3 (índice 2)
        df.iloc[2, col_idx] = valor_renglon_1
        columnas_modificadas += 1
    
    # Eliminar los renglones 1 y 2 (índices 0 y 1)
    df = df.drop(df.index[0:2])
    # Resetear el índice para que quede secuencial
    df = df.reset_index(drop=True)
    
    # Guardar los valores del renglón 1 (índice 0) para usarlos como encabezado
    valores_renglon_1 = df.iloc[0].tolist()
    
    # Verificar si hay un renglón 2 y si tiene los mismos valores que el renglón 1
    if len(df) > 1:
        valores_renglon_2 = df.iloc[1].tolist()
        # Comparar los valores (considerando NaN como iguales)
        son_iguales = True
        for i in range(len(valores_renglon_1)):
            val1 = valores_renglon_1[i]
            val2 = valores_renglon_2[i]
            # Comparar considerando NaN
            if pd.isna(val1) and pd.isna(val2):
                continue
            elif pd.isna(val1) or pd.isna(val2):
                son_iguales = False
                break
            elif str(val1).strip() != str(val2).strip():
                son_iguales = False
                break
        
        # Si son iguales, eliminar el renglón 2
        if son_iguales:
            df = df.drop(df.index[1])
            df = df.reset_index(drop=True)
    
    # Eliminar el renglón 1 (que ahora usaremos como encabezado)
    df = df.drop(df.index[0])
    df = df.reset_index(drop=True)
    
    # Buscar las columnas necesarias en los encabezados
    # Buscar "Customer"
    columna_customer = None
    # Buscar "Quantity OE/TO FT"
    columna_quantity_ft = None
    # Buscar "Quantity OE/TO M"
    columna_quantity_m = None
    # Buscar "Turnover w/o metal"
    columna_turnover_wo_metal_original = None
    # Buscar "OE/Turnover like FI"
    columna_turnover_like_fi_original = None
    
    for idx, encabezado in enumerate(valores_renglon_1):
        encabezado_str = str(encabezado).strip() if pd.notna(encabezado) else ''
        encabezado_lower = encabezado_str.lower()
        # Buscar "Customer"
        if columna_customer is None and 'customer' in encabezado_lower:
            columna_customer = idx
        # Buscar "Quantity OE/TO FT"
        if columna_quantity_ft is None and ('quantity' in encabezado_lower and 'ft' in encabezado_lower and 
            ('oe' in encabezado_lower or 'to' in encabezado_lower)):
            columna_quantity_ft = idx
        # Buscar "Quantity OE/TO M" (sin FT, solo M)
        if columna_quantity_m is None and ('quantity' in encabezado_lower and 'm' in encabezado_lower and 
            ('oe' in encabezado_lower or 'to' in encabezado_lower) and 'ft' not in encabezado_lower):
            columna_quantity_m = idx
        # Buscar "Turnover w/o metal"
        if columna_turnover_wo_metal_original is None and ('turnover' in encabezado_lower and 'metal' in encabezado_lower and 'w/o' in encabezado_lower):
            columna_turnover_wo_metal_original = idx
        # Buscar "OE/Turnover like FI"
        if columna_turnover_like_fi_original is None and ('turnover' in encabezado_lower and 'like' in encabezado_lower and 'fi' in encabezado_lower):
            columna_turnover_like_fi_original = idx
    
    # Calcular el índice actual de "Turnover w/o metal" después de eliminar columnas
    columna_turnover_wo_metal = None
    if columna_turnover_wo_metal_original is not None:
        # Crear un mapeo de índices originales a índices actuales
        # Las columnas eliminadas están en columnas_a_eliminar
        # Necesitamos contar cuántas columnas antes de la columna objetivo fueron eliminadas
        columnas_eliminadas_antes = sum(1 for col_idx in columnas_a_eliminar if col_idx < columna_turnover_wo_metal_original)
        # El índice actual es el índice original menos las columnas eliminadas antes
        columna_turnover_wo_metal = columna_turnover_wo_metal_original - columnas_eliminadas_antes
        # Verificar que el índice esté dentro del rango del DataFrame
        if columna_turnover_wo_metal < 0 or columna_turnover_wo_metal >= len(df.columns):
            columna_turnover_wo_metal = None
    
    # Calcular el índice actual de "OE/Turnover like FI" después de eliminar columnas
    columna_turnover_like_fi = None
    if columna_turnover_like_fi_original is not None:
        # Crear un mapeo de índices originales a índices actuales
        # Las columnas eliminadas están en columnas_a_eliminar
        # Necesitamos contar cuántas columnas antes de la columna objetivo fueron eliminadas
        columnas_eliminadas_antes = sum(1 for col_idx in columnas_a_eliminar if col_idx < columna_turnover_like_fi_original)
        # El índice actual es el índice original menos las columnas eliminadas antes
        columna_turnover_like_fi = columna_turnover_like_fi_original - columnas_eliminadas_antes
        # Verificar que el índice esté dentro del rango del DataFrame
        if columna_turnover_like_fi < 0 or columna_turnover_like_fi >= len(df.columns):
            columna_turnover_like_fi = None
    
    # Obtener todos los grupos de clientes de la base de datos
    grupos_clientes = await crud.list_cliente_grupos(db, limit=100000)
    # Crear un diccionario que mapee codigo_cliente -> grupo
    dict_grupos = {}
    for grupo_cliente in grupos_clientes:
        if grupo_cliente.codigo_cliente is not None:
            dict_grupos[grupo_cliente.codigo_cliente] = grupo_cliente.grupo or ''
    
    # Agregar la columna "Customer number" a la derecha de "Customer"
    if columna_customer is not None:
        # Obtener los valores de la columna Customer
        valores_customer = df.iloc[:, columna_customer]
        # Extraer los primeros 7 caracteres de cada valor
        customer_number = valores_customer.astype(str).str[:7]
        # Reemplazar 'nan' con string vacío
        customer_number = customer_number.replace('nan', '')
        # Insertar la columna justo después de Customer (posición columna_customer + 1)
        df.insert(columna_customer + 1, 'Customer number', customer_number.tolist())
        # Actualizar los índices de las otras columnas si Customer estaba antes de ellas
        if columna_quantity_ft is not None and columna_quantity_ft > columna_customer:
            columna_quantity_ft += 1
        if columna_quantity_m is not None and columna_quantity_m > columna_customer:
            columna_quantity_m += 1
        if columna_turnover_wo_metal is not None and columna_turnover_wo_metal > columna_customer:
            columna_turnover_wo_metal += 1
        if columna_turnover_like_fi is not None and columna_turnover_like_fi > columna_customer:
            columna_turnover_like_fi += 1
        # Actualizar valores_renglon_1 para incluir el nuevo encabezado
        valores_renglon_1.insert(columna_customer + 1, 'Customer number')
        
        # Agregar la columna "grupo" a la derecha de "Customer number"
        # Obtener los valores de la columna Customer number (ahora en posición columna_customer + 1)
        valores_customer_number = df.iloc[:, columna_customer + 1]
        # Convertir a entero para buscar en el diccionario, manejando errores
        grupos = []
        for valor in valores_customer_number:
            try:
                # Convertir a string y limpiar
                valor_str = str(valor).strip()
                # Verificar que no esté vacío ni sea 'nan'
                if not valor_str or valor_str.lower() == 'nan' or valor_str == '':
                    grupos.append('')
                    continue
                # Intentar convertir a entero
                codigo = int(float(valor_str))
                # Buscar el grupo en el diccionario
                if codigo in dict_grupos:
                    grupos.append(dict_grupos[codigo])
                else:
                    grupos.append('')
            except (ValueError, TypeError):
                grupos.append('')
        
        # Insertar la columna "grupo" justo después de Customer number (posición columna_customer + 2)
        df.insert(columna_customer + 2, 'grupo', grupos)
        # Actualizar los índices de las otras columnas si Customer estaba antes de ellas
        if columna_quantity_ft is not None and columna_quantity_ft > columna_customer:
            columna_quantity_ft += 1
        if columna_quantity_m is not None and columna_quantity_m > columna_customer:
            columna_quantity_m += 1
        if columna_turnover_wo_metal is not None and columna_turnover_wo_metal > columna_customer:
            columna_turnover_wo_metal += 1
        if columna_turnover_like_fi is not None and columna_turnover_like_fi > columna_customer:
            columna_turnover_like_fi += 1
        # Actualizar valores_renglon_1 para incluir el nuevo encabezado
        valores_renglon_1.insert(columna_customer + 2, 'grupo')
    
    # Agregar las 3 nuevas columnas
    num_filas = len(df)
    
    # Calcular "Conversion de FT a M" = Quantity OE/TO FT / 3.2808
    if columna_quantity_ft is not None:
        # Obtener los valores de la columna Quantity OE/TO FT
        valores_quantity_ft = df.iloc[:, columna_quantity_ft]
        # Convertir a numérico y dividir entre 3.2808
        valores_numericos_ft = pd.to_numeric(valores_quantity_ft, errors='coerce')
        conversion_ft_m = valores_numericos_ft / 3.2808
        df['Conversion de FT a M'] = conversion_ft_m
    else:
        # Si no se encuentra la columna, dejar vacío
        df['Conversion de FT a M'] = [None] * num_filas
    
    # Calcular "Sales total MTS" = Quantity OE/TO M + Conversion de FT a M
    if columna_quantity_m is not None:
        # Obtener los valores de la columna Quantity OE/TO M
        valores_quantity_m = df.iloc[:, columna_quantity_m]
        # Convertir a numérico
        valores_numericos_m = pd.to_numeric(valores_quantity_m, errors='coerce')
        # Obtener los valores de Conversion de FT a M (ya calculados)
        valores_conversion = pd.to_numeric(df['Conversion de FT a M'], errors='coerce')
        # Sumar ambos valores
        sales_total_mts = valores_numericos_m + valores_conversion
        df['Sales total MTS'] = sales_total_mts
    else:
        # Si no se encuentra la columna Quantity OE/TO M, usar solo Conversion de FT a M
        valores_conversion = pd.to_numeric(df['Conversion de FT a M'], errors='coerce')
        df['Sales total MTS'] = valores_conversion
    
    # Calcular "Sales KM" = Sales total MTS / 1000
    valores_sales_total = pd.to_numeric(df['Sales total MTS'], errors='coerce')
    sales_km = valores_sales_total / 1000
    df['Sales KM'] = sales_km
    
    # Buscar las columnas "Turnover w/o metal" y "OE/Turnover like FI" directamente en el DataFrame
    # después de todas las modificaciones, usando valores_renglon_1 que ya tiene las columnas insertadas
    columna_turnover_wo_metal_final = None
    columna_turnover_like_fi_final = None
    
    # Verificar que valores_renglon_1 tenga la misma longitud que las columnas del DataFrame
    num_columnas_df = len(df.columns)
    num_valores_renglon = len(valores_renglon_1)
    
    # Usar el mínimo para evitar errores de índice
    num_columnas_a_buscar = min(num_columnas_df, num_valores_renglon)
    
    for idx in range(num_columnas_a_buscar):
        encabezado = valores_renglon_1[idx] if idx < len(valores_renglon_1) else ''
        encabezado_str = str(encabezado).strip() if pd.notna(encabezado) else ''
        encabezado_lower = encabezado_str.lower()
        
        # Buscar "Turnover w/o metal" - debe tener "turnover", "metal" y "w/o" o "without", pero NO "like" ni "fi"
        # Patrón más específico: buscar "w/o" o "without" junto con "metal" pero sin "like" ni "fi"
        if columna_turnover_wo_metal_final is None:
            has_turnover = 'turnover' in encabezado_lower
            has_metal = 'metal' in encabezado_lower
            has_wo = 'w/o' in encabezado_lower or 'without' in encabezado_lower
            has_like = 'like' in encabezado_lower
            has_fi = 'fi' in encabezado_lower
            
            if has_turnover and has_metal and has_wo and not has_like and not has_fi:
                columna_turnover_wo_metal_final = idx
        
        # Buscar "OE/Turnover like FI" - debe tener "turnover", "like" y "fi"
        # Patrón más específico: buscar "like" y "fi" juntos, con "turnover"
        if columna_turnover_like_fi_final is None:
            has_turnover = 'turnover' in encabezado_lower
            has_like = 'like' in encabezado_lower
            has_fi = 'fi' in encabezado_lower
            
            if has_turnover and has_like and has_fi:
                columna_turnover_like_fi_final = idx
    
    # Calcular "Precio Exmetal KM" = Turnover w/o metal / Sales KM
    if columna_turnover_wo_metal_final is not None and 0 <= columna_turnover_wo_metal_final < len(df.columns):
        # Obtener los valores de la columna Turnover w/o metal
        valores_turnover = df.iloc[:, columna_turnover_wo_metal_final]
        # Convertir a numérico
        valores_numericos_turnover = pd.to_numeric(valores_turnover, errors='coerce')
        # Obtener los valores de Sales KM (ya calculados)
        valores_sales_km = pd.to_numeric(df['Sales KM'], errors='coerce')
        # Dividir: Precio Exmetal KM = Turnover w/o metal / Sales KM
        # Evitar división por cero
        precio_exmetal = valores_numericos_turnover / valores_sales_km.replace(0, pd.NA)
        df['Precio Exmetal KM'] = precio_exmetal
    else:
        # Si no se encuentra la columna, dejar vacío
        df['Precio Exmetal KM'] = [None] * num_filas
    
    # Calcular "Precio Full Metal KM" = OE/Turnover like FI / Sales KM
    if columna_turnover_like_fi_final is not None and 0 <= columna_turnover_like_fi_final < len(df.columns):
        # Verificar que no sea la misma columna que Turnover w/o metal
        if columna_turnover_like_fi_final != columna_turnover_wo_metal_final:
            # Obtener los valores de la columna OE/Turnover like FI
            valores_turnover_fi = df.iloc[:, columna_turnover_like_fi_final]
            # Convertir a numérico
            valores_numericos_turnover_fi = pd.to_numeric(valores_turnover_fi, errors='coerce')
            # Obtener los valores de Sales KM (ya calculados)
            valores_sales_km = pd.to_numeric(df['Sales KM'], errors='coerce')
            # Dividir: Precio Full Metal KM = OE/Turnover like FI / Sales KM
            # Evitar división por cero
            precio_full_metal = valores_numericos_turnover_fi / valores_sales_km.replace(0, pd.NA)
            df['Precio Full Metal KM'] = precio_full_metal
        else:
            # Si es la misma columna, dejar vacío
            df['Precio Full Metal KM'] = [None] * num_filas
    else:
        # Si no se encuentra la columna, dejar vacío
        df['Precio Full Metal KM'] = [None] * num_filas
    
    # Calcular "Precio Exmetal M" = Precio Exmetal KM / 1000
    valores_precio_exmetal_km = pd.to_numeric(df['Precio Exmetal KM'], errors='coerce')
    precio_exmetal_m = valores_precio_exmetal_km / 1000
    df['Precio Exmetal M'] = precio_exmetal_m
    
    # Calcular "Precio Full Metal M" = Precio Full Metal KM / 1000
    valores_precio_full_metal_km = pd.to_numeric(df['Precio Full Metal KM'], errors='coerce')
    precio_full_metal_m = valores_precio_full_metal_km / 1000
    df['Precio Full Metal M'] = precio_full_metal_m
    
    # Reemplazar NaN con string vacío en las columnas calculadas
    df['Conversion de FT a M'] = df['Conversion de FT a M'].fillna('')
    df['Sales total MTS'] = df['Sales total MTS'].fillna('')
    df['Sales KM'] = df['Sales KM'].fillna('')
    df['Precio Exmetal KM'] = df['Precio Exmetal KM'].fillna('')
    df['Precio Full Metal KM'] = df['Precio Full Metal KM'].fillna('')
    df['Precio Exmetal M'] = df['Precio Exmetal M'].fillna('')
    df['Precio Full Metal M'] = df['Precio Full Metal M'].fillna('')
    
    # Crear una fila de encabezados usando:
    # - Los valores del renglón 1 original para las columnas originales
    # - Los nombres de las nuevas columnas para las nuevas columnas
    nombres_encabezados = valores_renglon_1.copy()
    nombres_encabezados.extend(['Conversion de FT a M', 'Sales total MTS', 'Sales KM', 'Precio Exmetal KM', 'Precio Full Metal KM', 'Precio Exmetal M', 'Precio Full Metal M'])
    
    # Convertir los valores a strings para evitar problemas
    nombres_encabezados = [str(val) if pd.notna(val) else '' for val in nombres_encabezados]
    
    # Crear un DataFrame con una sola fila (los encabezados)
    fila_encabezados = pd.DataFrame([nombres_encabezados], columns=df.columns)
    # Concatenar la fila de encabezados al inicio del DataFrame
    df = pd.concat([fila_encabezados, df], ignore_index=True)
    
    # Actualizar y guardar los índices de las columnas según su nombre
    # Crear un diccionario que mapee el nombre de cada columna a su índice actualizado
    indices_columnas = {}
    for idx, nombre_columna in enumerate(df.columns):
        # Obtener el nombre real de la columna desde los encabezados (primera fila)
        nombre_real = str(df.iloc[0, idx]).strip() if pd.notna(df.iloc[0, idx]) else ''
        if nombre_real:
            # Guardar el índice con el nombre de la columna (case-insensitive para búsqueda)
            indices_columnas[nombre_real.lower()] = idx
            # También guardar con el nombre exacto
            indices_columnas[nombre_real] = idx
    
    # Actualizar los índices de las columnas conocidas después de todas las modificaciones
    # Buscar nuevamente las columnas en el DataFrame final para obtener sus índices actualizados
    columna_customer_actualizado = None
    columna_quantity_ft_actualizado = None
    columna_quantity_m_actualizado = None
    columna_customer_number_actualizado = None
    columna_grupo_actualizado = None
    columna_conversion_ft_m_actualizado = None
    columna_sales_total_mts_actualizado = None
    columna_sales_km_actualizado = None
    columna_turnover_wo_metal_actualizado = None
    columna_precio_exmetal_actualizado = None
    columna_turnover_like_fi_actualizado = None
    columna_precio_full_metal_actualizado = None
    columna_precio_exmetal_m_actualizado = None
    columna_precio_full_metal_m_actualizado = None
    
    # Buscar en la primera fila (encabezados) del DataFrame final
    primera_fila = df.iloc[0]
    for idx, encabezado in enumerate(primera_fila):
        encabezado_str = str(encabezado).strip() if pd.notna(encabezado) else ''
        encabezado_lower = encabezado_str.lower()
        
        # Buscar cada columna por nombre
        if columna_customer_actualizado is None and 'customer' in encabezado_lower and 'number' not in encabezado_lower:
            columna_customer_actualizado = idx
        if columna_customer_number_actualizado is None and 'customer number' in encabezado_lower:
            columna_customer_number_actualizado = idx
        if columna_grupo_actualizado is None and encabezado_lower == 'grupo':
            columna_grupo_actualizado = idx
        if columna_quantity_ft_actualizado is None and ('quantity' in encabezado_lower and 'ft' in encabezado_lower and 
            ('oe' in encabezado_lower or 'to' in encabezado_lower)):
            columna_quantity_ft_actualizado = idx
        if columna_quantity_m_actualizado is None and ('quantity' in encabezado_lower and 'm' in encabezado_lower and 
            ('oe' in encabezado_lower or 'to' in encabezado_lower) and 'ft' not in encabezado_lower):
            columna_quantity_m_actualizado = idx
        if columna_conversion_ft_m_actualizado is None and 'conversion de ft a m' in encabezado_lower:
            columna_conversion_ft_m_actualizado = idx
        if columna_sales_total_mts_actualizado is None and 'sales total mts' in encabezado_lower:
            columna_sales_total_mts_actualizado = idx
        if columna_sales_km_actualizado is None and 'sales km' in encabezado_lower:
            columna_sales_km_actualizado = idx
        if columna_turnover_wo_metal_actualizado is None and ('turnover' in encabezado_lower and 'metal' in encabezado_lower and 'w/o' in encabezado_lower):
            columna_turnover_wo_metal_actualizado = idx
        if columna_precio_exmetal_actualizado is None and 'precio exmetal km' in encabezado_lower:
            columna_precio_exmetal_actualizado = idx
        if columna_turnover_like_fi_actualizado is None and ('turnover' in encabezado_lower and 'like' in encabezado_lower and 'fi' in encabezado_lower):
            columna_turnover_like_fi_actualizado = idx
        if columna_precio_full_metal_actualizado is None and 'precio full metal km' in encabezado_lower:
            columna_precio_full_metal_actualizado = idx
        if columna_precio_exmetal_m_actualizado is None and 'precio exmetal m' in encabezado_lower:
            columna_precio_exmetal_m_actualizado = idx
        if columna_precio_full_metal_m_actualizado is None and 'precio full metal m' in encabezado_lower:
            columna_precio_full_metal_m_actualizado = idx
    
    # Crear diccionario con los índices actualizados de las columnas principales
    indices_columnas_principales = {
        'customer': columna_customer_actualizado,
        'customer_number': columna_customer_number_actualizado,
        'grupo': columna_grupo_actualizado,
        'quantity_ft': columna_quantity_ft_actualizado,
        'quantity_m': columna_quantity_m_actualizado,
        'conversion_ft_m': columna_conversion_ft_m_actualizado,
        'sales_total_mts': columna_sales_total_mts_actualizado,
        'sales_km': columna_sales_km_actualizado,
        'turnover_wo_metal': columna_turnover_wo_metal_actualizado,
        'precio_exmetal': columna_precio_exmetal_actualizado,
        'turnover_like_fi': columna_turnover_like_fi_actualizado,
        'precio_full_metal': columna_precio_full_metal_actualizado,
        'precio_exmetal_m': columna_precio_exmetal_m_actualizado,
        'precio_full_metal_m': columna_precio_full_metal_m_actualizado
    }
    
    # Modificar la columna Customer al final: eliminar los primeros 7 caracteres y espacios en blanco
    # Buscar la columna Customer en el DataFrame final
    columna_customer_final = None
    primera_fila_final = df.iloc[0]
    for idx, encabezado in enumerate(primera_fila_final):
        encabezado_str = str(encabezado).strip() if pd.notna(encabezado) else ''
        encabezado_lower = encabezado_str.lower()
        if columna_customer_final is None and 'customer' in encabezado_lower and 'number' not in encabezado_lower:
            columna_customer_final = idx
            break
    
    if columna_customer_final is not None:
        # Obtener todos los valores de la columna Customer (excluyendo la primera fila que es el encabezado)
        # Modificar desde la fila 1 en adelante (índice 1)
        for idx_fila in range(1, len(df)):
            valor_original = df.iloc[idx_fila, columna_customer_final]
            if pd.notna(valor_original):
                # Convertir a string
                valor_str = str(valor_original)
                # Eliminar los primeros 7 caracteres
                if len(valor_str) > 7:
                    valor_modificado = valor_str[7:]
                else:
                    valor_modificado = ''
                # Eliminar espacios en blanco al inicio y al final
                valor_modificado = valor_modificado.strip()
                # Actualizar el valor en el DataFrame
                df.iloc[idx_fila, columna_customer_final] = valor_modificado
    
    # Descomponer la columna "Product" en "Producto" y "Product description"
    # Buscar la columna Product en el DataFrame final
    columna_product_final = None
    primera_fila_final = df.iloc[0]
    for idx, encabezado in enumerate(primera_fila_final):
        encabezado_str = str(encabezado).strip() if pd.notna(encabezado) else ''
        encabezado_lower = encabezado_str.lower()
        if columna_product_final is None and encabezado_lower == 'product':
            columna_product_final = idx
            break
    
    if columna_product_final is not None:
        # Crear listas para almacenar los valores descompuestos
        valores_producto = []
        valores_product_description = []
        
        # Procesar cada fila (incluyendo el encabezado)
        for idx_fila in range(len(df)):
            valor_original = df.iloc[idx_fila, columna_product_final]
            if pd.notna(valor_original):
                # Convertir a string
                valor_str = str(valor_original).strip()
                # Dividir por el primer espacio
                partes = valor_str.split(' ', 1)  # split con maxsplit=1 divide solo en el primer espacio
                if len(partes) == 2:
                    # Hay dos partes: antes y después del primer espacio
                    valores_producto.append(partes[0])
                    # Eliminar espacios en blanco al inicio y al final de la descripción
                    descripcion_limpia = partes[1].strip()
                    valores_product_description.append(descripcion_limpia)
                elif len(partes) == 1:
                    # Solo hay una parte (no hay espacio)
                    valores_producto.append(partes[0])
                    valores_product_description.append('')
                else:
                    # Caso vacío
                    valores_producto.append('')
                    valores_product_description.append('')
            else:
                # Valor NaN
                valores_producto.append('')
                valores_product_description.append('')
        
        # Actualizar la columna Product con los valores de Producto
        df.iloc[:, columna_product_final] = valores_producto
        
        # Cambiar el nombre del encabezado de "Product" a "Producto" en la primera fila
        df.iloc[0, columna_product_final] = 'Producto'
        
        # Insertar la nueva columna "Product Description" justo después de Producto
        # Necesitamos insertar como una serie con el mismo índice que el DataFrame
        df.insert(columna_product_final + 1, 'Product Description', valores_product_description)
        
        # Asegurarse de que el encabezado en la primera fila sea "Product Description"
        df.iloc[0, columna_product_final + 1] = 'Product Description'
    
    # Mapear columnas del Excel a columnas de la BD
    # Primero, obtener los índices de todas las columnas necesarias desde los encabezados
    primera_fila = df.iloc[0]
    columnas_map = {}
    
    # Debug: imprimir todos los encabezados para ver qué hay disponible
    # print("Encabezados encontrados:", [str(h).strip() for h in primera_fila])
    
    # Buscar todas las columnas necesarias
    for idx, encabezado in enumerate(primera_fila):
        encabezado_str = str(encabezado).strip() if pd.notna(encabezado) else ''
        encabezado_lower = encabezado_str.lower()
        
        # Mapear según los nombres proporcionados
        if 'customer' in encabezado_lower and 'number' not in encabezado_lower:
            columnas_map['Customer'] = idx
        elif 'customer number' in encabezado_lower:
            columnas_map['Customer Number'] = idx
        elif encabezado_lower == 'grupo':
            columnas_map['grupo'] = idx
        elif 'business unit' in encabezado_lower:
            columnas_map['Business Unit'] = idx
        elif 'period' in encabezado_lower:
            if 'Period/Year' not in columnas_map:  # Solo tomar la primera coincidencia
                columnas_map['Period/Year'] = idx
        elif 'artnr condensed' in encabezado_lower:
            columnas_map['Artnr condensed'] = idx
        elif 'region asc' in encabezado_lower:
            columnas_map['Region ASC'] = idx
        elif encabezado_lower == 'plant':
            columnas_map['Plant'] = idx
        elif 'ship' in encabezado_lower and 'party' in encabezado_lower:
            if 'Shipt-to-party' not in columnas_map:  # Solo tomar la primera coincidencia
                columnas_map['Shipt-to-party'] = idx
        elif encabezado_lower == 'producto':
            columnas_map['Producto'] = idx
        elif 'product description' in encabezado_lower:
            columnas_map['Product Description'] = idx
        elif 'turnover' in encabezado_lower and 'metal' in encabezado_lower and 'w/o' in encabezado_lower:
            columnas_map['Turnover w/o metal'] = idx
        elif 'turnover' in encabezado_lower and 'like' in encabezado_lower and 'fi' in encabezado_lower:
            columnas_map['OE/Turnover like FI'] = idx
        elif 'copper sales' in encabezado_lower and 'cuv' in encabezado_lower:
            columnas_map['Copper Sales (CUV)'] = idx
        elif 'cu-sales effect' in encabezado_lower:
            columnas_map['CU-Sales effect'] = idx
        elif 'cu result' in encabezado_lower:
            columnas_map['CU result'] = idx
        elif 'quantity' in encabezado_lower and 'm' in encabezado_lower and 'ft' not in encabezado_lower:
            columnas_map['Quantity OE/TO M'] = idx
        elif 'quantity' in encabezado_lower and 'ft' in encabezado_lower:
            columnas_map['Quantity OE/TO FT'] = idx
        elif 'cu weight techn' in encabezado_lower and 'cut' in encabezado_lower:
            columnas_map['CU Weight techn. CUT'] = idx
        elif 'cu weight sales' in encabezado_lower and 'cuv' in encabezado_lower:
            columnas_map['CU weight Sales  CUV'] = idx
        elif 'conversion de ft a m' in encabezado_lower:
            columnas_map['Conversion de FT a M'] = idx
        elif 'sales total mts' in encabezado_lower:
            columnas_map['Sales total MTS'] = idx
        elif 'sales km' in encabezado_lower:
            columnas_map['Sales KM'] = idx
        elif 'precio exmetal km' in encabezado_lower:
            columnas_map['Precio Exmetal KM'] = idx
        elif 'precio full metal km' in encabezado_lower:
            columnas_map['Precio Full Metal KM'] = idx
        elif 'precio exmetal m' in encabezado_lower:
            columnas_map['Precio Exmetal M'] = idx
        elif 'precio full metal m' in encabezado_lower:
            columnas_map['Precio Full Metal M'] = idx
    
    # Debug: verificar qué columnas se encontraron
    # print(f"Columnas encontradas: {list(columnas_map.keys())}")
    # if 'Period/Year' not in columnas_map:
    #     print("ADVERTENCIA: Columna Period/Year no encontrada")
    # if 'Shipt-to-party' not in columnas_map:
    #     print("ADVERTENCIA: Columna Shipt-to-party no encontrada")
    
    # Obtener todos los grupos de clientes para mapear grupo_id
    grupos_clientes = await crud.list_cliente_grupos(db, limit=100000)
    dict_grupos_cliente = {}
    dict_grupos_id = {}
    for grupo_cliente in grupos_clientes:
        if grupo_cliente.codigo_cliente is not None:
            dict_grupos_cliente[grupo_cliente.codigo_cliente] = grupo_cliente.grupo or ''
            # Crear un diccionario para buscar grupo_id por codigo_cliente y grupo
            key = (grupo_cliente.codigo_cliente, grupo_cliente.grupo or '')
            dict_grupos_id[key] = grupo_cliente.id
    
    # Procesar cada fila del DataFrame (excluyendo el encabezado)
    ventas_data = []
    from datetime import date as date_class
    from decimal import Decimal
    
    for idx_fila in range(1, len(df)):
        venta_dict = {}
        # Inicializar valores por defecto
        venta_dict['periodo'] = None
        venta_dict['ship_to_party'] = None
        
        # Cliente - Customer
        if 'Customer' in columnas_map:
            cliente_val = df.iloc[idx_fila, columnas_map['Customer']]
            venta_dict['cliente'] = str(cliente_val).strip() if pd.notna(cliente_val) else None
        
        # Codigo_cliente - Customer Number
        if 'Customer Number' in columnas_map:
            codigo_val = df.iloc[idx_fila, columnas_map['Customer Number']]
            try:
                if pd.notna(codigo_val):
                    codigo_str = str(codigo_val).strip()
                    if codigo_str and codigo_str.lower() != 'nan':
                        venta_dict['codigo_cliente'] = int(float(codigo_str))
                    else:
                        venta_dict['codigo_cliente'] = None
                else:
                    venta_dict['codigo_cliente'] = None
            except (ValueError, TypeError):
                venta_dict['codigo_cliente'] = None
        
        # Grupo - buscar grupo_id basado en codigo_cliente y grupo
        grupo_id = None
        if 'grupo' in columnas_map and venta_dict.get('codigo_cliente'):
            grupo_val = df.iloc[idx_fila, columnas_map['grupo']]
            grupo_str = str(grupo_val).strip() if pd.notna(grupo_val) else ''
            if grupo_str and grupo_str.lower() != 'nan':
                key = (venta_dict['codigo_cliente'], grupo_str)
                grupo_id = dict_grupos_id.get(key)
        venta_dict['grupo_id'] = grupo_id
        
        # Unidad_negocio - Business Unit
        if 'Business Unit' in columnas_map:
            unidad_val = df.iloc[idx_fila, columnas_map['Business Unit']]
            venta_dict['unidad_negocio'] = str(unidad_val).strip() if pd.notna(unidad_val) else None
        
        # Periodo - Period/Year (convertir a Date)
        if 'Period/Year' in columnas_map:
            periodo_val = df.iloc[idx_fila, columnas_map['Period/Year']]
            venta_dict['periodo'] = None  # Valor por defecto
            if pd.notna(periodo_val):
                periodo_str = str(periodo_val).strip()
                if periodo_str and periodo_str.lower() not in ['nan', 'none', '']:
                    try:
                        import re
                        
                        # Buscar año de 4 dígitos (2020-2099)
                        año_match = re.search(r'\b(20\d{2})\b', periodo_str)
                        año = None
                        if año_match:
                            año = int(año_match.group(1))
                        
                        # Buscar mes (1-12) - puede estar antes o después del año
                        # Buscar números de 1-2 dígitos que representen meses válidos
                        mes_match = None
                        
                        # Patrón 1: Buscar "Period YYYY" seguido de un número que podría ser el mes
                        # Ejemplo: "Period 2025 4" o "4. Period 2025"
                        period_pattern = re.search(r'period\s+(\d{4})\s*[.\s]*(\d{1,2})', periodo_str, re.IGNORECASE)
                        if period_pattern:
                            año = int(period_pattern.group(1))
                            mes_candidato = int(period_pattern.group(2))
                            if 1 <= mes_candidato <= 12:
                                mes_match = mes_candidato
                        
                        # Patrón 2: Buscar número seguido de punto y año, luego otro número
                        # Ejemplo: "004.2025 4" -> mes=4, año=2025
                        dot_pattern = re.search(r'(\d{1,2})\.\s*(\d{4})\s+(\d{1,2})', periodo_str)
                        if dot_pattern and not mes_match:
                            año_candidato = int(dot_pattern.group(2))
                            mes_candidato = int(dot_pattern.group(3))
                            if 1 <= mes_candidato <= 12:
                                año = año_candidato
                                mes_match = mes_candidato
                        
                        # Patrón 3: Buscar número antes de "Period YYYY"
                        # Ejemplo: "4. Period 2025"
                        before_period = re.search(r'(\d{1,2})\s*\.\s*period\s+(\d{4})', periodo_str, re.IGNORECASE)
                        if before_period and not mes_match:
                            mes_candidato = int(before_period.group(1))
                            año_candidato = int(before_period.group(2))
                            if 1 <= mes_candidato <= 12:
                                año = año_candidato
                                mes_match = mes_candidato
                        
                        # Patrón 4: Buscar formato estándar MM/YYYY o YYYY-MM
                        if '/' in periodo_str:
                            partes = periodo_str.split('/')
                            if len(partes) == 2:
                                try:
                                    mes_candidato = int(partes[0])
                                    año_candidato = int(partes[1])
                                    if año_candidato < 100:
                                        año_candidato += 2000
                                    if 1 <= mes_candidato <= 12 and 2000 <= año_candidato <= 2099:
                                        año = año_candidato
                                        mes_match = mes_candidato
                                except (ValueError, TypeError):
                                    pass
                        
                        if '-' in periodo_str and not mes_match:
                            partes = periodo_str.split('-')
                            if len(partes) == 2:
                                try:
                                    año_candidato = int(partes[0])
                                    mes_candidato = int(partes[1])
                                    if año_candidato < 100:
                                        año_candidato += 2000
                                    if 1 <= mes_candidato <= 12 and 2000 <= año_candidato <= 2099:
                                        año = año_candidato
                                        mes_match = mes_candidato
                                except (ValueError, TypeError):
                                    pass
                        
                        # Patrón 5: Buscar todos los números y encontrar mes y año
                        if not mes_match and año:
                            numeros = re.findall(r'\d+', periodo_str)
                            for num_str in numeros:
                                num = int(num_str)
                                # Si es un número de 1-12, podría ser el mes
                                if 1 <= num <= 12 and not mes_match:
                                    mes_match = num
                                # Si es un número de 4 dígitos entre 2000-2099, es el año
                                elif 2000 <= num <= 2099 and not año:
                                    año = num
                        
                        # Patrón 6: Formato YYYYMM o MMYYYY (6 dígitos)
                        if not mes_match and len(periodo_str.replace(' ', '').replace('.', '').replace('-', '')) == 6:
                            periodo_limpio = re.sub(r'[^\d]', '', periodo_str)
                            if len(periodo_limpio) == 6:
                                try:
                                    # Intentar YYYYMM primero
                                    año_candidato = int(periodo_limpio[:4])
                                    mes_candidato = int(periodo_limpio[4:])
                                    if 1 <= mes_candidato <= 12 and 2000 <= año_candidato <= 2099:
                                        año = año_candidato
                                        mes_match = mes_candidato
                                except (ValueError, TypeError):
                                    try:
                                        # Intentar MMYYYY
                                        mes_candidato = int(periodo_limpio[:2])
                                        año_candidato = int(periodo_limpio[2:])
                                        if año_candidato < 100:
                                            año_candidato += 2000
                                        if 1 <= mes_candidato <= 12 and 2000 <= año_candidato <= 2099:
                                            año = año_candidato
                                            mes_match = mes_candidato
                                    except (ValueError, TypeError):
                                        pass
                        
                        # Patrón 7: Intentar parsear como fecha completa usando pandas
                        if not mes_match or not año:
                            try:
                                fecha = pd.to_datetime(periodo_str, errors='coerce')
                                if pd.notna(fecha):
                                    año = fecha.year
                                    mes_match = fecha.month
                            except:
                                pass
                        
                        # Si tenemos mes y año válidos, crear la fecha
                        if año and mes_match and 1 <= mes_match <= 12 and 2000 <= año <= 2099:
                            venta_dict['periodo'] = date_class(año, mes_match, 1)
                            
                    except Exception as e:
                        # Si hay algún error, dejar como None
                        pass
        else:
            venta_dict['periodo'] = None
        
        # Producto_condensado - Artnr condensed
        if 'Artnr condensed' in columnas_map:
            producto_cond_val = df.iloc[idx_fila, columnas_map['Artnr condensed']]
            venta_dict['producto_condensado'] = str(producto_cond_val).strip() if pd.notna(producto_cond_val) else None
        
        # Region_asc - Region ASC
        if 'Region ASC' in columnas_map:
            region_val = df.iloc[idx_fila, columnas_map['Region ASC']]
            venta_dict['region_asc'] = str(region_val).strip() if pd.notna(region_val) else None
        
        # Planta - Plant
        if 'Plant' in columnas_map:
            planta_val = df.iloc[idx_fila, columnas_map['Plant']]
            venta_dict['planta'] = str(planta_val).strip() if pd.notna(planta_val) else None
        
        # ship_to_party - Shipt-to-party
        if 'Shipt-to-party' in columnas_map:
            ship_val = df.iloc[idx_fila, columnas_map['Shipt-to-party']]
            if pd.notna(ship_val):
                ship_str = str(ship_val).strip()
                venta_dict['ship_to_party'] = ship_str if ship_str and ship_str.lower() != 'nan' else None
            else:
                venta_dict['ship_to_party'] = None
        else:
            venta_dict['ship_to_party'] = None
        
        # Producto - Producto
        if 'Producto' in columnas_map:
            producto_val = df.iloc[idx_fila, columnas_map['Producto']]
            venta_dict['producto'] = str(producto_val).strip() if pd.notna(producto_val) else None
        
        # descripcion_producto - Product Description
        if 'Product Description' in columnas_map:
            desc_val = df.iloc[idx_fila, columnas_map['Product Description']]
            venta_dict['descripcion_producto'] = str(desc_val).strip() if pd.notna(desc_val) else None
        
        # Campos numéricos - convertir a Decimal o None
        campos_numericos = {
            'Turnover w/o metal': 'turnover_wo_metal',
            'OE/Turnover like FI': 'oe_turnover_like_fi',
            'Copper Sales (CUV)': 'copper_sales_cuv',
            'CU-Sales effect': 'cu_sales_effect',
            'CU result': 'cu_result',
            'Quantity OE/TO M': 'quantity_oe_to_m',
            'Quantity OE/TO FT': 'quantity_oe_to_ft',
            'CU Weight techn. CUT': 'cu_weight_techn_cut',
            'CU weight Sales  CUV': 'cu_weight_sales_cuv',
            'Conversion de FT a M': 'conversion_ft_a_m',
            'Sales total MTS': 'sales_total_mts',
            'Sales KM': 'sales_km',
            'Precio Exmetal KM': 'precio_exmetal_km',
            'Precio Full Metal KM': 'precio_full_metal_km',
            'Precio Exmetal M': 'precio_exmetal_m',
            'Precio Full Metal M': 'precio_full_metal_m'
        }
        
        for col_excel, col_bd in campos_numericos.items():
            if col_excel in columnas_map:
                val = df.iloc[idx_fila, columnas_map[col_excel]]
                if pd.notna(val):
                    try:
                        val_str = str(val).strip()
                        if val_str and val_str.lower() != 'nan' and val_str != '':
                            venta_dict[col_bd] = Decimal(str(val))
                        else:
                            venta_dict[col_bd] = None
                    except (ValueError, TypeError):
                        venta_dict[col_bd] = None
                else:
                    venta_dict[col_bd] = None
            else:
                venta_dict[col_bd] = None
        
        ventas_data.append(venta_dict)
    
    # Insertar los datos en la base de datos (verificando duplicados)
    resultado = await crud.bulk_create_ventas(db, ventas_data)
    registros_insertados = resultado.get("insertados", 0)
    registros_duplicados = resultado.get("duplicados", 0)
    errores = resultado.get("errores", [])
    
    mensaje = f"Archivo procesado exitosamente. {registros_insertados} registros insertados."
    if registros_duplicados > 0:
        mensaje += f" {registros_duplicados} registros duplicados fueron omitidos (ya existían con la misma combinación de código cliente, producto y período)."
    if errores:
        mensaje += f" Se encontraron {len(errores)} error(es) al procesar algunos registros."
    
    # Determinar si fue exitoso (al menos algunos registros se insertaron o había duplicados)
    success = registros_insertados > 0 or (registros_duplicados > 0 and len(errores) == 0)
    return {
        "success": success,
        "message": mensaje,
        "registros_insertados": registros_insertados,
        "registros_duplicados": registros_duplicados,
        "total_registros": len(ventas_data),
        "errores": errores
    }


async def _procesar_archivo_ventas_desde_ruta(
    db,
    path_ventas: "Path",
    execution_id: Optional[int] = None,
    user_id: Optional[int] = None,
):
    """Lee el Excel desde path_ventas, procesa con _procesar_ventas_core y actualiza la ejecución.
    Si execution_id es None, crea una nueva ejecución (requiere user_id).
    Devuelve el dict con success, message, registros_insertados, etc. En error de validación/lectura actualiza ejecución y devuelve dict con success=False."""
    from datetime import timezone as tz
    import pandas as pd
    if execution_id is None and user_id is not None:
        ahora_utc = datetime.now(tz.utc)
        execution_id = await crud.create_sales_execution(
            db=db,
            user_id=user_id,
            fecha_inicio_periodo=ahora_utc,
            fecha_fin_periodo=ahora_utc,
            sistema_sap="Carga manual",
            transaccion="Procesar archivo",
            maquina=socket.gethostname() if hasattr(socket, 'gethostname') else (platform.node() or "unknown"),
        )
        await crud.update_sales_execution_status(
            db=db, execution_id=execution_id, estado=ExecutionStatus.RUNNING, archivo_nombre=path_ventas.name
        )
    extension = path_ventas.suffix.lower()
    if extension not in (".xlsx", ".xls"):
        msg_ext = f"Extensión no permitida: {extension}"
        if execution_id is not None:
            await crud.update_sales_execution_status(
                db=db, execution_id=execution_id, estado=ExecutionStatus.FAILED,
                mensaje_error=msg_ext, detalles=msg_ext
            )
        return {"success": False, "message": msg_ext, "registros_insertados": 0, "registros_duplicados": 0, "total_registros": 0, "errores": []}
    try:
        if extension == ".xlsx":
            df = pd.read_excel(path_ventas, engine="openpyxl", header=None)
        else:
            df = pd.read_excel(path_ventas, engine="xlrd", header=None)
    except Exception as e:
        msg = f"Error al leer el archivo Excel: {str(e)}"
        if execution_id is not None:
            await crud.update_sales_execution_status(db=db, execution_id=execution_id, estado=ExecutionStatus.FAILED, mensaje_error=msg, detalles=msg)
        return {"success": False, "message": msg, "registros_insertados": 0, "registros_duplicados": 0, "total_registros": 0, "errores": []}
    if len(df) < 3:
        msg = f"El archivo debe tener al menos 3 filas. Tiene {len(df)}."
        if execution_id is not None:
            await crud.update_sales_execution_status(db=db, execution_id=execution_id, estado=ExecutionStatus.FAILED, mensaje_error=msg, detalles=msg)
        return {"success": False, "message": msg, "registros_insertados": 0, "registros_duplicados": 0, "total_registros": 0, "errores": []}
    try:
        ahora_utc = datetime.now(tz.utc)
        resultado_dict = await _procesar_ventas_core(db, df)
        # Actualizar precios de venta desde ventas (misma lógica que el botón Actualizar en Precios de Venta)
        try:
            resultado_pv = await crud.sincronizar_precios_venta_desde_ventas(db=db, user_id=user_id)
            resultado_dict["precios_venta_actualizados"] = resultado_pv.get("actualizados", 0)
            resultado_dict["precios_venta_sin_cambios"] = resultado_pv.get("sin_cambios", 0)
            resultado_dict["precios_venta_sin_venta"] = resultado_pv.get("sin_venta", 0)
            msg_pv = (
                f" Precios de venta: {resultado_pv.get('actualizados', 0)} actualizados, "
                f"{resultado_pv.get('sin_cambios', 0)} sin cambios, "
                f"{resultado_pv.get('sin_venta', 0)} sin venta aplicable."
            )
            resultado_dict["message"] = (resultado_dict.get("message") or "") + msg_pv
        except Exception as e_pv:
            resultado_dict["precios_venta_error"] = str(e_pv)
            resultado_dict["message"] = (resultado_dict.get("message") or "") + f" Actualización precios de venta falló: {e_pv}."
        fin_utc = datetime.now(tz.utc)
        duracion = int((fin_utc - ahora_utc).total_seconds())
        await crud.update_sales_execution_status(
            db=db,
            execution_id=execution_id,
            estado=ExecutionStatus.SUCCESS,
            fecha_inicio_ejecucion=ahora_utc,
            fecha_fin_ejecucion=fin_utc,
            duracion_segundos=duracion,
            mensaje_error=resultado_dict["message"],
            detalles=resultado_dict["message"],
            archivo_nombre=path_ventas.name,
        )
        return resultado_dict
    except Exception as e:
        import traceback
        msg = f"Error al procesar ventas: {str(e)}"
        if execution_id is not None:
            await crud.update_sales_execution_status(
                db=db, execution_id=execution_id, estado=ExecutionStatus.FAILED,
                mensaje_error=msg, detalles=msg, stack_trace=traceback.format_exc()
            )
        return {"success": False, "message": msg, "registros_insertados": 0, "registros_duplicados": 0, "total_registros": 0, "errores": []}


@app.post("/api/ventas/procesar-archivos")
async def procesar_archivos_ventas(
    request: Request,
    archivo_ventas: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Procesa un archivo Excel (Reporte de ventas) y sube los datos a la base de datos."""
    import tempfile
    import shutil
    import os
    from pathlib import Path
    import pandas as pd
    import traceback
    
    execution_id = None
    ahora_utc = None
    try:
        # Validar que sea un archivo Excel
        extensiones_permitidas = ['.xlsx', '.xls']
        
        nombre_archivo_ventas = archivo_ventas.filename or ''
        
        extension_ventas = Path(nombre_archivo_ventas).suffix.lower()
        
        if extension_ventas not in extensiones_permitidas:
            return JSONResponse(
                status_code=400,
                content={
                    "success": False,
                    "error": f"El archivo de Reporte de ventas debe ser un archivo Excel (.xlsx o .xls). El archivo proporcionado tiene la extensión: {extension_ventas}"
                }
            )
        
        
        # Crear directorio temporal para el archivo
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            
            # Guardar archivo temporalmente
            archivo_ventas_path = temp_path / f"ventas_{nombre_archivo_ventas}"
            
            with open(archivo_ventas_path, "wb") as f:
                shutil.copyfileobj(archivo_ventas.file, f)
            
            result = await _procesar_archivo_ventas_desde_ruta(db, archivo_ventas_path, None, current_user.id)
            return JSONResponse(
                status_code=200 if result["success"] else 207,
                content=result
            )
        
        
    except Exception as e:
        import traceback
        error_traceback = traceback.format_exc()
        error_message = str(e)
        
        # Si se creó el registro de ejecución, actualizarlo a FAILED
        try:
            if execution_id is not None:
                await crud.update_sales_execution_status(
                    db=db,
                    execution_id=execution_id,
                    estado=ExecutionStatus.FAILED,
                    mensaje_error=str(e),
                    detalles=str(e),
                    stack_trace=error_traceback
                )
        except Exception:
            pass  # No fallar si no se puede actualizar el historial
        
        # Mensaje de error más descriptivo en español
        mensaje_error = f"Error al procesar el archivo de ventas: {error_message}"
        
        # Mensajes de error más descriptivos en español según el tipo de error
        error_lower = error_message.lower()
        
        if "duplicate" in error_lower or "unique" in error_lower or "ya existe" in error_lower:
            mensaje_error = "Error: Se intentó insertar un registro duplicado. El sistema ya detectó y omitió los duplicados, pero puede haber un problema con la verificación. Por favor, verifica que no existan registros con la misma combinación de código cliente, producto y período."
        elif "foreign key" in error_lower or "constraint" in error_lower or "violates foreign key" in error_lower:
            mensaje_error = "Error: Violación de restricción de base de datos. Verifica que los datos de referencia (grupo de cliente) existan en la base de datos antes de insertar las ventas."
        elif ("null" in error_lower and "not null" in error_lower) or "required" in error_lower:
            mensaje_error = "Error: Se intentó insertar un registro con campos requeridos vacíos. Por favor, verifica que todos los campos obligatorios tengan valores válidos en el archivo Excel."
        elif "connection" in error_lower or "timeout" in error_lower or "database" in error_lower:
            mensaje_error = "Error: No se pudo conectar con la base de datos. Por favor, verifica la conexión e intenta nuevamente."
        elif "value" in error_lower and "type" in error_lower:
            mensaje_error = "Error: Tipo de dato incorrecto. Verifica que los valores en el archivo Excel sean del tipo correcto (números, fechas, texto)."
        elif "periodo" in error_lower or "date" in error_lower:
            mensaje_error = "Error al procesar el período. Verifica que el formato del período en el archivo Excel sea correcto (mes y año)."
        else:
            mensaje_error = f"Error al procesar el archivo de ventas: {error_message}"
        
        # Asegurar que siempre haya un mensaje de error en español
        if not mensaje_error or mensaje_error.strip() == "":
            mensaje_error = f"Error inesperado al procesar el archivo de ventas. Tipo de error: {type(e).__name__}"
        
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "error": mensaje_error,
                "detalle_tecnico": error_message,
                "tipo_error": type(e).__name__,
                "traceback": error_traceback
            }
        )


@app.get("/api/select-folder")
async def select_folder(request: Request, current_user: User = Depends(get_current_user)):
    """Abre un diálogo nativo para seleccionar una carpeta."""
    import os
    import platform
    import subprocess
    
    try:
        system = platform.system()
        folder_path = None
        
        if system == "Darwin":  # macOS
            # Usar AppleScript para abrir el diálogo nativo de macOS
            script = '''
            tell application "System Events"
                activate
                set folderPath to choose folder with prompt "Seleccionar carpeta de salida"
                return POSIX path of folderPath
            end tell
            '''
            try:
                result = subprocess.run(
                    ["osascript", "-e", script],
                    capture_output=True,
                    text=True,
                    timeout=60
                )
                if result.returncode == 0:
                    folder_path = result.stdout.strip()
                    # Convertir ruta POSIX a formato estándar
                    if folder_path and not folder_path.endswith(os.sep):
                        folder_path += os.sep
            except subprocess.TimeoutExpired:
                return JSONResponse(
                    status_code=400,
                    content={"error": "Tiempo de espera agotado"}
                )
            except Exception as e:
                return JSONResponse(
                    status_code=500,
                    content={"error": f"Error al abrir el diálogo: {str(e)}"}
                )
        
        elif system == "Windows":
            # Usar tkinter para Windows (puede ejecutarse en hilo)
            try:
                import tkinter as tk
                from tkinter import filedialog
                
                selected_folder = [None]
                
                def open_dialog():
                    try:
                        root = tk.Tk()
                        root.withdraw()
                        root.attributes('-topmost', True)
                        folder = filedialog.askdirectory(
                            title="Seleccionar carpeta de salida",
                            mustexist=True
                        )
                        selected_folder[0] = folder
                        root.destroy()
                    except Exception as e:
                        selected_folder[0] = f"Error: {str(e)}"
                
                thread = threading.Thread(target=open_dialog)
                thread.daemon = True
                thread.start()
                thread.join(timeout=60)
                
                if selected_folder[0] and not selected_folder[0].startswith("Error"):
                    folder_path = selected_folder[0]
                elif selected_folder[0] and selected_folder[0].startswith("Error"):
                    return JSONResponse(
                        status_code=500,
                        content={"error": selected_folder[0]}
                    )
            except ImportError:
                return JSONResponse(
                    status_code=500,
                    content={"error": "tkinter no está disponible"}
                )
        
        else:  # Linux y otros
            # Usar tkinter para Linux
            try:
                import tkinter as tk
                from tkinter import filedialog
                
                selected_folder = [None]
                
                def open_dialog():
                    try:
                        root = tk.Tk()
                        root.withdraw()
                        folder = filedialog.askdirectory(
                            title="Seleccionar carpeta de salida",
                            mustexist=True
                        )
                        selected_folder[0] = folder
                        root.destroy()
                    except Exception as e:
                        selected_folder[0] = f"Error: {str(e)}"
                
                thread = threading.Thread(target=open_dialog)
                thread.daemon = True
                thread.start()
                thread.join(timeout=60)
                
                if selected_folder[0] and not selected_folder[0].startswith("Error"):
                    folder_path = selected_folder[0]
                elif selected_folder[0] and selected_folder[0].startswith("Error"):
                    return JSONResponse(
                        status_code=500,
                        content={"error": selected_folder[0]}
                    )
            except ImportError:
                return JSONResponse(
                    status_code=500,
                    content={"error": "tkinter no está disponible"}
                )
        
        if folder_path:
            # Asegurar que termine con el separador correcto
            if not folder_path.endswith(os.sep):
                folder_path += os.sep
            
            return JSONResponse(
                content={"folder_path": folder_path}
            )
        else:
            return JSONResponse(
                status_code=400,
                content={"error": "No se seleccionó ninguna carpeta"}
            )
            
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": f"Error inesperado: {str(e)}"}
        )


@app.get("/admin")
async def admin(request: Request, current_user: User = Depends(get_current_user), db: AsyncSession = Depends(get_db)):
    """Panel de administración - requiere rol admin."""
    if current_user.rol != "admin":
        return RedirectResponse(url="/dashboard", status_code=302)
    
    from sqlalchemy import select
    from app.db.models import User
    
    # Obtener todos los usuarios
    result = await db.execute(select(User).order_by(User.created_at.desc()))
    users = result.scalars().all()
    
    return templates.TemplateResponse(
        "admin.html",
        {
            "request": request,
            "current_user": current_user,
            "users": users,
            "active_page": "admin"
        }
    )


@app.post("/admin/users/create")
async def admin_create_user(
    request: Request,
    email: str = Form(...),
    password: str = Form(...),
    nombre: Optional[str] = Form(None),
    rol: str = Form("operador"),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Crear nuevo usuario - solo admin."""
    if current_user.rol != "admin":
        return RedirectResponse(url="/dashboard", status_code=302)
    
    from app.db import crud
    
    # Validaciones
    if len(password) < 8:
        return RedirectResponse(
            url=f"/admin?error=El password debe tener al menos 8 caracteres",
            status_code=302
        )
    
    if len(password.encode('utf-8')) > 72:
        return RedirectResponse(
            url=f"/admin?error=El password es demasiado largo (máximo 72 bytes)",
            status_code=302
        )
    
    # Validar rol
    if rol not in ["admin", "operador", "auditor"]:
        rol = "operador"
    
    # Verificar si el email ya existe
    existing_user = await crud.get_user_by_email(db, email)
    if existing_user:
        return RedirectResponse(
            url=f"/admin?error=El email ya está registrado",
            status_code=302
        )
    
    try:
        new_user = await crud.create_user(
            db=db,
            email=email,
            password=password,
            nombre=nombre,
            rol=rol
        )
        return RedirectResponse(url="/admin?success=Usuario creado exitosamente", status_code=302)
    except Exception as e:
        return RedirectResponse(
            url=f"/admin?error=Error al crear usuario: {str(e)}",
            status_code=302
        )


@app.post("/admin/users/{user_id}/update-role")
async def admin_update_user_role(
    user_id: int,
    request: Request,
    rol: str = Form(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Actualizar rol de usuario - solo admin."""
    if current_user.rol != "admin":
        return RedirectResponse(url="/dashboard", status_code=302)
    
    from app.db import crud
    
    # Validar rol
    if rol not in ["admin", "operador", "auditor"]:
        return RedirectResponse(url="/admin?error=Rol inválido", status_code=302)
    
    try:
        await crud.update_user_role(db, user_id, rol)
        return RedirectResponse(url="/admin?success=Rol actualizado exitosamente", status_code=302)
    except Exception as e:
        return RedirectResponse(
            url=f"/admin?error=Error al actualizar rol: {str(e)}",
            status_code=302
        )


@app.post("/admin/users/{user_id}/update")
async def admin_update_user(
    user_id: int,
    request: Request,
    email: str = Form(...),
    nombre: Optional[str] = Form(None),
    rol: str = Form("operador"),
    password: Optional[str] = Form(None),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Actualizar usuario - solo admin."""
    if current_user.rol != "admin":
        return RedirectResponse(url="/dashboard", status_code=302)
    
    from app.db import crud
    
    # Validar rol
    if rol not in ["admin", "operador", "auditor"]:
        return RedirectResponse(url="/admin?error=Rol inválido", status_code=302)
    
    # Validar contraseña si se proporciona
    if password:
        if len(password) < 8:
            return RedirectResponse(
                url=f"/admin?error=El password debe tener al menos 8 caracteres",
                status_code=302
            )
        
        if len(password.encode('utf-8')) > 72:
            return RedirectResponse(
                url=f"/admin?error=El password es demasiado largo (máximo 72 bytes)",
                status_code=302
            )
    
    try:
        updated_user = await crud.update_user(
            db=db,
            user_id=user_id,
            email=email,
            nombre=nombre,
            rol=rol,
            password=password if password else None
        )
        
        if not updated_user:
            return RedirectResponse(
                url=f"/admin?error=Usuario no encontrado",
                status_code=302
            )
        
        return RedirectResponse(url="/admin?success=Usuario actualizado exitosamente", status_code=302)
    except ValueError as e:
        return RedirectResponse(
            url=f"/admin?error={str(e)}",
            status_code=302
        )
    except Exception as e:
        return RedirectResponse(
            url=f"/admin?error=Error al actualizar usuario: {str(e)}",
            status_code=302
        )


@app.get("/hello/{name}")
async def say_hello(name: str):
    """Endpoint de ejemplo - no requiere autenticación."""
    return {"message": f"Hello {name}"}
